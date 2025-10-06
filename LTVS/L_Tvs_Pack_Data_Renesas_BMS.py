# -*- coding: utf-8 -*-
"""
Created on Sat May 18 14:30:40 2024

@author: IITM
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import docx
import matplotlib.colors as mcolors
from cycler import cycler
import datetime
import os
import glob

#%% Get all the files with names matching the naming pattern of the files to be processed
files=glob.glob(r"D:\Benisha\LTVS\10kWh\Renesas_BMS\Drive_Test\15-04-2024\New folder\15-4-24 2  tvs 10kwh pack drive testteraterm 1.xlsx")

#%%
# Specify tentative column names
# cols=['CV_CT_PC_tag','Cell 1','Cell 2','Cell 3','Cell 4','Cell 5','Cell 6','Cell 7','Cell 8','Cell 9','Cell 10','Cell 11','Cell 12','Cell 13','Cell 14','Cell 15','Cell 16','BR_ID_tag','BR_ID_th','tm','ts']
cols=['Date','Hours','Minutes','Seconds','CV_CT_PC_tag','Cell 1','Cell 2','Cell 3','Cell 4','Cell 5','Cell 6','Cell 7','Cell 8','Cell 9','Cell 10','Cell 11','Cell 12','Cell 13','Cell 14','Cell 15','Cell 16','BR_ID_tag','BR_ID_th','tm','ts']

# Mention sheet name if the input putty log has more than one sheet
# "sheet" should have the raw putty log
# sheet="log2_raw" # Uncomment if needed

#%%
# Mention number of cells in pack
v_no=17
t_no=17

#%% Loop through all the files
for file in files:
    # Read file using read_excel() function of the pandas library
    data=pd.read_excel(file,names=cols,index_col=False) #,sheet_name=sheet # include this in paranthesis if needed
    
    # Get destination file path and name of the file from the input file path
    img_path=file.rsplit('\\',1)[0]+'\\'+file.rsplit('\\',1)[1].rsplit('.',1)[0]
    filename=file.rsplit('\\',1)[1].rsplit('.',1)[0]
    
    # If the destination folder path does not exist, create it
    if not os.path.exists(img_path):
        os.makedirs(img_path)
    
    #%% Create MS Word document object using docx library
    document=docx.Document()
    document.styles['Normal'].font.name='Calibri'
    H1=document.add_heading()
    H1.alignment=1
    H1.add_run('LTVS Pack Data').font.size=docx.shared.Pt(18)
    document.add_paragraph().add_run('\n\nPack Specifications:').bold=True
    
    # Create a table in the word document containing the pack specification details
    table1=document.add_table(rows=5,cols=2)
    
    table1.cell(0,0).text='Chemistry'
    table1.cell(0,1).text='LFP'
    table1.cell(1,0).text='Pack Capacity'
    table1.cell(1,1).text='210Ah (2*105Ah)'
    table1.cell(2,0).text='Configuration'
    table1.cell(2,1).text='2P16S'
    table1.cell(3,0).text='UV cut-off'
    table1.cell(3,1).text='2.8V'
    table1.cell(4,0).text='OV cut-off'
    table1.cell(4,1).text='3.5V'
    
    table1.style='Table Grid'
    
    #%% Clean and process date and timestamp of dataframe
    data[['S','mS']]=data['Seconds'].str.split('.',expand=True)
    data['Date']=data['Date'].str.replace('[','',regex=True)
    data['DateTime']=np.nan
    data['DateTime']=data['Date']+' '+data['Hours'].astype(str)+':'+data['Minutes'].astype(str)+':'+data['S']
    data=data.drop(['S','mS'],axis=1)
    
    data=data.replace('|','',regex=True)
    
    # Since the data is usually huge, a portion of the data is sliced
    # This is to view the structure of the data logging, which will come in handy in the following section    
    d1=data.head(1000)
    
    data['DateTime']=pd.to_datetime(data['DateTime'],format='%Y-%m-%d %H:%M:%S',errors='coerce')
    
    #%% Separate dataset into separate dataframes
    
    # This is done as the data is in putty log format, which does not have data separated into proper columns
    # Note the values to which the "CV_CT_PC_tag" is checked against (eg, '|CV_1[1..16](mV)', '|Min_temp')
    # These values need to be checked if they are right, using the 'd1' dataframe in the previous section
    # Change the values if they are different in the input data file
    
    voltage_data=data[data['CV_CT_PC_tag']=='|CV_1[1..16](mV)'].dropna(how='all',axis=1).reset_index(drop=True)
    temperature_data=data[data['CV_CT_PC_tag']=='|Cell_Temperatures[1..16]'].dropna(how='all',axis=1).reset_index(drop=True)
    max_min_temp_data=data[data['CV_CT_PC_tag']=='|Min_temp'].dropna(how='all',axis=1).reset_index(drop=True)
    max_min_V_data=data[data['CV_CT_PC_tag']=='|Max'].dropna(how='all',axis=1).reset_index(drop=True)
    PV_PC_data=data[data['CV_CT_PC_tag']=='|PV(mV)'].dropna(how='all',axis=1).reset_index(drop=True)
    misc_details=data[data['CV_CT_PC_tag']=='|Bal'].dropna(how='all',axis=1).reset_index(drop=True)
    energy_details=data[data['CV_CT_PC_tag']=='|Pack_Energy'].dropna(how='all',axis=1).reset_index(drop=True)
    relay_details=data[data['CV_CT_PC_tag']=='|RLY'].dropna(how='all',axis=1).reset_index(drop=True)
    
    #%% Clean PV and PP data for processing <<>>
    PV_PC_data[['PV','junk']]=PV_PC_data['Cell 1'].str.split('|',expand=True)
    PV_PC_data[['PP','junk']]=PV_PC_data['Cell 2'].str.split('|',expand=True)
    
    #%% Column names for the voltage_data and temperature_data dataframes are properly defined
    
    # "junk" is added as the first element to represent the column 'CV_CT_PC_tag'
    # v_names/t_names is a copy as the list with only the cell number names will be used later to name the legend while plotting
    v_names_1=[f'V{x}' for x in range(1,v_no)] 
    t_names_1=[f'T{x}' for x in range(1,t_no)]
    
    #%% If the dataframe does not have date and time, uncomment this block
    # mm_temp=['junk','Min_T [Cell_No]','junk','junk','Max_T [Cell_No]','junk']
    # mm_v=['junk','junk','Max_Cell_V_BMS(mV)','Max_V_Cell_No_BMS','junk','junk','junk','Min_Cell_V_BMS(mV)','Min_V_Cell_No_BMS','junk','junk','deltaV(mV)_BMS']
    # pv=['junk','junk','junk','Pack_Current','Pack_Voltage','junk','Pack_Power(W)']
    # misc=['junk','Balancing_Status','junk','junk','Afe_1','junk','junk','Fault_status','junk','junk','junk','Latch_counter','junk']
    # engy=['junk','Pack_Energy_BMS','junk','junk','Pack_Ah_BMS','junk','junk','SoC(%)','junk','junk','SoH(%)','junk','junk','Cycles','junk','junk','Premature_Cycles']
    # rly=['junk','junk','junk','Main_Relay_Status','junk','Precharge','junk','junk','Pre-discharge','junk']
    
    #%% DT
    # mm_temp=['junk','junk','junk','junk','junk','Min_T [Cell_No]','junk','junk','Max_T [Cell_No]','junk','DateTime']
    mm_temp=['junk','junk','junk','junk','junk','Min_T [Cell_No]','junk','junk','Max_T [Cell_No]','junk','Flash Usage(%)','DateTime']
    # mm_v=['junk','junk','junk','junk','junk','junk','Max_Cell_V_BMS(mV)','Max_V_Cell_No_BMS','junk','junk','junk','Min_Cell_V_BMS(mV)','Min_V_Cell_No_BMS','junk','junk','deltaV(mV)_BMS','DateTime']
    # mm_v=['junk','junk','junk','junk','junk','junk','Max_Cell_V_BMS(mV)','Max_V_Cell_No_BMS','junk','junk','junk','Min_Cell_V_BMS(mV)[Cell_No]','junk','junk','deltaV(mV)_BMS','junk','DateTime'] # 12/04/2024
    mm_v=['junk','junk','junk','junk','junk','junk','Max_Cell_V_BMS(mV)','Max_V_Cell_No_BMS','junk','junk','junk','Min_Cell_V_BMS(mV)','Min_V_Cell_No_BMS','junk','deltaV(mV)_BMS','DateTime']
    pv=['junk','junk','junk','junk','junk','junk','junk','Pack_Current','DateTime','Pack_Voltage','junk','Pack_Power(W)']
    # pv=['junk','junk','junk','junk','junk','junk','junk','Pack_Current','junk','DateTime','Pack_Voltage','junk','Pack_Power(W)'] # 12/04/2024
    # misc=['junk','junk','junk','junk','junk','Balancing_Status','junk','junk','Afe_1','junk','junk','Fault_status','junk','junk','junk','Latch_counter','junk','DateTime']
    misc=['junk','junk','junk','junk','junk','Balancing_Status','junk','junk','Afe_1','junk','DateTime']
    # engy=['junk','junk','junk','junk','junk','Pack_Energy_BMS','junk','junk','Pack_Ah_BMS','junk','junk','SoC(%)','junk','junk','SoH(%)','junk','junk','Cycles','junk','junk','Premature_Cycles','DateTime']
    engy=['junk','junk','junk','junk','junk','Pack_Energy_BMS','junk','junk','Pack_Ah_BMS','junk','SoC(%)','junk','junk','SoH(%)','junk','Cycles','junk','junk','Premature_Cycles','DateTime']
    rly=['junk','junk','junk','junk','junk','junk','junk','Main_Relay_Status','junk','Precharge','junk','junk','Pre-discharge','junk','DateTime']
    
    
    #%%
    
    # v_names=v_names_1.copy()
    # v_names.insert(0,'junk')
    # t_names=t_names_1.copy()
    # t_names.insert(0,'junk')
    # t_names.append('junk')
    
    #%% DT
    
    v_names=v_names_1.copy()
    v_names.insert(0,'junk')
    v_names.insert(0,'junk')
    v_names.insert(0,'junk')
    v_names.insert(0,'junk')
    v_names.insert(0,'junk')
    v_names.append('DateTime')
    t_names=t_names_1.copy()
    t_names.append('junk')
    t_names.insert(0,'junk')
    t_names.insert(0,'junk')
    t_names.insert(0,'junk')
    t_names.insert(0,'junk')
    t_names.insert(0,'junk')
    t_names.append('DateTime')
    
    #%%
    voltage_data.columns=v_names
    temperature_data.columns=t_names
    max_min_temp_data.columns=mm_temp
    max_min_V_data.columns=mm_v 
    PV_PC_data.columns=pv
    misc_details.columns=misc
    energy_details.columns=engy
    relay_details.columns=rly
    
    #%%
    voltage_data=voltage_data.drop([col for col in voltage_data.columns if "junk" in col], axis=1)
    temperature_data=temperature_data.drop([col for col in temperature_data.columns if "junk" in col], axis=1)
    
    max_min_temp_data=max_min_temp_data.drop([col for col in max_min_temp_data.columns if "junk" in col], axis=1)
    max_min_V_data=max_min_V_data.drop([col for col in max_min_V_data.columns if "junk" in col], axis=1)
    PV_PC_data=PV_PC_data.drop([col for col in PV_PC_data.columns if "junk" in col], axis=1)
    misc_details=misc_details.drop([col for col in misc_details.columns if "junk" in col], axis=1)
    energy_details=energy_details.drop([col for col in energy_details.columns if "junk" in col], axis=1)
    relay_details=relay_details.drop([col for col in relay_details.columns if "junk" in col], axis=1)
    
    #%%
    voltage_data['V'+str(v_no-1)]=voltage_data['V'+str(v_no-1)].str.replace('|','',regex=True)
    PV_PC_data['Pack_Current']=PV_PC_data['Pack_Current'].str.replace('|','',regex=True)
    # PV_PC_data['Pack_Current']=PV_PC_data['Pack_Current'].astype(str).str.extract(r'(\-{0,}\d{1,})')
    
    for x in v_names_1: 
        voltage_data[x]=voltage_data[x].astype(str).str.extract(r'(\d{1,})')
        voltage_data[x]=voltage_data[x].fillna(method='ffill')
        voltage_data[x]=voltage_data[x].astype(int)
        voltage_data=voltage_data[(voltage_data[x]>=2000) & (voltage_data[x]<=4300)]
    
    for x in t_names_1: 
        temperature_data[x]=temperature_data[x].astype(str).str.extract(r'(\d{1,})')
    #%% 
    
    # voltage_data.loc[:,'V1':'V'+str(v_no-1)]=voltage_data.loc[:,'V1':'V'+str(v_no-1)].fillna(method='ffill')
    # voltage_data.loc[:,'V1':'V'+str(v_no-1)]=voltage_data.loc[:,'V1':'V'+str(v_no-1)].astype(int)
    voltage_data['Mean_V']=voltage_data.loc[:,'V1':'V'+str(v_no-1)].mean(axis=1)
    voltage_data['Max_V']=voltage_data.loc[:,'V1':'V'+str(v_no-1)].max(axis=1)
    voltage_data['Max_V_Cell']=voltage_data.loc[:,'V1':'V'+str(v_no-1)].astype(int).idxmax(axis=1)
    voltage_data['Min_V']=voltage_data.loc[:,'V1':'V'+str(v_no-1)].min(axis=1)
    voltage_data['Min_V_Cell']=voltage_data.loc[:,'V1':'V'+str(v_no-1)].astype(int).idxmin(axis=1)
    
    voltage_data['delV']=voltage_data['Max_V']-voltage_data['Min_V']
    
    #%% DT
    # voltage_data['ElapsedTime']=(voltage_data['DateTime']-voltage_data['DateTime'].iloc[0])
    
    #%%
    temperature_data.loc[:,'T1':'T'+str(t_no-1)]=temperature_data.loc[:,'T1':'T'+str(t_no-1)].astype(float)
    temperature_data.loc[:,'T1':'T'+str(t_no-1)]=temperature_data.loc[:,'T1':'T'+str(t_no-1)].fillna(method='ffill')
    temperature_data['Mean_T']=temperature_data.loc[:,'T1':'T'+str(t_no-1)].mean(axis=1)
    temperature_data['Max_T']=temperature_data.loc[:,'T1':'T'+str(t_no-1)].max(axis=1)
    temperature_data['Max_TCell']=temperature_data.loc[:,'T1':'T'+str(t_no-1)].astype(float).idxmax(axis=1)
    temperature_data['Min_T']=temperature_data.loc[:,'T1':'T'+str(t_no-1)].min(axis=1)
    temperature_data['Min_TCell']=temperature_data.loc[:,'T1':'T'+str(t_no-1)].astype(float).idxmin(axis=1)
    
    temperature_data['delT']=temperature_data['Max_T']-temperature_data['Min_T']
    
    #%% DT
    # temperature_data['ElapsedTime']=(temperature_data['DateTime']-temperature_data['DateTime'].iloc[0])
    
    #%% 
    # PV_PC_data['Pack_Current']=PV_PC_data['Pack_Current'].astype(str).str.extract(r'(\-{0,}\d{1,}.\d{1,})')
    # PV_PC_data['Pack_Current']=PV_PC_data['Pack_Current'].astype(float)
    
    #%% DT
    # PV_PC_data['ElapsedTime']=(PV_PC_data['DateTime']-PV_PC_data['DateTime'].iloc[0])
    
    
    #%%
    
    relay_details['Precharge']=relay_details['Precharge'].astype(str).str.extract(r'(\d{1,})')
    relay_details['Main_Relay_Status']=relay_details['Main_Relay_Status'].astype(str).str.extract(r'(\d{1,})')
    relay_details['Pre-discharge']=relay_details['Pre-discharge'].astype(str).str.extract(r'(\d{1,})')
    
    relay_details['Main_Relay_Status']=relay_details['Main_Relay_Status'].fillna(method='ffill')
    relay_details['Main_Relay_Status']=relay_details['Main_Relay_Status'].fillna(method='bfill')
    
    relay_details['Precharge']=relay_details['Precharge'].fillna(method='ffill')
    relay_details['Precharge']=relay_details['Precharge'].fillna(method='bfill')
    
    relay_details['Pre-discharge']=relay_details['Pre-discharge'].fillna(method='ffill')
    relay_details['Pre-discharge']=relay_details['Pre-discharge'].fillna(method='bfill')
    
    
    #%%
    energy_details['Pack_Ah_BMS']=energy_details['Pack_Ah_BMS'].astype(str).str.replace('0e','',regex=True).str.extract(r'(\-{0,}\d{1,}.\d{0,2})')
    energy_details['Pack_Energy_BMS']=energy_details['Pack_Energy_BMS'].astype(str).str.extract(r'(\-{0,}\d{1,}.\d{0,2})')
    
    energy_details['Pack_Ah_BMS']=energy_details['Pack_Ah_BMS'].fillna(method='ffill')
    energy_details['Pack_Energy_BMS']=energy_details['Pack_Energy_BMS'].fillna(method='ffill')
    
    #%%
    # PV_PC_data['Pack_Current']=PV_PC_data['Pack_Current'].astype(str).str.extract(r'(\-{0,}\d{1,}.\d{1,})')
    # PV_PC_data['Pack_Current']=PV_PC_data['Pack_Current'].astype(float)
    PV_PC_data['Pack_Voltage']=PV_PC_data['Pack_Voltage'].astype(float)
    PV_PC_data['Pack_Voltage']=PV_PC_data['Pack_Voltage']*0.001
    
    #%% DT
    details=pd.DataFrame()
    # details=pd.concat([PV_PC_data,energy_details],axis=1,ignore_index=True,keys='DateTime')
    details=pd.merge(left=PV_PC_data,right=energy_details,on='DateTime') # DT # left=PV_PC_data,right=PV_PC_data,

    
    details['Time_in_sec_s']=(details['DateTime']-details['DateTime'].shift(1))/np.timedelta64(1,'s')# time difference b/w every datapt
    details['Time_in_sec']=(details['DateTime']-details['DateTime'].iloc[0])/np.timedelta64(1,'s')
    
    
    #%%
    # details=pd.DataFrame()
    # details=pd.concat([PV_PC_data,energy_details],axis=1,ignore_index=True)
    # p=list(PV_PC_data.columns)
    # r=list(energy_details.columns)
    # p.extend(r)
    # details.columns=p
    # details['Time_in_sec_s']=1
    
    #%%
    details['Pack_Current']=details['Pack_Current'].astype(str).str.extract(r'(\-{0,}\d{1,}\.\d{1,})')
    # details['Pack_Voltage']=details['Pack_Voltage_BMS(mV)'].astype(str).str.extract(r'(\d{1,}.\d{1,})')
    details['Pack_Current']=details['Pack_Current'].astype(float)
    # details['Pack_Voltage']=details['Pack_Voltage'].astype(float)
    # details['Pack_Voltage']=details['Pack_Voltage']*0.001
    
    #%%
    details['State']=np.nan
    details['Time_in_sec_s_cap']=np.where((details['Time_in_sec_s']>300),np.nan,details['Time_in_sec_s'])
    details['State']=np.where(details['Pack_Current']>0,0,details['State']) # 1
    details['State']=np.where(details['Pack_Current']<0,1,details['State']) # 2
    details['State']=np.where(details['Pack_Current']==0,2,details['State'])
    
    #%%
    details['Cap_inst']=details['Time_in_sec_s_cap']*abs(details['Pack_Current'])/3600
    details['Capacity_calculated']=details['Cap_inst'].groupby(details['State']).cumsum()
    details['Capacity_calculated_chg']=(details[details['State']==0]['Cap_inst']).cumsum()
    details['Capacity_calculated_dchg']=(details[details['State']==1]['Cap_inst']).cumsum()
    
    details['Capacity_calculated_chg']=details['Capacity_calculated_chg'].fillna(method='bfill')
    details['Capacity_calculated_dchg']=details['Capacity_calculated_dchg'].fillna(method='bfill')
    
    details['Energy_calculated']=details['Capacity_calculated']*details['Pack_Voltage']
    details['Energy_calculated_chg']=details['Capacity_calculated_chg']*details['Pack_Voltage']
    details['Energy_calculated_dchg']=details['Capacity_calculated_dchg']*details['Pack_Voltage']
    
    #%%
    l=7.5
    b=4.5
    
    #%%
    document.add_page_break()
    document.add_paragraph().add_run('Precharge:').bold=True
    
    plt.figure(figsize=(l,b))
    
    plt.plot(range(len(relay_details)),relay_details['Precharge'])
    
    plt.grid(linestyle='dotted')
    plt.title('Precharge',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Precharge.png',dpi=1200)
    
    document.add_picture(img_path+'\\Precharge.png')
    
    #%%
    document.add_paragraph().add_run('Pre-discharge:').bold=True
    
    plt.figure(figsize=(l,b))
    
    plt.plot(range(len(relay_details)),relay_details['Pre-discharge'])
    
    plt.grid(linestyle='dotted')
    plt.title('Pre-discharge',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Pre-discharge.png',dpi=1200)
    
    document.add_picture(img_path+'\\Pre-discharge.png')
    
    #%%
    document.add_paragraph().add_run('Main Relay').bold=True
    
    plt.figure(figsize=(l,b))
    
    plt.plot(range(len(relay_details)),relay_details['Main_Relay_Status'])
    
    plt.grid(linestyle='dotted')
    plt.title('Main Relay',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Main Relay.png',dpi=1200)
    
    document.add_picture(img_path+'\\Main Relay.png')
    
    #%%
    # document.add_paragraph().add_run('Pack Current:').bold=True
    # plt.figure(figsize=(l,b))
    # plt.plot(range(len(PV_PC_data)),PV_PC_data['Pack_Current'])
    
    # plt.ylabel('Current(A)',fontweight='bold')
    
    # plt.grid(linestyle='dotted')
    # plt.title('Pack Current',fontweight='bold')
    # plt.tight_layout()
    # plt.savefig(img_path+'\\Pack Current.png',dpi=1200)
    
    # document.add_picture(img_path+'\\Pack Current.png')
    
    #%%
    document.add_paragraph().add_run('Pack Capacity (BMS):').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(range(len(energy_details)),energy_details['Pack_Ah_BMS'])
    
    plt.ylabel('Capacity(Ah)',fontweight='bold')
    
    plt.grid(linestyle='dotted')
    plt.title('Pack Capacity (BMS)',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Pack Capacity_BMS.png',dpi=1200)
     
    document.add_picture(img_path+'\\Pack Capacity_BMS.png')
    document.add_paragraph('In this session, the capacity expended by the pack is about Ah ()' )
    
    #%%
    document.add_paragraph().add_run('Pack Energy (BMS):').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(range(len(energy_details)),energy_details['Pack_Energy_BMS'])
    
    plt.ylabel('Energy(Wh)',fontweight='bold')
    
    plt.grid(linestyle='dotted')
    plt.title('Pack Energy (BMS)',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Pack Energy_BMS.png',dpi=1200)
    
    document.add_picture(img_path+'\\Pack Energy_BMS.png')
    p_bms_d_engy=document.add_paragraph('The ')
    p_bms_d_engy.add_run('discharging pack energy ').bold=True
    p_bms_d_engy.add_run('is about ')
    p_bms_d_engy.add_run('Wh.').bold=True
    
    #%%
    document.add_paragraph().add_run('Pack Voltage:').bold=True
    
    plt.figure(figsize=(l,b))
    # plt.plot(range(len(PV_PC_data)),PV_PC_data['Pack_Voltage_BMS(mV)'])
    plt.plot(range(len(details)),details['Pack_Voltage']) # DT
    
    plt.ylabel('Voltage(mV)',fontweight='bold')
    
    plt.grid(linestyle='dotted')
    plt.title('Pack Voltage',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Pack Voltage.png',dpi=1200)
    
    #%%
    document.add_picture(img_path+'\\Pack Voltage.png')
    pv=document.add_paragraph('At ')
    pv.add_run('discharge, ').bold=True
    pv.add_run('during the first cycle, the pack goes from to . The pack voltage at rest after discharge is . At ')
    pv.add_run('charge, ').bold=True
    pv.add_run('the pack goes from to . ')
    pv.add_run('In the end, the voltage stabilizes at V during the end rest.')
    
    #%%
    document.add_paragraph().add_run('Cell Voltage:').bold=True
    
    #%%
    plt.figure(figsize=(l,b))
    plt.rc('axes',prop_cycle=cycler('color',list(set(mcolors.XKCD_COLORS.values()))))
    plt.plot(range(len(voltage_data)),voltage_data.loc[:,'V1':'V'+str(v_no-1)])
    
    plt.ylabel('Voltage(mV)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Cell Voltage',fontweight='bold')
    plt.legend(v_names_1)
    plt.tight_layout()
    plt.savefig(img_path+'\\Cell Voltage.png',dpi=1200)
    
    #%%
    document.add_paragraph().add_run('Average Voltage:').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(range(len(voltage_data)),(voltage_data['Mean_V'])) #*0.0001
    
    plt.ylabel('Voltage(mV)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Average Voltage',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Average Voltage.png',dpi=1200)
    
    #%%
    document.add_picture(img_path+'\\Average Voltage.png')
    av=document.add_paragraph('The average cell voltage during ')
    av.add_run('discharge ').bold=True
    av.add_run('goes from ')
    av.add_run('mV to mV. ').bold=True
    av.add_run('The average cell voltage at the rest after discharge is mV. At ')
    av.add_run('charge').bold=True
    av.add_run(', the average cell voltage goes from ')
    av.add_run('mV to mV.').bold=True
    av.add_run('In the end rest, it  stabilizes at %smV.' % (voltage_data['Mean_V'].iloc[-1]))
    
    #%%
    # document.add_paragraph().add_run('Calculated Voltage Difference (delV):').bold=True
    # plt.figure(figsize=(l,b))
    # plt.plot(range(len(voltage_data)),(voltage_data['delV']))
    
    # plt.ylabel('Voltage(mV)',fontweight='bold')
    # plt.grid(linestyle='dotted')
    # plt.title('Voltage difference (delV)',fontweight='bold')
    # plt.tight_layout()
    # plt.savefig(img_path+'\\delV.png',dpi=1200)
    
    # document.add_picture(img_path+'\\delV.png')
    # dv=document.add_paragraph('During discharge, the ')
    # dv.add_run('maximum delV (difference between the highest and lowest cell voltage), ').bold=True
    # dv.add_run('is about ')
    # dv.add_run('mV, at the end of discharge. ').bold=True
    # dv.add_run('At the rest after discharging, the delmV is mV. In the charging session, the maximum delV is mV. ')
    # dv.add_run('The pack stabilizes at %smV delV at the end rest. '% (voltage_data['delV'].iloc[-1]))
    
    #%%
    max_min_V_data['deltaV(mV)_BMS']=max_min_V_data['deltaV(mV)_BMS'].astype(str).str.extract(r'(\d{1,})')
    max_min_V_data['deltaV(mV)_BMS']=max_min_V_data['deltaV(mV)_BMS'].fillna(method='ffill')
    document.add_paragraph().add_run('BMS Voltage Difference (delV):').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(range(len(max_min_V_data)),(max_min_V_data['deltaV(mV)_BMS']))
    
    plt.ylabel('Voltage(mV)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Voltage difference (delV)',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\delV_bms.png',dpi=1200)
    
    document.add_picture(img_path+'\\delV_bms.png')
    dv=document.add_paragraph('During discharge, the ')
    dv.add_run('maximum delV (difference between the highest and lowest cell voltage), ').bold=True
    dv.add_run('is about ')
    dv.add_run('mV, at the end of discharge. ').bold=True
    dv.add_run('At the rest after discharging, the delmV is mV. In the charging session, the maximum delV is mV. ')
    dv.add_run('The pack stabilizes at %smV delV at the end rest. '% (voltage_data['delV'].iloc[-1]))

    #%%
    document.add_paragraph().add_run('Cell Temperature:').bold=True
    
    plt.figure(figsize=(l,b))
    plt.rc('axes',prop_cycle=cycler('color',list(set(mcolors.XKCD_COLORS.values()))))
    plt.plot(range(len(temperature_data)),temperature_data.loc[:,'T1':'T'+str(t_no-1)])
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.legend(t_names_1)
    plt.grid(linestyle='dotted')
    plt.title('Cell Temperature',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Cell Temperature.png',dpi=1200)
    
    # document.add_picture(img_path+'\\Cell Temperature.png')
    ct=document.add_paragraph('The maximum temperature reached is around ')
    ct.add_run('%s'%(temperature_data.loc[:,'T1':'T'+str(t_no-1)].max(axis=1).max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True
    
    #%%
    document.add_paragraph().add_run('Average temperature:').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(range(len(temperature_data)),temperature_data['Mean_T'])
    
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Average Temperature',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Average Temperature.png',dpi=1200)
    
    document.add_picture(img_path+'\\Average Temperature.png')
    at=document.add_paragraph('As can be seen, the ')
    at.add_run('average cell temperature ').bold=True
    at.add_run('stays between ')
    at.add_run(u'\N{DEGREE SIGN}'+'C to '+u'\N{DEGREE SIGN}'+'C ').bold=True
    at.add_run('during ')
    at.add_run('discharging').bold=True
    at.add_run(' and ')
    at.add_run(u'\N{DEGREE SIGN}'+'C to '+u'\N{DEGREE SIGN}'+'C ').bold=True
    at.add_run('during ')
    at.add_run('charging.').bold=True
    
    #%%
    document.add_paragraph().add_run('Temperature Difference (delT):').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(range(len(temperature_data)),temperature_data['delT'])
    
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Temperature difference (delT)',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\delT.png',dpi=1200)
    
    document.add_picture(img_path+'\\delT.png')
    dt=document.add_paragraph('The ')
    dt.add_run('maximum delT ').bold=True
    dt.add_run('is about ')
    dt.add_run('%s' %temperature_data['delT'].max()).bold=True
    dt.add_run(u'\N{DEGREE SIGN}'+'C.').bold=True
    
    #%%
    document.add_paragraph().add_run('Max T:').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(range(len(temperature_data)),temperature_data['Max_T'])
    
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Maximum Temperature',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\MaxT.png',dpi=1200)
    
    document.add_picture(img_path+'\\MaxT.png')
    dt=document.add_paragraph('The ')
    dt.add_run('maximum Temperature ').bold=True
    dt.add_run('is about ')
    dt.add_run(u'\N{DEGREE SIGN}'+'C.').bold=True
    
    #%%
    path=img_path+'\\'+filename+'_modified.xlsx'
    document.save(img_path+'\\'+filename+'_observations_document.docx')
    
    #%%
    writer = pd.ExcelWriter(path, engine = 'xlsxwriter')
    voltage_data.to_excel(writer, sheet_name = 'Cell Voltage',index=False)
    temperature_data.to_excel(writer, sheet_name = 'Cell Temperature',index=False)
    PV_PC_data.to_excel(writer,sheet_name='Pack Voltage Details',index=False)
    relay_details.to_excel(writer,sheet_name='Relay Status',index=False)
    max_min_V_data.to_excel(writer,sheet_name='Voltage Max Min',index=False)
    max_min_temp_data.to_excel(writer,sheet_name='Temp Max Min',index=False)
    misc_details.to_excel(writer,sheet_name='Fault Status',index=False)
    energy_details.to_excel(writer,sheet_name='Energy Details',index=False)
    details.to_excel(writer,sheet_name='Pack Details',index=False) #DT
    writer.close()
    
    #%% DT
    f=open(file.rsplit('.')[0]+'_observations.txt',"w")
    vol_col_check=v_names_1
    temp_col_check=t_names_1
    
    # voltage_data['Time_in_sec_s']=(voltage_data['DateTime']-voltage_data['DateTime'].shift(1))/np.timedelta64(1,'s')# time difference b/w every datapt
    # voltage_data['Time_in_sec']=(voltage_data['DateTime']-voltage_data['DateTime'].iloc[0])/np.timedelta64(1,'s')
    
    print('======================================================================',file=f)
    print('------------------------------------------\n' ,file=f)
    print('Datapoints:%s' % (len(voltage_data)),file=f)
    # print('Timeframe: %s to %s' % (voltage_data['DateTime'].iloc[0],voltage_data['DateTime'].iloc[-1]),file=f)
    print('Total Time Spent:%s' % (datetime.timedelta(seconds=(voltage_data['DateTime'].iloc[-1]-voltage_data['DateTime'].iloc[0])/np.timedelta64(1,'s'))),file=f)
    print('---------------------------',file=f)
    
    if len(details[details['State']==0])!=0:
        c_dt=pd.to_datetime(details[details['State']==0]['DateTime'].tolist())
        
        chg_vol=voltage_data[voltage_data['DateTime'].isin(c_dt)]
        print('Chg Time Spent:%s' % (datetime.timedelta(seconds=(chg_vol['Time_in_sec_s'].sum()))),file=f)
        print('Maximum Charging Current:%.3fA ; Minimum Charging Current:%.3fA' % (details[details['State']==0]['Pack_Current'].max(),details[details['State']==0]['Pack_Current'].min()),file=f)
        print('Chg Pack Voltage range:%sV to %sV' %(details[details['State']==0]['Pack_Voltage'].iloc[0],details[details['State']==0]['Pack_Voltage'].iloc[-1]),file=f)
        print('---------------------------',file=f)
    
        print('Chg Min Vol:%.3fV' %((chg_vol.loc[:,'V1':'V'+str(v_no-1)]*0.001).min(axis=1).min()),file=f)
        print('Chg Max Vol:%.3fV' %((chg_vol.loc[:,'V1':'V'+str(v_no-1)]*0.001).max(axis=1).max()),file=f)
        print('Chg Average Vol:%.3fV' %((chg_vol['Mean_V'])*0.001).mean(),file=f)
        
        print('Chg Minimum delV:%.3fmV' %((chg_vol['delV'])).min(),file=f)
        print('Chg Maximum delV:%.3fmV' %((chg_vol['delV'])).max(),file=f)
        print('Chg Average delV:%.3fmV' %((chg_vol['delV'])).mean(),file=f)
        print('---------------------------',file=f)
     
        
        chg_temp=temperature_data[temperature_data['DateTime'].isin(c_dt)]
        print('Chg Min Temp:%.3f degC' %(chg_temp.loc[:,'T1':'T'+str(t_no-1)].min(axis=1).min()),file=f)
        print('Chg Max Temp:%.3f degC' %(chg_temp.loc[:,'T1':'T'+str(t_no-1)].max(axis=1).max()),file=f)
        print('Chg Average Temp:%.3f degC' %(chg_temp['Mean_T'].mean()),file=f)
        
        print('Chg Minimum delT:%.3f degC' %(chg_temp['delT'].min()),file=f)
        print('Chg Maximum delT:%.3f degC' %(chg_temp['delT'].max()),file=f)
        print('Chg Average delT:%.3f degC' %(chg_temp['delT'].mean()),file=f)
        print('---------------------------',file=f)
        print('Cells with (lower)erroneous chg voltage values: %s' %([col for col in vol_col_check if chg_vol[col].lt(2300).any()]),file=f)
        print('Cells with (higher)erroneous chg voltage values: %s' %([col for col in vol_col_check if chg_vol[col].gt(4300).any()]),file=f)
        
        print('Cells with (lower)erroneous chg temp values: %s' %([col for col in temp_col_check if chg_temp[col].lt(17).any()]),file=f)
        print('Cells with (higher)erroneous chg temp values: %s' %([col for col in temp_col_check if chg_temp[col].gt(35).any()]),file=f)
    print('---------------------------',file=f)
    if len(details[details['State']==1])!=0:
        d_dt=pd.to_datetime(details[details['State']==1]['DateTime'].tolist())
    
        dchg_vol=voltage_data[voltage_data['DateTime'].isin(d_dt)]
        print('Dchg Time Spent:%s' % (datetime.timedelta(seconds=(dchg_vol['Time_in_sec_s'].sum()))),file=f)
        print('Maximum Discharging Current:%.3fA ; Minimum Discharging Current:%.3fA' % (details[details['State']==1]['Pack_Current'].max(),details[details['State']==1]['Pack_Current'].min()),file=f)
        print('DChg Pack Voltage range:%sV to %sV' %(details[details['State']==1]['Pack_Voltage'].iloc[0],details[details['State']==1]['Pack_Voltage'].iloc[-1]),file=f)
        print('---------------------------',file=f)
    
        print('Dchg Min Vol:%.3fV' %((dchg_vol.loc[:,'V1':'V'+str(v_no-1)]*0.001).min(axis=1).min()),file=f)
        print('Dchg Max Vol:%.3fV' %((dchg_vol.loc[:,'V1':'V'+str(v_no-1)]*0.001).max(axis=1).max()),file=f)
        print('Dchg Average Vol:%.3fV' %((dchg_vol['Mean_V'])*0.001).mean(),file=f)
        
        print('Dchg Minimum delV:%.3fmV' %((dchg_vol['delV'])).min(),file=f) 
        print('Dchg Maximum delV:%.3fmV' %((dchg_vol['delV'])).max(),file=f) 
        print('Dchg Average delV:%.3fmV' %((dchg_vol['delV'])).mean(),file=f) 
        print('---------------------------',file=f)
        
        
        dchg_temp=temperature_data[temperature_data['DateTime'].isin(d_dt)]  
        print('Dchg Min Temp:%.3f degC' %(dchg_temp.loc[:,'T1':'T'+str(t_no-1)].min(axis=1).min()),file=f)
        print('Dchg Max Temp:%.3f degC' %(dchg_temp.loc[:,'T1':'T'+str(t_no-1)].max(axis=1).max()),file=f)
        print('Dchg Average Temp:%.3f degC' %(dchg_temp['Mean_T'].mean()),file=f)
        print('Dchg Minimum delT:%.3f degC' %(dchg_temp['delT'].min()),file=f)
        print('Dchg Maximum delT:%.3f degC' %(dchg_temp['delT'].max()),file=f)
        print('Dchg Average delT:%.3f degC' %(dchg_temp['delT'].mean()),file=f)
        print('---------------------------',file=f)
        print('Cells with (lower)erroneous dchg voltage values: %s' %([col for col in vol_col_check if dchg_vol[col].lt(2300).any()]),file=f)
        print('Cells with (higher)erroneous dchg voltage values: %s' %([col for col in vol_col_check if dchg_vol[col].gt(4300).any()]),file=f)
        
        print('Cells with (lower)erroneous dchg temp values: %s' %([col for col in temp_col_check if dchg_temp[col].lt(17).any()]),file=f)
        print('Cells with (higher)erroneous dchg temp values: %s' %([col for col in temp_col_check if dchg_temp[col].gt(35).any()]),file=f)
    print('---------------------------',file=f)
    if len(details[details['State']==2])!=0:
    
        r_dt=pd.to_datetime(details[details['State']==2]['DateTime'].tolist())
        rst_vol=voltage_data[voltage_data['DateTime'].isin(r_dt)]
        print('Rest Time Spent:%s' % (datetime.timedelta(seconds=(rst_vol['Time_in_sec_s'].sum()))),file=f)
        print('Rest Pack Voltage range:%sV to %sV' %(details[details['State']==2]['Pack_Voltage'].iloc[0],details[details['State']==2]['Pack_Voltage'].iloc[-1]),file=f)
        print('---------------------------',file=f)
        print('Rest Min Vol:%.3fV' %((rst_vol.loc[:,'V1':'V'+str(v_no-1)]*0.001).min(axis=1).min()),file=f)
        print('Rest Max Vol:%.3fV' %((rst_vol.loc[:,'V1':'V'+str(v_no-1)]*0.001).max(axis=1).max()),file=f)
        print('Rest Average Vol:%.3fV' %((rst_vol['Mean_V'])*0.001).mean(),file=f)
        
        print('Rest Minimum delV:%.3fmV' %((rst_vol['delV'])).min(),file=f)
        print('Rest Maximum delV:%.3fmV' %((rst_vol['delV'])).max(),file=f) 
        print('Rest Average delV:%.3fmV' %((rst_vol['delV'])).mean(),file=f) 
        print('---------------------------',file=f)
        
        
        rst_temp=temperature_data[temperature_data['DateTime'].isin(r_dt)]
        print('Rest Min Temp:%.3f degC' %(rst_temp.loc[:,'T1':'T'+str(t_no-1)].min(axis=1).min()),file=f)
        print('Rest Max Temp:%.3f degC' %(rst_temp.loc[:,'T1':'T'+str(t_no-1)].max(axis=1).max()),file=f)
        print('Rest Average Temp:%.3f degC' %(rst_temp['Mean_T'].mean()),file=f)
        
        print('Rest Minimum delT:%.3f degC' %(rst_temp['delT'].min()),file=f)
        print('Rest Maximum delT:%.3f degC' %(rst_temp['delT'].max()),file=f)
        print('Rest Average delT:%.3f degC' %(rst_temp['delT'].mean()),file=f)
        print('---------------------------',file=f)
        print('Cells with (lower)erroneous Rest voltage values: %s' %([col for col in vol_col_check if rst_vol[col].lt(2300).any()]),file=f)
        print('Cells with (higher)erroneous Rest voltage values: %s' %([col for col in vol_col_check if rst_vol[col].gt(4300).any()]),file=f)
        
        print('Cells with (lower)erroneous Rest temp values: %s' %([col for col in temp_col_check if rst_temp[col].lt(17).any()]),file=f)
        print('Cells with (higher)erroneous Rest temp values: %s' %([col for col in temp_col_check if rst_temp[col].gt(35).any()]),file=f)
    print('---------------------------',file=f)
    print('Maximum Charging Capacity:%.3fAh ; Minimum Charging Capacity:%.3fAh' % (details['Capacity_calculated_chg'].max(),details['Capacity_calculated_chg'].min()),file=f)
    print('Maximum Discharging Capacity:%.3fAh ; Minimum Discharging Capacity:%.3fAh' % (details['Capacity_calculated_dchg'].max(),details['Capacity_calculated_dchg'].min()),file=f)
    print('Maximum Charging Energy:%.3fkWh ; Minimum Charging Energy:%.3fkWh' % ((details['Energy_calculated_chg']*0.001).max(),(details['Energy_calculated_chg']*0.001).min()),file=f)
    print('Maximum Discharging Energy:%.3fkWh ; Minimum Discharging Energy:%.3fkWh' % ((details['Energy_calculated_dchg']*0.001).max(),(details['Energy_calculated_dchg']*0.001).min()),file=f)
    
    print('Cells with (lower)erroneous voltage values: %s' %([col for col in vol_col_check if voltage_data[col].lt(2300).any()]),file=f)
    print('Cells with (higher)erroneous voltage values: %s' %([col for col in vol_col_check if voltage_data[col].gt(4300).any()]),file=f)
    
    print('Cells with (lower)erroneous temp values: %s' %([col for col in temp_col_check if temperature_data[col].lt(17).any()]),file=f)
    print('Cells with (higher)erroneous temp values: %s' %([col for col in temp_col_check if temperature_data[col].gt(35).any()]),file=f)
    
    print('======================================================================\n',file=f)
    print('\n',file=f)
    f.close()
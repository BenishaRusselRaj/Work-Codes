# -*- coding: utf-8 -*-
"""
Created on Mon Jun 24 11:52:50 2024

@author: IITM
"""
#%% Code Description

"""
USE THIS CODE IF:
    1. Putty log data needs to be processed and report generated
    2. Tester data is available
    3. Putty log HAS date and time    
"""

"""
PRE-REQUISITES:
    1. The putty log need to be imported into MS Excel first
    2. While opening in Excel, open it by mentioning the "comma", "space" and in "other" - ":", as the delimiters
    3. Save the file as an Excel file (i.e., with .xlsx extension)
"""

"""
NOTE:
    1. This code generally needs multiple changes for it to work without errors
    2. It is because the data collection process is highly irregular
    3. Proceed with caution at every cell, otherwise useful data may be lost.
    4. Check and double-check the data after executing every cell

"""

#%% Import necessary libraries
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import docx
import os

#%% 
# Add path to the tester and BMS files
t_f = r"D:\Benisha\LTVS\10kWh\Re_In_house_testing\New_Cycles\Farhan_Testing\06-09-2024\Tester_LTVS 10kWh Pack All Deg testing log.xlsx"
file = r"D:\Benisha\LTVS\10kWh\Re_In_house_testing\New_Cycles\Farhan_Testing\06-09-2024\LTVS_10kWh_pack_current_test_06-09-2024_BMS.xlsx"

step = 'Discharge/Charge'

# Define the tentative column names
# The number of columns must be equal to the number of fields in the input file
cols=['Date','Hours','Minutes','Seconds','CV_CT_PC_tag','Cell 1','Cell 2','Cell 3','Cell 4','Cell 5','Cell 6','Cell 7','Cell 8','Cell 9','Cell 10','Cell 11','Cell 12','Cell 13','Cell 14','Cell 15','Cell 16','BR_ID_tag','BR_ID_th','tm','ts']

#%%
# Number of cells in the pack
v_no = 16
t_no = 16

# Read the excel files using read_excel() function in pandas library
data=pd.read_excel(file, names=cols, index_col=False) #, sheet_name=sheet 

tester_data = pd.read_excel(t_f, sheet_name = 'record')

tester_data['Date'] = pd.to_datetime(tester_data['Date'], errors = 'coerce')

# Get destination folder path and name of the file from the input file path
img_path=file.rsplit('\\',1)[0]+'\\'+file.rsplit('\\',1)[1].rsplit('.',1)[0]
filename=file.rsplit('\\',1)[1].rsplit('.',1)[0]

# If the destination folder does not exist, create the folder
if not os.path.exists(img_path):
    os.makedirs(img_path)

#%% Create a MS Word document using docx library
document=docx.Document()
document.styles['Normal'].font.name='Calibri'
H1=document.add_heading()
H1.alignment=1
H1.add_run('LTVS Pack Data').font.size=docx.shared.Pt(18)

# Add a table in the word document with pack specification details
document.add_paragraph().add_run('\n\nPack Specifications:').bold=True
table1=document.add_table(rows=5,cols=2)

table1.cell(0,0).text='Chemistry'
table1.cell(0,1).text='LFP'
table1.cell(1,0).text='Pack Capacity'
table1.cell(1,1).text='210Ah (2*105Ah)'
table1.cell(2,0).text='Configuration'
table1.cell(2,1).text='2P16S'
table1.cell(3,0).text='UV cut-off'
table1.cell(3,1).text='2.8V'
table1.cell(4,0).text='OV cut-off'
table1.cell(4,1).text='3.5V'

table1.style='Table Grid'

#%% Cleaning and processing of datetime column
data[['S','mS']]=data['Seconds'].str.split('.',expand=True)
data['Date']=data['Date'].str.replace('[','',regex=True)
data['DateTime']=np.nan
data['DateTime']=data['Date']+' '+data['Hours'].astype(str)+':'+data['Minutes'].astype(str)+':'+data['S']
data=data.drop(['S','mS'],axis=1)

data=data.replace('|','',regex=True)

# Since the data is usually huge, a portion of the data is sliced
# This is to view the structure of the data logging, which will come in handy in the next block
d1=data.head(1000)

data['DateTime']=pd.to_datetime(data['DateTime'],format='%Y-%m-%d %H:%M:%S',errors='coerce')

#%% Separate data into separate dataframes

# This is done as the data is in putty log format, which does not have data separated into proper columns
# Note the values to which the "CV_CT_PC_tag" is checked against (eg, '|CV_1[1..16](mV)', '|Min_temp')
# These values need to be checked if they are right, using the 'd1' dataframe in the previous section
# Change the values if they are different in the input data file

voltage_data=data[data['CV_CT_PC_tag']=='|CV_1[1..16](mV)'].dropna(how='all',axis=1).reset_index(drop=True)
temperature_data=data[data['CV_CT_PC_tag']=='|Cell_Temperatures[1..16]'].dropna(how='all',axis=1).reset_index(drop=True)
max_min_temp_data=data[data['CV_CT_PC_tag']=='|Min_temp'].dropna(how='all',axis=1).reset_index(drop=True)
max_min_V_data=data[data['CV_CT_PC_tag']=='|Max'].dropna(how='all',axis=1).reset_index(drop=True)
PV_data=data[data['CV_CT_PC_tag']=='|PV(mV)'].dropna(how='all',axis=1).reset_index(drop=True)
misc_details=data[data['CV_CT_PC_tag']=='|Fault'].dropna(how='all',axis=1).reset_index(drop=True)
energy_details=data[data['CV_CT_PC_tag']=='|Pack_Energy'].dropna(how='all',axis=1).reset_index(drop=True)
relay_details=data[data['CV_CT_PC_tag']=='|RLY'].dropna(how='all',axis=1).reset_index(drop=True)
amb_T = data[data['CV_CT_PC_tag']=='AMB'].dropna(how='all',axis=1).reset_index(drop=True)
shunt_data = data[data['CV_CT_PC_tag']=='END|*|Shunt_instant_value'].dropna(how='all',axis=1).reset_index(drop=True)
gpio_data = data[data['CV_CT_PC_tag']=='*|AFE_GPIO_instant_value'].dropna(how='all',axis=1).reset_index(drop=True)
cc_start = data[data['CV_CT_PC_tag']=='*|CC_start'].dropna(how='all',axis=1).reset_index(drop=True)
k_soc = data[data['CV_CT_PC_tag']=='EKF_SoC'].dropna(how='all',axis=1).reset_index(drop=True)
new_soc = data[data['CV_CT_PC_tag']=='Charging_Ah'].dropna(how='all',axis=1).reset_index(drop=True)

#%% Define Proper column names for each dataframe
    
# Change column names if necessary
# Give the name "junk" to columns with useless data (like Nan or some symbol); It will be removed later
# Changes often with each input dataset as the data logging process is usually irregular

mm_temp=['junk','junk','junk','junk','junk','Min_T [Cell_No]','junk','junk','Max_T [Cell_No]','junk','Flash_Usage','DateTime']
mm_v=['junk','junk','junk','junk','junk','junk','Max_Cell_V_BMS(mV)','Max_V_Cell_No_BMS','junk','junk','junk','Min_Cell_V_BMS(mV)','Min_V_Cell_No_BMS','junk','deltaV(mV)_BMS','DateTime']# CFD Test Data
pv=['junk','junk','junk','junk','junk','Pack_Voltage_PP','PP_PC','Pack_Current','DateTime']
i_names=['junk','junk','junk','junk','mA_A_tag','Pack_Current','junk','junk','junk','Raw_ADC','junk','junk','Offset_Average','junk','DateTime']
misc=['junk','junk','junk','junk','junk','junk','Fault_Status','junk','junk','junk','Latch_Counter','junk','DateTime']
engy=['junk','junk','junk','junk','junk','Pack_Energy_BMS','junk','Pack_Ah_BMS','junk','SoC(%)','junk','junk','SoH(%)','junk','Cycles','junk','junk','Premature_Cycles','DateTime']
rly=['junk','junk','junk','junk','junk','junk','junk','Main_Relay_Status','junk','Precharge','junk','junk','Pre-discharge','junk','DateTime']
amb = ['junk','junk','junk','junk','junk','amb_RT','RT','DateTime']

shunt_col = ['junk','junk','junk','junk','junk','Shunt_instant_value','junk','junk','Offset_average_shunt','junk','junk','ADC_Samples','DateTime']
gpio_col = ['junk','junk','junk','junk','junk','AFE_GPIO_instant_value','junk','junk','Offset_average_ADC_GPIO','junk','junk','AFE_Samples','DateTime']
cc_col = ['junk','junk','junk','junk','junk','CC_Start','junk','junk','Accumulated_As','junk','Negative_Ah_state','state_rbc','RBC_ast','junk','DateTime']
k_soc_col = ['junk','junk','junk','junk','junk','EKF_SoC','junk','junk','CC_SoC','junk','Calib_Current','junk','junk','Energy','junk','junk','junk','DateTime']
n_soc_col = ['junk','junk','junk','junk','junk', 'Charging_Capacity','junk','junk','Discharging_Capacity','DateTime']

# The cell voltage and cell temperature values are large depending upon the number of cells
# It is also quite repetitive, hence the naming process is automated here.
v_names_1=[f'Cell_{x}' for x in range(1,v_no + 1)] 
t_names_1=[f'T{x}' for x in range(1,t_no + 1)]

#%% Column names for the voltage_data and temperature_data dataframes are properly defined
    
# Five "junk" columns are added at first to represent the 'Date','Hours', 'Minutes', 'Seconds' and 'CV_CT_PC_tag' columns
# 'DateTime' is added at the end, since we created the datetime column previously in this code
# v_names/t_names is a copy as the list with only the cell number names will be used later to name the legend while plotting
   
v_names=v_names_1.copy()
v_names.insert(0,'junk')
v_names.insert(0,'junk')
v_names.insert(0,'junk')
v_names.insert(0,'junk')
v_names.insert(0,'junk')
v_names.append('DateTime')
t_names=t_names_1.copy()
t_names.append('junk')
t_names.insert(0,'junk')
t_names.insert(0,'junk')
t_names.insert(0,'junk')
t_names.insert(0,'junk')
t_names.insert(0,'junk')
t_names.append('DateTime')

#%% Assign correct column names to the respective dataframes
voltage_data.columns=v_names
temperature_data.columns=t_names
max_min_temp_data.columns=mm_temp
max_min_V_data.columns=mm_v 
PV_data.columns=pv
misc_details.columns=misc
energy_details.columns=engy
relay_details.columns=rly
amb_T.columns = amb
shunt_data.columns = shunt_col
gpio_data.columns = gpio_col
cc_start.columns = cc_col
k_soc.columns = k_soc_col
new_soc.columns = n_soc_col

#%% Data Cleaning and Processing
amb_T['Ambient_Temperature'] = amb_T['amb_RT'].str.split('|',expand=True)[0]
amb_T['junk'] = amb_T['amb_RT'].str.split('|',expand=True)[1]
PV_data['Pack_Voltage'] = PV_data['Pack_Voltage_PP'].str.split('|', expand=True)[0]
PV_data['junk'] = PV_data['Pack_Voltage_PP'].str.split('|', expand=True)[1]
PV_data['Pack_Power(W)'] = PV_data['PP_PC'].str.split('|',expand=True)[0]
PV_data['junk'] = PV_data['PP_PC'].str.split('|',expand=True)[1]

cc_start['Negative_Ah'] = cc_start['Negative_Ah_state'].str.split('|',expand=True)[0]
cc_start['junk'] = cc_start['Negative_Ah_state'].str.split('|',expand=True)[1]
cc_start['State'] = cc_start['state_rbc'].str.split('|',expand=True)[0]
cc_start['junk'] = cc_start['state_rbc'].str.split('|',expand=True)[1]
cc_start['RBC'] = cc_start['RBC_ast'].str.split('|',expand=True)[0]
cc_start['junk'] = cc_start['RBC_ast'].str.split('|',expand=True)[1]

cc_start = cc_start.drop(columns = ['Negative_Ah_state', 'state_rbc', 'RBC_ast'])

voltage_data=voltage_data.drop([col for col in voltage_data.columns if "junk" in col], axis=1)
temperature_data=temperature_data.drop([col for col in temperature_data.columns if "junk" in col], axis=1)

max_min_temp_data=max_min_temp_data.drop([col for col in max_min_temp_data.columns if "junk" in col], axis=1)
max_min_V_data=max_min_V_data.drop([col for col in max_min_V_data.columns if "junk" in col], axis=1)
PV_data=PV_data.drop([col for col in PV_data.columns if "junk" in col], axis=1)
misc_details=misc_details.drop([col for col in misc_details.columns if "junk" in col], axis=1)
energy_details=energy_details.drop([col for col in energy_details.columns if "junk" in col], axis=1)
relay_details=relay_details.drop([col for col in relay_details.columns if "junk" in col], axis=1)
amb_T=amb_T.drop([col for col in relay_details.columns if "junk" in col], axis=1)

shunt_data = shunt_data.drop([col for col in shunt_data.columns if "junk" in col], axis=1)
gpio_data = gpio_data.drop([col for col in gpio_data.columns if "junk" in col], axis=1)
cc_start = cc_start.drop([col for col in cc_start.columns if "junk" in col], axis=1)
k_soc = k_soc.drop([col for col in k_soc.columns if "junk" in col], axis=1)

new_soc = new_soc.drop([col for col in new_soc.columns if "junk" in col], axis=1)

voltage_data['Cell_'+str(v_no)]=voltage_data['Cell_'+str(v_no)].str.replace('|','',regex=True)
PV_data['Pack_Current']=PV_data['Pack_Current'].str.replace('|','',regex=True)
PV_data['Pack_Current']=PV_data['Pack_Current'].astype(str).str.extract(r'(\-{0,}\d{1,})')
gpio_data['AFE_Samples'] = gpio_data['AFE_Samples'].str.replace('\|\*','',regex=True)
shunt_data['ADC_Samples'] = shunt_data['ADC_Samples'].str.replace('\|\*','',regex=True)

for x in v_names_1: 
    voltage_data[x]=voltage_data[x].astype(str).str.extract(r'(\d{1,})')
    voltage_data[x]=voltage_data[x].fillna(method='ffill')
    voltage_data[x]=voltage_data[x].astype(int)
    voltage_data=voltage_data[(voltage_data[x]>=2000) & (voltage_data[x]<=4300)]

for x in t_names_1: 
    temperature_data[x]=temperature_data[x].astype(str).str.extract(r'(\d{1,})')
    temperature_data[x] = temperature_data[x].fillna(method='ffill')
    temperature_data[x] = temperature_data[x].astype(int)
    temperature_data[x] = np.where(temperature_data[x]>70,np.nan,temperature_data[x])
    temperature_data[x] = temperature_data[x].fillna(method='ffill')

#%% Calculate Maximum, Minimum, Average and delV from the cell voltage data
voltage_data.loc[:,'Cell_1':'Cell_'+str(v_no)] = voltage_data.loc[:,'Cell_1':'Cell_'+str(v_no)] * 0.001
voltage_data['Mean_V']=voltage_data.loc[:,'Cell_1':'Cell_'+str(v_no)].mean(axis=1)
voltage_data['Max_V']=voltage_data.loc[:,'Cell_1':'Cell_'+str(v_no)].max(axis=1)
voltage_data['Max_V_Cell']=voltage_data.loc[:,'Cell_1':'Cell_'+str(v_no)].astype(int).idxmax(axis=1)
voltage_data['Min_V']=voltage_data.loc[:,'Cell_1':'Cell_'+str(v_no)].min(axis=1)
voltage_data['Min_V_Cell']=voltage_data.loc[:,'Cell_1':'Cell_'+str(v_no)].astype(int).idxmin(axis=1)
voltage_data['delV']=(voltage_data['Max_V']-voltage_data['Min_V'])*1000

#%% DT
# voltage_data['ElapsedTime']=(voltage_data['DateTime']-voltage_data['DateTime'].iloc[0])

#%% Calculate Maximum, Minimum, Average and delT from the cell temperature data
temperature_data.loc[:,'T1':'T'+str(t_no)]=temperature_data.loc[:,'T1':'T'+str(t_no)].astype(float)
temperature_data.loc[:,'T1':'T'+str(t_no)]=temperature_data.loc[:,'T1':'T'+str(t_no)].fillna(method='ffill')
temperature_data['Mean_T']=temperature_data.loc[:,'T1':'T'+str(t_no)].mean(axis=1)
temperature_data['Max_T']=temperature_data.loc[:,'T1':'T'+str(t_no)].max(axis=1)
temperature_data['Max_TCell']=temperature_data.loc[:,'T1':'T'+str(t_no)].astype(float).idxmax(axis=1)
temperature_data['Min_T']=temperature_data.loc[:,'T1':'T'+str(t_no)].min(axis=1)
temperature_data['Min_TCell']=temperature_data.loc[:,'T1':'T'+str(t_no)].astype(float).idxmin(axis=1)
temperature_data['delT']=temperature_data['Max_T']-temperature_data['Min_T']

#%% DT
# temperature_data['ElapsedTime']=(temperature_data['DateTime']-temperature_data['DateTime'].iloc[0])

#%%  Data Cleaning and Preparation
# PV_data['ElapsedTime']=(PV_data['DateTime']-PV_data['DateTime'].iloc[0])
PV_data['Pack_Current']=PV_data['Pack_Current'].astype(float)
gpio_data['AFE_GPIO_instant_value'] = gpio_data['AFE_GPIO_instant_value'].replace('(NaN)', 0)
gpio_data['AFE_GPIO_instant_value'] = gpio_data['AFE_GPIO_instant_value'].astype(float)
gpio_data['AFE_GPIO_instant_value'] = gpio_data['AFE_GPIO_instant_value']*0.001

relay_details['Precharge']=relay_details['Precharge'].astype(str).str.extract(r'(\d{1,})')
relay_details['Main_Relay_Status']=relay_details['Main_Relay_Status'].astype(str).str.extract(r'(\d{1,})')
relay_details['Pre-discharge']=relay_details['Pre-discharge'].astype(str).str.extract(r'(\d{1,})')

relay_details['Main_Relay_Status']=relay_details['Main_Relay_Status'].fillna(method='ffill')
relay_details['Main_Relay_Status']=relay_details['Main_Relay_Status'].fillna(method='bfill')

relay_details['Precharge']=relay_details['Precharge'].fillna(method='ffill')
relay_details['Precharge']=relay_details['Precharge'].fillna(method='bfill')

relay_details['Pre-discharge']=relay_details['Pre-discharge'].fillna(method='ffill')
relay_details['Pre-discharge']=relay_details['Pre-discharge'].fillna(method='bfill')

energy_details['Pack_Ah_BMS']=energy_details['Pack_Ah_BMS'].astype(str).str.replace('0e','',regex=True).str.extract(r'(\-{0,}\d{1,}.\d{0,2})')
energy_details['Pack_Energy_BMS']=energy_details['Pack_Energy_BMS'].astype(str).str.extract(r'(\-{0,}\d{1,6}.\d{0,2})')
energy_details['Pack_Ah_BMS']=energy_details['Pack_Ah_BMS'].astype(float)
energy_details['Pack_Energy_BMS']=energy_details['Pack_Energy_BMS'].astype(float)
energy_details=energy_details[energy_details['Pack_Energy_BMS']<12000]

energy_details['Pack_Ah_BMS']=energy_details['Pack_Ah_BMS'].fillna(method='ffill')
energy_details['Pack_Energy_BMS']=energy_details['Pack_Energy_BMS'].fillna(method='ffill')

# PV_data['Pack_Current']=PV_data['Pack_Current'].astype(str).str.extract(r'(\-{0,}\d{1,}.\d{1,})')
PV_data['Pack_Voltage']=PV_data['Pack_Voltage'].astype(str).str.extract(r'(\d{1,}.\d{1,})')
PV_data['Pack_Current']=PV_data['Pack_Current'].astype(float)
PV_data['Pack_Voltage']=PV_data['Pack_Voltage'].astype(float)
PV_data['Pack_Voltage']=PV_data['Pack_Voltage']*0.001

energy_details['Premature_Cycles'] = energy_details['Premature_Cycles'].str.replace('|','',regex=True)

PV_data = PV_data.drop(columns=['Pack_Voltage_PP', 'PP_PC'])
amb_T = amb_T.drop(columns=['amb_RT'])
amb_T['RT'] = amb_T['RT'].str.replace('|','',regex=True)

#%% details dataframe is made by combining PV_data and energy_details dataframe
# This is done in order to view current and capacity/energy details in the same dataframe
details=pd.DataFrame()
details=pd.merge(left= PV_data, right= energy_details, on='DateTime')
p=list(PV_data.columns)
r=list(energy_details.columns)
r.remove('DateTime')

p.extend(r)
details.columns=p
# Time difference between every datapoint
details['Time_in_sec_s']=(details['DateTime']-details['DateTime'].shift(1))/np.timedelta64(1,'s')
details['Time_in_sec']=(details['DateTime']-details['DateTime'].iloc[0])/np.timedelta64(1,'s')

#%% Get SoC details into a separate dataframe
soc_details = pd.DataFrame()
soc_details = pd.merge(left= voltage_data, right= details, on='DateTime')

soc_details = pd.merge(left= soc_details, right= temperature_data, on='DateTime')

p = list(voltage_data.columns)
r = list(details.columns)
s = list(temperature_data.columns)
r.remove('DateTime')
s.remove('DateTime')

p.extend(r)
p.extend(s)

soc_details.columns=p

#%% Tester data and details are merged together
# This is done in order to view tester capacity/energy data and bms capacity/energy data in the same dataframe
tester_pack_details=pd.merge(left= tester_data, right= details, left_on='Date', right_on='DateTime')
p=list(tester_data.columns)
r=list(details.columns)
p.extend(r)
tester_pack_details.columns = p

#%%
tester_pack_cv=pd.merge(left= tester_data, right= voltage_data, left_on='Date', right_on='DateTime')
p=list(tester_data.columns)
r=list(voltage_data.columns)
p.extend(r)
tester_pack_cv.columns = p

#%% The pack current and voltage data are extracted into correct format
# Most of the times, the values have appropriate format; so Uncomment only if necessary


# details['Pack_Current']=details['Pack_Current'].astype(str).str.extract(r'(\-{0,}\d{1,}.\d{1,})')
# details['Pack_Voltage']=details['Pack_Voltage_BMS(mV)'].astype(str).str.extract(r'(\d{1,}.\d{1,})')
# details['Pack_Current']=details['Pack_Current'].astype(float)
# details['Pack_Voltage']=details['Pack_Voltage'].astype(float)
# details['Pack_Voltage']=details['Pack_Voltage']*0.001

#%% The state of the pack is assigned based on the current value
# '0' refers to "Charge" state
# '1' refers to "Discharge" state
# '2' refers to "Rest" state
# Why these values? My code, my rules.

details['State']=np.nan
details['Time_in_sec_s_cap']=np.where((details['Time_in_sec_s']>300),np.nan,details['Time_in_sec_s'])
details = details[details['Pack_Voltage']<100]
details['State']=np.where(details['Pack_Current']>0,0,details['State'])
details['State']=np.where(details['Pack_Current']<0,1,details['State'])
details['State']=np.where(details['Pack_Current']==0,2,details['State'])

#%% Capacity and energy values are calculated based on "Coulomb-counting" technique
# Current and time are used to calculate capacity
# Current, time and voltage are used to calculate energy
# This is done in order to validate the values calculated in the BMS


details['Cap_inst']=details['Time_in_sec_s_cap']*abs(details['Pack_Current'])/3600
details['Capacity_calculated']=details['Cap_inst'].groupby(details['State']).cumsum()
details['Capacity_calculated_chg']=(details[details['State']==0]['Cap_inst']).cumsum()
details['Capacity_calculated_dchg']=(details[details['State']==1]['Cap_inst']).cumsum()

details['Capacity_calculated_chg']=details['Capacity_calculated_chg'].fillna(method='bfill')
details['Capacity_calculated_dchg']=details['Capacity_calculated_dchg'].fillna(method='bfill')

details['Energy_calculated']=details['Capacity_calculated']*details['Pack_Voltage']
details['Energy_calculated_chg']=details['Capacity_calculated_chg']*details['Pack_Voltage']
details['Energy_calculated_dchg']=details['Capacity_calculated_dchg']*details['Pack_Voltage']

#%% Calculate time difference between every data point
# This is done to check if any datapoint is missing and,  
# to get the instantaneous "time" variable needed to calculate capacity/energy

tester_data['Time_in_sec_s']=(tester_data['Date']-tester_data['Date'].shift(1))/np.timedelta64(1,'s')# time difference b/w every data pt.
tester_data['Time_in_sec']=(tester_data['Date']-tester_data['Date'].iloc[0])/np.timedelta64(1,'s')

tester_data['Time_in_sec_s_cap']=np.where((tester_data['Time_in_sec_s']>300),np.nan,tester_data['Time_in_sec_s'])

#%% If the test conducted has a pulse charging technique, the capacity values are not accumulated
# Hence, the full value is not immediately known
# When that is the case, uncomment this section to get the cumulative values of the capacity/energy

# tester_data['Cap_inst']=tester_data['Time_in_sec_s_cap']*abs(tester_data['Current(A)'])/3600
# tester_data['Capacity_calculated']=tester_data['Cap_inst'].groupby(tester_data['Step Type']).cumsum()
# tester_data['Capacity_calculated_chg']=(tester_data[tester_data['Step Type']=='CCCV Chg']['Cap_inst']).cumsum()
# tester_data['Capacity_calculated_dchg']=(tester_data[tester_data['Step Type']=='CC DChg']['Cap_inst']).cumsum()

# tester_data['Capacity_calculated_chg']=tester_data['Capacity_calculated_chg'].fillna(method='bfill')
# tester_data['Capacity_calculated_dchg']=tester_data['Capacity_calculated_dchg'].fillna(method='bfill')

# tester_data['Energy_calculated']=tester_data['Capacity_calculated']*tester_data['Voltage(V)']
# tester_data['Energy_calculated_chg']=tester_data['Capacity_calculated_chg']*tester_data['Voltage(V)']
# tester_data['Energy_calculated_dchg']=tester_data['Capacity_calculated_dchg']*tester_data['Voltage(V)']

# tester_data['SoC_Calculated'] = (tester_data['Capacity_calculated'] / 210) * 100

# tester_data = tester_data.drop(columns=['Cap_inst'])

#%% The ambient temperature and Relay temperature are changed into appropriate formats
# any impossible values like temperatures less than or equal to 0 are removed to eliminate junk values
# The junk values are usually due to wire connection/operational issues 

amb_T['Ambient_Temperature'] = amb_T['Ambient_Temperature'].astype(int)
amb_T['RT'] = amb_T['RT'].astype(int)
amb_T = amb_T[amb_T['Ambient_Temperature']>0]
amb_T = amb_T[amb_T['RT']>0]

#%% The length and breadth of the plot window to be opened; values selected on trial and error
l=7.5
b=4.5

#%% Plot Pre-Charge data
document.add_page_break()
document.add_paragraph().add_run('Precharge:').bold=True

plt.figure(figsize=(l,b))
plt.plot(relay_details['DateTime'],relay_details['Precharge'])
plt.grid(linestyle='dotted')
plt.xlabel('DateTime',fontweight='bold')
plt.title('Precharge',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Precharge.png',dpi=1200)

document.add_picture(img_path+'\\Precharge.png')

#%% Plot Pre-discharge data
document.add_paragraph().add_run('Pre-discharge:').bold=True

plt.figure(figsize=(l,b))
plt.plot(relay_details['DateTime'],relay_details['Pre-discharge'])
plt.grid(linestyle='dotted')
plt.xlabel('DateTime',fontweight='bold')
plt.title('Pre-discharge',fontweight='bold')
plt.tight_layout()

plt.savefig(img_path+'\\Pre-discharge.png',dpi=1200)

document.add_picture(img_path+'\\Pre-discharge.png')
#%% Plot Main Relay data
document.add_paragraph().add_run('Main Relay:').bold=True

plt.figure(figsize=(l,b))
plt.plot(relay_details['DateTime'],relay_details['Main_Relay_Status'])
plt.grid(linestyle='dotted')
plt.xlabel('DateTime',fontweight='bold')
plt.title('Main Relay',fontweight='bold')
plt.tight_layout()

plt.savefig(img_path+'\\Main Relay.png',dpi=1200)

document.add_picture(img_path+'\\Main Relay.png')

#%% Plot BMS Pack Current data
document.add_paragraph().add_run('Pack Current:').bold=True

plt.figure(figsize=(l,b))
plt.plot(PV_data['DateTime'],PV_data['Pack_Current']*0.001)
plt.ylabel('Current(A)',fontweight='bold')
plt.xlabel('DateTime',fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Pack Current',fontweight='bold')
plt.tight_layout()

plt.savefig(img_path+'\\Pack Current.png',dpi=1200)

document.add_picture(img_path+'\\Pack Current.png')

#%% Plot Tester Current data
document.add_paragraph().add_run('Pack Current(Tester data):').bold=True

plt.figure(figsize=(l,b))
plt.plot(tester_data['Date'],tester_data['Current(A)'])
plt.ylabel('Current(A)',fontweight='bold')
plt.xlabel('DateTime',fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Tester Current',fontweight='bold')
plt.tight_layout()

plt.savefig(img_path+'\\Pack Current_tester.png',dpi=1200)

document.add_picture(img_path+'\\Pack Current_tester.png')

#%% Plot BMS Capacity
document.add_paragraph().add_run('Pack Capacity (BMS):').bold=True

plt.figure(figsize=(l,b))
plt.plot(energy_details['DateTime'],energy_details['Pack_Ah_BMS'])
plt.ylabel('Capacity(Ah)',fontweight='bold')
plt.xlabel('DateTime',fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Pack Capacity (BMS)',fontweight='bold')
plt.tight_layout()

plt.savefig(img_path+'\\Pack Capacity_BMS.png',dpi=1200)
 
document.add_picture(img_path+'\\Pack Capacity_BMS.png')
document.add_paragraph('In this session, the capacity expended by the pack is about %sAh ()' %energy_details['Pack_Ah_BMS'].max() )

#%%
tester_data['Capacity_diff'] = tester_data['Capacity(Ah)'] - tester_data['Capacity(Ah)'].shift()
tester_data['Capacity_diff'] = np.where(tester_data['Capacity_diff']<0, 0, 
                                    tester_data['Capacity_diff'])

tester_data['Capacity_inst'] = tester_data['Capacity_diff'].groupby(tester_data['Step Type']).cumsum()

tester_data['Capacity_cumulative'] = np.where(tester_data['Step Type']=='CCCV Chg',
                                          tester_data['Capacity_inst'],tester_data['Capacity(Ah)'])

tester_data['SoC_Tester'] = (tester_data['Capacity_cumulative']/210) * 100

tester_data = tester_data.drop(columns=['Capacity_inst', 'Capacity_diff'])

#%% Plot Tester Capacity
document.add_paragraph().add_run('Pack Capacity (Tester data):').bold=True

plt.figure(figsize=(l,b))
plt.plot(tester_data['Date'], tester_data['Capacity(Ah)'])
plt.xlabel('DateTime',fontweight='bold')

plt.ylabel('Capacity(Ah)',fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('Pack Capacity (Tester data)',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Pack Capacity_tester.png',dpi=1200)
 
document.add_picture(img_path+'\\Pack Capacity_tester.png')
document.add_paragraph('In this session, the capacity expended by the pack, according to the tester measurement, is about %sAh ()' %tester_data['Capacity_cumulative'].max() )

#%% Plot BMS Pack Energy
document.add_paragraph().add_run('Pack Energy (BMS):').bold=True

plt.figure(figsize=(l,b))
plt.plot(energy_details['DateTime'],energy_details['Pack_Energy_BMS'])

plt.ylabel('Energy(Wh)',fontweight='bold')
plt.xlabel('DateTime',fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('Pack Energy (BMS)',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Pack Energy_BMS.png',dpi=1200)

document.add_picture(img_path+'\\Pack Energy_BMS.png')
p_bms_d_engy=document.add_paragraph('The ')
p_bms_d_engy.add_run('discharging pack energy ').bold=True
p_bms_d_engy.add_run('is about ')
p_bms_d_engy.add_run('%s Wh.' % energy_details['Pack_Energy_BMS'].max()).bold=True

#%% 
tester_data['Energy_diff'] = tester_data['Energy(Wh)'] - tester_data['Energy(Wh)'].shift()
tester_data['Energy_diff'] = np.where(tester_data['Energy_diff']<0, 0, 
                                  tester_data['Energy_diff'])

tester_data['Energy_inst'] = tester_data['Energy_diff'].groupby(tester_data['Step Type']).cumsum()

tester_data['Energy_cumulative'] = np.where(tester_data['Step Type']=='CCCV Chg',
                                          tester_data['Energy_inst'],tester_data['Energy(Wh)'])


tester_data = tester_data.drop(columns=['Energy_inst', 'Energy_diff'])

#%% Plot Tester Pack Energy
document.add_paragraph().add_run('Pack Energy (Tester data):').bold=True

plt.figure(figsize=(l,b))
plt.plot(tester_data['Date'],tester_data['Energy(Wh)'])

plt.ylabel('Energy(Wh)',fontweight='bold')
plt.xlabel('DateTime',fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('Pack Energy (Tester)',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Pack Energy_tester.png',dpi=1200)

document.add_picture(img_path+'\\Pack Energy_tester.png')
p_bms_d_engy=document.add_paragraph('The ')
p_bms_d_engy.add_run('pack energy, ').bold=True
p_bms_d_engy.add_run('according to the tester observations, is about ')
p_bms_d_engy.add_run('%s Wh.' % tester_data['Energy_cumulative'].max()).bold=True

#%% Plot Pack Voltage
document.add_paragraph().add_run('Pack Voltage:').bold=True

plt.figure(figsize=(l,b))
plt.plot(details['DateTime'],details['Pack_Voltage'])

plt.ylabel('Voltage(V)',fontweight='bold')
plt.xlabel('DateTime',fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('Pack Voltage',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Pack Voltage.png',dpi=1200)

document.add_picture(img_path+'\\Pack Voltage.png')
pv=document.add_paragraph('At ')
pv.add_run('discharge, ').bold=True
pv.add_run('during the first cycle, the pack goes from to . The pack voltage at rest after discharge is . At ')
pv.add_run('charge, ').bold=True
pv.add_run('the pack goes from to . ')
pv.add_run('In the end, the voltage stabilizes at V during the end.')
#%% Plot Tester Pack Voltage
document.add_paragraph().add_run('Pack Voltage (Tester data):').bold=True

plt.figure(figsize=(l,b))
plt.plot(tester_data['Date'],tester_data['Voltage(V)']) 

plt.ylabel('Voltage(V)',fontweight='bold')
plt.xlabel('DateTime',fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('Pack Voltage (Tester)',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Pack Voltage_tester.png',dpi=1200)

document.add_picture(img_path+'\\Pack Voltage_tester.png')

#%% Plot Cell Voltage data
document.add_paragraph().add_run('Cell Voltage:').bold=True

plt.figure(figsize=(l,b))
plt.plot(voltage_data['DateTime'],voltage_data.loc[:,'Cell_1':'Cell_'+str(v_no)])
plt.ylabel('Voltage(mV)',fontweight='bold')
plt.xlabel('DateTime',fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Cell Voltage',fontweight='bold')
plt.legend(v_names_1)
plt.tight_layout()

# The savefig command is commented as the cell voltage generally has too many values in the legend, which makes it illegible
# plt.savefig(img_path+'\\Cell Voltage.png',dpi=1200)

#%% Plot Average Voltage
document.add_paragraph().add_run('Average Voltage:').bold=True

plt.figure(figsize=(l,b))
plt.plot(voltage_data['DateTime'],(voltage_data['Mean_V']))
plt.ylabel('Voltage(V)',fontweight='bold')
plt.xlabel('DateTime',fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Average Voltage',fontweight='bold')
plt.tight_layout()

plt.savefig(img_path+'\\Average Voltage.png',dpi=1200)

document.add_picture(img_path+'\\Average Voltage.png')
av=document.add_paragraph('The average cell voltage during ')
av.add_run('discharge ').bold=True
av.add_run('goes from ')
av.add_run('V to V. ').bold=True
av.add_run('The average cell voltage at the rest after discharge is V. At ')
av.add_run('charge').bold=True
av.add_run(', the average cell voltage goes from ')
av.add_run('V to V.').bold=True
av.add_run('In the end rest, it stabilizes at %sV.' % (voltage_data['Mean_V'].iloc[-1]))


#%% Extract delV values calculated by the BMS into the proper format

max_min_V_data['deltaV(mV)_BMS']=max_min_V_data['deltaV(mV)_BMS'].astype(str).str.extract(r'(\d{1,})')
max_min_V_data['deltaV(mV)_BMS']=max_min_V_data['deltaV(mV)_BMS'].fillna(method='ffill')
max_min_V_data['deltaV(mV)_BMS'] = max_min_V_data['deltaV(mV)_BMS'].astype(int)

#%% Plot BMS delV data

document.add_paragraph().add_run('BMS Voltage Difference (delV):').bold=True
plt.figure(figsize=(l,b))
plt.plot(max_min_V_data['DateTime'],(max_min_V_data['deltaV(mV)_BMS']))

plt.ylabel('Voltage(mV)',fontweight='bold')
plt.xlabel('DateTime',fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Voltage difference (delV)',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\delV_bms.png',dpi=1200)

document.add_picture(img_path+'\\delV_bms.png')
dv=document.add_paragraph('During discharge, the ')
dv.add_run('maximum delV (difference between the highest and lowest cell voltage), ').bold=True
dv.add_run('is about ')
dv.add_run('mV, at the end of discharge. ').bold=True
dv.add_run('At the rest after discharging, the delmV is mV. In the charging session, the maximum delV is mV. ')
dv.add_run('The pack stabilizes at %smV delV at the end rest. '% (voltage_data['delV'].iloc[-1]))

document.add_paragraph().add_run('Cell Temperature:').bold=True

#%% Plot Cell Temperature

plt.figure(figsize=(l,b))
plt.plot(temperature_data['DateTime'],temperature_data.loc[:,'T1':'T'+str(t_no)])
plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
plt.xlabel('DateTime',fontweight='bold')
plt.legend(t_names_1)
plt.grid(linestyle='dotted')
plt.title('Cell Temperature',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Cell Temperature.png',dpi=1200)

# The savefig command is commented as the cell temperature generally has too many values in the legend, which makes it illegible
# document.add_picture(img_path+'\\Cell Temperature.png') 
ct=document.add_paragraph('The maximum temperature reached is around ')
ct.add_run('%s'%(temperature_data.loc[:,'T1':'T'+str(t_no)].max(axis=1).max())).bold=True
ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True

#%% Plot Average Temperature
document.add_paragraph().add_run('Average temperature:').bold=True
plt.figure(figsize=(l,b))
plt.plot(temperature_data['DateTime'],temperature_data['Mean_T'])

plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
plt.xlabel('DateTime',fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Average Temperature',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Average Temperature.png',dpi=1200)

document.add_picture(img_path+'\\Average Temperature.png')
at=document.add_paragraph('As can be seen, the ')
at.add_run('average cell temperature ').bold=True
at.add_run('stays between ')
at.add_run(u'\N{DEGREE SIGN}'+'C to '+u'\N{DEGREE SIGN}'+'C ').bold=True
at.add_run('during ')
at.add_run('discharging').bold=True
at.add_run(' and ')
at.add_run(u'\N{DEGREE SIGN}'+'C to '+u'\N{DEGREE SIGN}'+'C ').bold=True
at.add_run('during ')
at.add_run('charging.').bold=True

#%% Plot Cell Temperature difference (delT)
document.add_paragraph().add_run('Temperature Difference (delT):').bold=True
plt.figure(figsize=(l,b))
plt.plot(temperature_data['DateTime'],temperature_data['delT'])

plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
plt.xlabel('DateTime',fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Temperature difference (delT)',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\delT.png',dpi=1200)

document.add_picture(img_path+'\\delT.png')
dt=document.add_paragraph('The ')
dt.add_run('maximum delT ').bold=True
dt.add_run('is about ')
dt.add_run('%s' %temperature_data['delT'].max()).bold=True
dt.add_run(u'\N{DEGREE SIGN}'+'C.').bold=True


#%% Plot Ambient Temperature
document.add_paragraph().add_run('Ambient Temperature:').bold=True
plt.figure(figsize=(l,b))
plt.plot(amb_T['DateTime'],amb_T['Ambient_Temperature'])

plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
plt.xlabel('DateTime',fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Ambient Temperature',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\ambT.png',dpi=1200)

document.add_picture(img_path+'\\ambT.png')
dt=document.add_paragraph('The ')
dt.add_run('ambient temperature ').bold=True
dt.add_run('is about ')
dt.add_run(u'\N{DEGREE SIGN}'+'C.').bold=True

#%% Plot Relay Temperature
document.add_paragraph().add_run('Relay Temperature:').bold=True
plt.figure(figsize=(l,b))
plt.plot(amb_T['DateTime'],amb_T['RT'])

plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
plt.xlabel('DateTime',fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Relay Temperature',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\RT.png',dpi=1200)

document.add_picture(img_path+'\\RT.png')
dt=document.add_paragraph('The ')
dt.add_run('relay temperature ').bold=True
dt.add_run('is about ')
dt.add_run(u'\N{DEGREE SIGN}'+'C.').bold=True

#%% Plot SoC Comparison between EKF_SoC, CC_SoC and BMS SoC
document.add_paragraph().add_run('SoC Comparison:').bold=True
plt.figure(figsize=(l,b))
plt.plot(k_soc['DateTime'],k_soc['EKF_SoC'],label='EKF_SoC')
plt.plot(k_soc['DateTime'],k_soc['CC_SoC'],label='CC_SoC')
plt.plot(details['DateTime'],details['SoC(%)'],label='SoC (%)')

plt.ylabel('SoC (%)',fontweight='bold')
plt.xlabel('DateTime',fontweight='bold')
plt.grid(linestyle='dotted')
plt.legend()
plt.title('SoC Comparison',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\SoCs.png',dpi=1200)

document.add_picture(img_path+'\\SoCs.png')

#%% Add final observations table in the word document
document.add_paragraph().add_run('Tabulation:').bold=True

table2=document.add_table(rows=3,cols=22)

table2.cell(0,0).text='Test Date'
table2.cell(0,1).text='Step'
table2.cell(0,2).text='Current'
table2.cell(0,3).text='Pack Voltage'
table2.cell(0,4).text='Capacity (Tester)'
table2.cell(0,5).text='Capacity (BMS)'
table2.cell(0,6).text='Energy (Tester)'
table2.cell(0,7).text='Energy (BMS)'
table2.cell(0,8).text='Tester SoC(%)'
table2.cell(0,9).text='BMS EKF SoC(%) range'
table2.cell(0,10).text='BMS CC SoC(%) range'
table2.cell(0,11).text='Average Cell Voltage'
table2.cell(0,12).text='delV'
table2.cell(0,14).text='Temperature'
table2.cell(0,16).text='delT'
table2.cell(0,18).text='Ambient Temperature'
table2.cell(0,20).text='Relay Temperature'

table2.cell(1,12).text='Start'
table2.cell(1,13).text='End'
table2.cell(1,14).text='Start'
table2.cell(1,15).text='End'
table2.cell(1,16).text='Start'
table2.cell(1,17).text='End'
table2.cell(1,18).text='Start'
table2.cell(1,19).text='End'
table2.cell(1,20).text='Start'
table2.cell(1,21).text='End'

table2.cell(2,0).text= voltage_data['DateTime'].iloc[-1] - voltage_data['DateTime'].iloc[0]
table2.cell(2,1).text = step
table2.cell(2,2).text = str((PV_data['Pack_Current'].max())*0.001)+'A/'+str((PV_data['Pack_Current'].min())*0.001) +'A'
table2.cell(2,3).text = str(details['Pack_Voltage'].iloc[0]) + 'V to ' + str(details['Pack_Voltage'].iloc[-1]) +'V'
table2.cell(2,4).text = str(tester_data['Capacity_cumulative'].max()) + 'Ah'
table2.cell(2,5).text = str(energy_details['Pack_Ah_BMS'].max()) + 'Ah' 
table2.cell(2,6).text = str(tester_data['Energy_cumulative'].max()) + 'Wh'
table2.cell(2,7).text = str(energy_details['Pack_Energy_BMS'].max()) + 'Wh' 
table2.cell(2,11).text = str(voltage_data['Mean_V'].iloc[0]) + 'V to ' + str(voltage_data['Mean_V'].iloc[-1]) +'V'
table2.cell(2,12).text = str(voltage_data['delV'].iloc[0]) + 'mV'
table2.cell(2,13).text = str(voltage_data['delV'].iloc[-1]) + 'mV'
table2.cell(2,14).text = str(temperature_data['Min_T'].iloc[0]) + u'\N{DEGREE SIGN}'+'C'
table2.cell(2,15).text = str(temperature_data['Max_T'].iloc[-1]) + u'\N{DEGREE SIGN}'+'C'
table2.cell(2,16).text = str(temperature_data['delT'].iloc[0]) + u'\N{DEGREE SIGN}'+'C'
table2.cell(2,17).text = str(temperature_data['delT'].iloc[-1]) + u'\N{DEGREE SIGN}'+'C'
table2.cell(2,18).text = str(amb_T['Ambient_Temperature'].iloc[0]) + u'\N{DEGREE SIGN}'+'C'
table2.cell(2,19).text = str(amb_T['Ambient_Temperature'].iloc[-1]) + u'\N{DEGREE SIGN}'+'C'
table2.cell(2,20).text = str(amb_T['RT'].iloc[0]) + u'\N{DEGREE SIGN}'+'C'
table2.cell(2,21).text = str(amb_T['RT'].iloc[-1]) + u'\N{DEGREE SIGN}'+'C'

table2.style='Table Grid'

#%% 
# Get path of the final excel file to be saved from the path of the input file
path=img_path+'\\'+filename+'_modified.xlsx'

# Save the word document in the same path as the plots
document.save(img_path+'\\'+filename+'_observations_document.docx')

#%%
chg_soc_details = soc_details[soc_details['Pack_Current']>2000]
dchg_soc_details = soc_details[soc_details['Pack_Current']<-2000]

#%% Save all the dataframes into separate sheets in the same excel
writer = pd.ExcelWriter(path, engine = 'xlsxwriter')
voltage_data.to_excel(writer, sheet_name = 'Cell Voltage',index=False)
temperature_data.to_excel(writer, sheet_name = 'Cell Temperature',index=False)
PV_data.to_excel(writer,sheet_name='Pack Voltage Details',index=False)
relay_details.to_excel(writer,sheet_name='Relay Status',index=False)
max_min_V_data.to_excel(writer,sheet_name='Voltage Max Min',index=False)
max_min_temp_data.to_excel(writer,sheet_name='Temp Max Min',index=False)
misc_details.to_excel(writer,sheet_name='Fault Status',index=False)
energy_details.to_excel(writer,sheet_name='Energy Details',index=False)
details.to_excel(writer,sheet_name='Pack Details',index=False) #DT
amb_T.to_excel(writer, sheet_name = 'Ambient Temperature', index = False)
shunt_data.to_excel(writer, sheet_name = 'Shunt Data', index = False)
gpio_data.to_excel(writer, sheet_name = 'GPIO Data', index = False)
cc_start.to_excel(writer, sheet_name = 'CC Start Data', index = False)
k_soc.to_excel(writer, sheet_name = 'EKF_SoC', index = False)
tester_data.to_excel(writer, sheet_name = 'Tester data', index = False)
chg_soc_details.to_excel(writer, sheet_name = 'Chg SoC Details', index = False)
dchg_soc_details.to_excel(writer, sheet_name = 'Dchg SoC Details', index = False)
tester_pack_details.to_excel(writer, sheet_name = 'Tester and Pack data', index = False)
new_soc.to_excel(writer, sheet_name = 'New SoC Values', index = False)

writer.close()

#%% Create and save a text file containing some of the observations about the file
f=open(img_path+'\\'+filename+'_observations.txt',"w")
vol_col_check=v_names_1
temp_col_check=t_names_1

print('======================================================================',file=f)
print('Datapoints:%s' % (len(voltage_data)),file=f)
print('Maximum Charging Current:%s ; Minimum Charging Current:%s' % (details[details['State']==0]['Pack_Current'].max(),details[details['State']==0]['Pack_Current'].min()),file=f)
print('Maximum Discharging Current:%s ; Minimum Discharging Current:%s' % (details[details['State']==1]['Pack_Current'].max(),details[details['State']==1]['Pack_Current'].min()),file=f)
print('Maximum Charging Capacity:%s ; Minimum Charging Capacity:%s' % (details['Capacity_calculated_chg'].max(),details['Capacity_calculated_chg'].min()),file=f)
print('Maximum Discharging Capacity:%s ; Minimum Discharging Capacity:%s' % (details['Capacity_calculated_dchg'].max(),details['Capacity_calculated_dchg'].min()),file=f)
print('Maximum Charging Energy:%s ; Minimum Charging Energy:%s' % (details['Energy_calculated_chg'].max(),details['Energy_calculated_chg'].min()),file=f)
print('Maximum Discharging Energy:%s ; Minimum Discharging Energy:%s' % (details['Energy_calculated_dchg'].max(),details['Energy_calculated_dchg'].min()),file=f)
print('Maximum delT:%s degC' %(temperature_data['delT'].max()),file=f)
print('======================================================================',file=f)

f.close()
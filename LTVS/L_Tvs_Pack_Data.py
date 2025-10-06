# -*- coding: utf-8 -*-
"""
Created on Thu Dec 21 16:19:40 2023

@author: IITM
"""
#%% Code Description

"""
USE THIS CODE IF:
    1. Putty log data needs to be processed and report generated
    2. Tester data is available
    3. Putty log does NOT have date and time    
"""

"""
PRE-REQUISITES:
    1. The putty log need to be imported into MS Excel first
    2. While opening in Excel, open it by mentioning the "comma", "space" and in "other" - ":", as the delimiters
    3. Save the file as an Excel file (i.e., with .xlsx extension)
"""

"""
NOTE:
    1. This code generally needs multiple changes for it to work without errors
    2. It is because the data collection process is highly irregular
    3. Proceed with caution at every cell, otherwise useful data may be lost.
    4. Check and double-check the data after executing every cell

"""

#%% Import necessary libraries
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import docx
import os
import glob

#%% Get all the files with names that match the pattern of the names of the files that needs processing

# Care must be taken that the files that need processing should have the same naming format
files = glob.glob(r"C:\Users\IITM\Downloads\10kwh pack 0.3C cyc test data 29.10.2024 1.xlsx")

# Will be added in the "Step" column in the final word document
# After the file is saved, open the MS Word file and change according to the data received
step = 'Charge/Discharge'

# Generally, it is just "record"; only change if the received file has been altered
testerSheetName = "Bms raw data " 

cols=['CV_CT_PC_tag','Cell 1','Cell 2','Cell 3','Cell 4','Cell 5','Cell 6','Cell 7',
      'Cell 8','Cell 9','Cell 10','Cell 11','Cell 12','Cell 13','Cell 14','Cell 15',
      'Cell 16','BR_ID_tag','BR_ID_th','tm','ts']

# Number of cells
v_no = 16
t_no = 16

# chosen through trial and error; optimal value chosen for displaying in a word document
l=7.5
b=4.5

#%%

for file in files:
    
    # Read the excel files using read_excel() function of pandas library
    data=pd.read_excel(file,names=cols,index_col=False, sheet_name=testerSheetName)
    
    tester_data = pd.read_excel(r"C:\Users\IITM\Downloads\10kwh pack 0.3C cyc test data 29.10.2024 1.xlsx", 
                                sheet_name = 'record')
    
    # Change the date and time column (called "Date") to the pandas datetime format
    # Only change the name of the "Date" column if it has been altered beforehand
    tester_data['Date'] = pd.to_datetime(tester_data['Date'], errors = 'coerce')
    
    # Create destination folder path from the input file's path
    # This folder is in the same path as the input file
    # All images and the observation word document are saved in this folder.
    img_path=file.rsplit('\\',1)[0]+'\\'+file.rsplit('\\',1)[1].rsplit('.',1)[0]
    
    # Separate the name of the file from the file path
    filename=file.rsplit('\\',1)[1].rsplit('.',1)[0]
    
    # If the destination folder path does not exist, the folder is created
    if not os.path.exists(img_path):
        os.makedirs(img_path)
    
    #%% Word document creation
    
    # Create the word document using the docx library
    document=docx.Document()
    document.styles['Normal'].font.name='Calibri'
    H1=document.add_heading()
    H1.alignment=1
    H1.add_run('LTVS Pack Data').font.size=docx.shared.Pt(18)
    
    # Add a table with the title "Pack Specifications" and input all relevant information into the table
    document.add_paragraph().add_run('\n\nPack Specifications:').bold=True
    table1=document.add_table(rows=5,cols=2)
    
    table1.cell(0,0).text='Chemistry'
    table1.cell(0,1).text='LFP'
    table1.cell(1,0).text='Pack Capacity'
    table1.cell(1,1).text='210Ah (2*105Ah)'
    table1.cell(2,0).text='Configuration'
    table1.cell(2,1).text='2P16S'
    table1.cell(3,0).text='UV cut-off'
    table1.cell(3,1).text='2.8V'
    table1.cell(4,0).text='OV cut-off'
    table1.cell(4,1).text='3.5V'
    
    table1.style='Table Grid'
    
    #%%
    # Replace unrelated symbols like "|" and ":" to clean data for further processing
    data=data.replace('|','',regex=True)
    data=data.replace(':','',regex=True)
    
    # Since the data is usually huge, a portion of the data is sliced
    # This is to view the structure of the data logging, which will come in handy in the following section
    d1=data.head(1000)
    
    #%% Separate data into different dataframes
    
    # This is done as the data is in putty log format, which does not have data separated into proper columns
    # Note the values to which the "CV_CT_PC_tag" is checked against (eg, '|CV_1[1..16](mV)', '|Min_temp')
    # These values need to be checked if they are right, using the 'd1' dataframe in the previous section
    # Change the values if they are different in the input data file
    
    voltage_data=data[data['CV_CT_PC_tag']=='|CV_1[1..16](mV)'].dropna(how='all',axis=1).reset_index(drop=True)
    temperature_data=data[data['CV_CT_PC_tag']=='|Cell_Temperatures[1..16]'].dropna(how='all',axis=1).reset_index(drop=True)
    max_min_temp_data=data[data['CV_CT_PC_tag']=='|Min_temp'].dropna(how='all',axis=1).reset_index(drop=True)
    max_min_V_data=data[data['CV_CT_PC_tag']=='|Max'].dropna(how='all',axis=1).reset_index(drop=True)
    PV_data=data[data['CV_CT_PC_tag']=='|PV(mV)'].dropna(how='all',axis=1).reset_index(drop=True)
    misc_details=data[data['CV_CT_PC_tag']=='|Fault'].dropna(how='all',axis=1).reset_index(drop=True)
    energy_details=data[data['CV_CT_PC_tag']=='|Pack_Energy'].dropna(how='all',axis=1).reset_index(drop=True)
    relay_details=data[data['CV_CT_PC_tag']=='|RLY'].dropna(how='all',axis=1).reset_index(drop=True)
    amb_T = data[data['CV_CT_PC_tag']=='AMB'].dropna(how='all',axis=1).reset_index(drop=True)
    soc_data = data[data['CV_CT_PC_tag']=='Calib_current'].dropna(how='all',axis=1).reset_index(drop=True)

    
    #%% Define Proper column names for each dataframe
    
    # Change column names if necessary
    # Give the name "junk" to columns with useless data (like Nan or some symbol); It will be removed later
    # Changes often with each input dataset as the data logging process is usually irregular
     
    mm_temp=['junk','junk','Min_T [Cell_No]','junk','junk','junk','Max_T [Cell_No]','junk','junk','Flash_Usage(%)']
    mm_v=['junk','junk','Max_Cell_V_BMS(mV)','Max_V_Cell_No_BMS','junk','junk','junk','Min_Cell_V_BMS(mV)','Min_V_Cell_No_BMS','junk','deltaV(mV)_BMS']# CFD Test Data
    pv=['junk','Pack_Voltage_PP','PP_PC','junk','Pack_Current']
    misc=['junk','junk','Fault_Status','junk','junk','junk','junk','Latch_Counter','junk']
    engy=['junk','junk','Pack_Energy_BMS','junk','junk','Pack_Ah_BMS','junk','junk','SoC(%)','junk','junk','junk','SoH(%)','junk','junk','Cycles','junk','junk','junk','Premature_Cycles']
    rly=['junk','junk','junk','junk','Main_Relay_Status','junk','junk','Precharge','junk','junk','junk','Pre-discharge','junk']
    amb = ['junk','junk','amb_RT','junk','RT']
    soc = ['junk', 'Calibration_Current', 'junk', 'junk', 'EKF_SoC', 'junk', 'junk', 'CC_SoC','junk','Energy','junk']

    # The cell voltage and cell temperature values are large depending upon the number of cells
    # It is also quite repetitive, hence the naming process is automated here.
    v_names_1=[f'Cell_{x}' for x in range(1,v_no + 1)] 
    t_names_1=[f'T{x}' for x in range(1,t_no + 1)]
    
    #%% Column names for the voltage_data and temperature_data dataframes are properly defined
    
    # "junk" is added as the first element to represent the column 'CV_CT_PC_tag'
    # v_names/t_names is a copy as the list with only the cell number names will be used later to name the legend while plotting
    v_names=v_names_1.copy()
    v_names.insert(0,'junk')
    
    t_names=t_names_1.copy()
    t_names.insert(0,'junk')
    t_names.append('junk') # Only if the input dataset format has an unnecessary empty column in the end
    
    #%% Assign properly defined column names to the dataframe columns
    
    voltage_data.columns=v_names
    temperature_data.columns=t_names
    max_min_temp_data.columns=mm_temp
    max_min_V_data.columns=mm_v 
    PV_data.columns=pv
    misc_details.columns=misc
    energy_details.columns=engy
    relay_details.columns=rly
    amb_T.columns = amb
    soc_data.columns = soc
    
    #%% Data Cleaning and Processing - I
    
    # 'amb_RT' has a ambient temperature and a string "RT" separated by '|' in the same column
    # So it is separated and ambient temperature is taken
    # Same goes for pack voltage and pack power
    
    amb_T['Ambient_Temperature'] = amb_T['amb_RT'].str.split('|',expand=True)[0]
    amb_T['junk'] = amb_T['amb_RT'].str.split('|',expand=True)[1]
    PV_data['Pack_Voltage'] = PV_data['Pack_Voltage_PP'].str.split('|', expand=True)[0]
    PV_data['junk'] = PV_data['Pack_Voltage_PP'].str.split('|', expand=True)[1]
    PV_data['Pack_Power(W)'] = PV_data['PP_PC'].str.split('|',expand=True)[0]
    PV_data['junk'] = PV_data['PP_PC'].str.split('|',expand=True)[1]
    
    #%% Remove all "junk" columns from all dataframes
    
    voltage_data=voltage_data.drop([col for col in voltage_data.columns if "junk" in col], axis=1)
    temperature_data=temperature_data.drop([col for col in temperature_data.columns if "junk" in col], axis=1)
    
    max_min_temp_data=max_min_temp_data.drop([col for col in max_min_temp_data.columns if "junk" in col], axis=1)
    max_min_V_data=max_min_V_data.drop([col for col in max_min_V_data.columns if "junk" in col], axis=1)
    PV_data=PV_data.drop([col for col in PV_data.columns if "junk" in col], axis=1)
    misc_details=misc_details.drop([col for col in misc_details.columns if "junk" in col], axis=1)
    energy_details=energy_details.drop([col for col in energy_details.columns if "junk" in col], axis=1)
    relay_details=relay_details.drop([col for col in relay_details.columns if "junk" in col], axis=1)
    amb_T=amb_T.drop([col for col in amb_T.columns if "junk" in col], axis=1)
    soc_data = soc_data.drop([col for col in soc_data.columns if "junk" in col], axis=1)
    
    #%% Data Cleaning and Processing - II
    voltage_data['Cell_'+str(v_no)]=voltage_data['Cell_'+str(v_no)].str.replace('|','',regex=True)
    PV_data['Pack_Current']=PV_data['Pack_Current'].str.replace('|','',regex=True)
    PV_data['Pack_Current']=PV_data['Pack_Current'].astype(str).str.extract(r'(\-{0,}\d{1,})')
    
    for x in v_names_1: 
        voltage_data[x]=voltage_data[x].astype(str).str.extract(r'(\d{1,})')
        voltage_data[x]=voltage_data[x].fillna(method='ffill')
        voltage_data[x]=voltage_data[x].astype(int)
        voltage_data=voltage_data[(voltage_data[x]>=2000) & (voltage_data[x]<=4300)]
    
    for x in t_names_1: 
        temperature_data[x]=temperature_data[x].astype(str).str.extract(r'(\d{1,})')
        
    PV_data['Pack_Current']=PV_data['Pack_Current'].astype(float)
    
    
    relay_details['Precharge']=relay_details['Precharge'].astype(str).str.extract(r'(\d{1,})')
    relay_details['Main_Relay_Status']=relay_details['Main_Relay_Status'].astype(str).str.extract(r'(\d{1,})')
    relay_details['Pre-discharge']=relay_details['Pre-discharge'].astype(str).str.extract(r'(\d{1,})')
    
    relay_details['Main_Relay_Status']=relay_details['Main_Relay_Status'].fillna(method='ffill')
    relay_details['Main_Relay_Status']=relay_details['Main_Relay_Status'].fillna(method='bfill')
    
    relay_details['Precharge']=relay_details['Precharge'].fillna(method='ffill')
    relay_details['Precharge']=relay_details['Precharge'].fillna(method='bfill')
    
    relay_details['Pre-discharge']=relay_details['Pre-discharge'].fillna(method='ffill')
    relay_details['Pre-discharge']=relay_details['Pre-discharge'].fillna(method='bfill')
    
    
    #%% Create Maximum, minimum, average and delV columns from the cell voltage data
    voltage_data.loc[:,'Cell_1':'Cell_'+str(v_no)] = voltage_data.loc[:,'Cell_1':'Cell_'+str(v_no)] * 0.001
    voltage_data['Mean_V']=voltage_data.loc[:,'Cell_1':'Cell_'+str(v_no)].mean(axis=1)
    voltage_data['Max_V']=voltage_data.loc[:,'Cell_1':'Cell_'+str(v_no)].max(axis=1)
    voltage_data['Max_V_Cell']=voltage_data.loc[:,'Cell_1':'Cell_'+str(v_no)].astype(int).idxmax(axis=1)
    voltage_data['Min_V']=voltage_data.loc[:,'Cell_1':'Cell_'+str(v_no)].min(axis=1)
    voltage_data['Min_V_Cell']=voltage_data.loc[:,'Cell_1':'Cell_'+str(v_no)].astype(int).idxmin(axis=1)
    
    voltage_data['delV'] = (voltage_data['Max_V']-voltage_data['Min_V'])*1000
        
    #%% Create Maximum, minimum, average and delT columns from the cell temperature data
    temperature_data.loc[:,'T1':'T'+str(t_no)]=temperature_data.loc[:,'T1':'T'+str(t_no)].astype(float)
    temperature_data.loc[:,'T1':'T'+str(t_no)]=temperature_data.loc[:,'T1':'T'+str(t_no)].fillna(method='ffill')
    
    temperature_data = temperature_data[temperature_data.loc[:,'T1':'T'+str(t_no)]>20]
    temperature_data = temperature_data[temperature_data.loc[:,'T1':'T'+str(t_no)]<80]
    
    temperature_data['Mean_T']=temperature_data.loc[:,'T1':'T'+str(t_no)].mean(axis=1)
    temperature_data['Max_T']=temperature_data.loc[:,'T1':'T'+str(t_no)].max(axis=1)
    temperature_data['Max_TCell']=temperature_data.loc[:,'T1':'T'+str(t_no)].astype(float).idxmax(axis=1)
    temperature_data['Min_T']=temperature_data.loc[:,'T1':'T'+str(t_no)].min(axis=1)
    temperature_data['Min_TCell']=temperature_data.loc[:,'T1':'T'+str(t_no)].astype(float).idxmin(axis=1)
    
    temperature_data['delT']=temperature_data['Max_T']-temperature_data['Min_T']
    
    #%% Data Cleaning and Processing - III
    
    energy_details['Pack_Ah_BMS']=energy_details['Pack_Ah_BMS'].astype(str).str.replace('0e','',regex=True).str.extract(r'(\-{0,}\d{1,}.\d{0,2})')
    energy_details['Pack_Energy_BMS']=energy_details['Pack_Energy_BMS'].astype(str).str.extract(r'(\-{0,}\d{1,6}.\d{0,2})')
    energy_details['Pack_Ah_BMS']=energy_details['Pack_Ah_BMS'].astype(float)
    energy_details['Pack_Energy_BMS']=energy_details['Pack_Energy_BMS'].astype(float)
    energy_details=energy_details[energy_details['Pack_Energy_BMS']<12000]
    
    energy_details['Pack_Ah_BMS']=energy_details['Pack_Ah_BMS'].fillna(method='ffill')
    energy_details['Pack_Energy_BMS']=energy_details['Pack_Energy_BMS'].fillna(method='ffill')
    
    # Uncomment next line if pack current data is in improper format:
    # PV_data['Pack_Current']=PV_data['Pack_Current'].astype(str).str.extract(r'(\-{0,}\d{1,}.\d{1,})')
    PV_data['Pack_Voltage']=PV_data['Pack_Voltage'].astype(str).str.extract(r'(\d{1,}.\d{1,})')
    PV_data['Pack_Current']=PV_data['Pack_Current'].astype(float)
    PV_data['Pack_Voltage']=PV_data['Pack_Voltage'].astype(float)
    PV_data['Pack_Voltage']=PV_data['Pack_Voltage']*0.001
    PV_data = PV_data[PV_data['Pack_Voltage']<100]
    
    energy_details['Premature_Cycles'] = energy_details['Premature_Cycles'].str.replace('|','',regex=True)
    
    PV_data = PV_data.drop(columns=['Pack_Voltage_PP', 'PP_PC'])
    amb_T = amb_T.drop(columns=['amb_RT'])
    amb_T['RT'] = amb_T['RT'].str.replace('|','',regex=True)
    amb_T['RT'] = amb_T['RT'].astype(float)
    amb_T['Ambient_Temperature'] = amb_T['Ambient_Temperature'].astype(float)
    amb_T= amb_T[amb_T['Ambient_Temperature']>0]
    amb_T = amb_T[amb_T['RT']>0]
    
    soc_data['Calibration_Current'] = soc_data['Calibration_Current'].replace('(NaN)', np.nan)
    soc_data['Calibration_Current'] = soc_data['Calibration_Current'].fillna(method = 'ffill')
    
    max_min_V_data['deltaV(mV)_BMS']=max_min_V_data['deltaV(mV)_BMS'].astype(str).str.extract(r'(\d{1,})')
    max_min_V_data['deltaV(mV)_BMS']=max_min_V_data['deltaV(mV)_BMS'].fillna(method='ffill')
    max_min_V_data['deltaV(mV)_BMS'] = max_min_V_data['deltaV(mV)_BMS'].astype(int)
    
    #%% Combine pack voltage, energy and ambient temperature dataframes for better visualization and understanding
    
    details=pd.DataFrame()
    details=pd.concat([PV_data,energy_details,amb_T],axis=1,ignore_index=True)
    p=list(PV_data.columns)
    q=list(energy_details.columns)
    r=list(amb_T.columns)
    p.extend(q)
    p.extend(r)
    details.columns=p
    
    # Since timestamp/Date and time is not available, we set the number of seconds approximately as one for every datapoint
    details['Time_in_sec_s']=1
    
    #%% Assign state values as 0 for charge, 1 for discharge and 2 for rest
    
    details['State']=np.nan
    details['Time_in_sec_s_cap']=np.where((details['Time_in_sec_s']>300),np.nan,details['Time_in_sec_s'])
    details['State']=np.where(details['Pack_Current']>0,0,details['State'])
    details['State']=np.where(details['Pack_Current']<0,1,details['State'])
    details['State']=np.where(details['Pack_Current']==0,2,details['State'])
    
    #%% Calculate capacity and energy of shared data using coulomb counting method
    
    details['Cap_inst']=details['Time_in_sec_s_cap']*abs(details['Pack_Current'])/3600
    details['Capacity_calculated']=details['Cap_inst'].groupby(details['State']).cumsum()
    details['Capacity_calculated_chg']=(details[details['State']==0]['Cap_inst']).cumsum()
    details['Capacity_calculated_dchg']=(details[details['State']==1]['Cap_inst']).cumsum()
    
    details['Capacity_calculated_chg']=details['Capacity_calculated_chg'].fillna(method='bfill')
    details['Capacity_calculated_dchg']=details['Capacity_calculated_dchg'].fillna(method='bfill')
    
    details['Energy_calculated']=details['Capacity_calculated']*details['Pack_Voltage']
    details['Energy_calculated_chg']=details['Capacity_calculated_chg']*details['Pack_Voltage']
    details['Energy_calculated_dchg']=details['Capacity_calculated_dchg']*details['Pack_Voltage']

    #%% Plot graphs, save them and add them in the word document
    # Pre-charge graph
    document.add_page_break()
    document.add_paragraph().add_run('Precharge:').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(range(len(relay_details)),relay_details['Precharge'])
    plt.grid(linestyle='dotted')
    plt.title('Precharge',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Precharge.png',dpi=1200)
    
    document.add_picture(img_path+'\\Precharge.png')
    
    #%% Pre-discharge graph
    document.add_paragraph().add_run('Pre-discharge:').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(range(len(relay_details)),relay_details['Pre-discharge'])
    plt.grid(linestyle='dotted')
    plt.title('Pre-discharge',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Pre-discharge.png',dpi=1200)
    
    document.add_picture(img_path+'\\Pre-discharge.png')
    
    #%% Main Relay graph
    document.add_paragraph().add_run('Main Relay:').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(range(len(relay_details)),relay_details['Main_Relay_Status'])
    plt.grid(linestyle='dotted')
    plt.title('Main Relay',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Main Relay.png',dpi=1200)
    
    document.add_picture(img_path+'\\Main Relay.png')
    
    #%% Pack Current
    document.add_paragraph().add_run('Pack Current:').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(range(len(PV_data)),PV_data['Pack_Current']*0.001)
    plt.ylabel('Current(A)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Pack Current',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Pack Current.png',dpi=1200)
    
    document.add_picture(img_path+'\\Pack Current.png')
    
    #%% BMS Pack Current and Calibration current comparison graph
    document.add_paragraph().add_run('Pack Current and Calibration Current:').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(range(len(PV_data)),PV_data['Pack_Current']*0.001, label = 'Pack Current')
    plt.plot(range(len(soc_data)),soc_data['Calibration_Current'], label = 'Calibration Current')
    plt.ylabel('Current(A)',fontweight='bold')
    plt.legend()
    plt.grid(linestyle='dotted')
    plt.title('Pack Current and Calibration Current',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\BMS_Tester and Calibration Current.png',dpi=1200)
    
    document.add_picture(img_path+'\\BMS_Tester and Calibration Current.png')
    
    #%% Tester pack current graph
    document.add_paragraph().add_run('Pack Current(Tester data):').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(range(len(tester_data)),tester_data['Current(A)'])
    plt.ylabel('Current(A)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Tester Current',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Pack Current_tester.png',dpi=1200)
    
    document.add_picture(img_path+'\\Pack Current_tester.png')
    
    #%% BMS pack capacity graph
    document.add_paragraph().add_run('Pack Capacity (BMS):').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(range(len(energy_details)),energy_details['Pack_Ah_BMS'])
    plt.ylabel('Capacity(Ah)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Pack Capacity (BMS)',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Pack Capacity_BMS.png',dpi=1200)
     
    document.add_picture(img_path+'\\Pack Capacity_BMS.png')
    document.add_paragraph('In this session, the capacity expended by the pack is about %sAh ()' %energy_details['Pack_Ah_BMS'].max() )
    
    #%% Tester capacity data processing
    # Due to pulse charging, the tester data capacity is calculated individually for each pulse
    # Hence, we accumulate the capacity to get proper estimate
    tester_data['Capacity_diff'] = tester_data['Capacity(Ah)'] - tester_data['Capacity(Ah)'].shift()
    tester_data['Capacity_diff'] = np.where(tester_data['Capacity_diff']<0, 0, 
                                            tester_data['Capacity_diff'])
    
    tester_data['Capacity_inst'] = tester_data['Capacity_diff'].groupby(tester_data['Step Type']).cumsum()
    
    tester_data['Capacity_cumulative'] = np.where(tester_data['Step Type']=='CCCV Chg',
                                                  tester_data['Capacity_inst']
                                                  ,tester_data['Capacity(Ah)'])
    
    tester_data['SoC_Tester'] = (tester_data['Capacity_cumulative']/210) * 100
    
    tester_data = tester_data.drop(columns=['Capacity_inst', 'Capacity_diff'])
    
    #%% Tester capacity graph
    document.add_paragraph().add_run('Pack Capacity (Tester data):').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(range(len(tester_data)), tester_data['Capacity(Ah)'])
    plt.ylabel('Capacity(Ah)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Pack Capacity (Tester data)',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Pack Capacity_tester.png',dpi=1200)
     
    document.add_picture(img_path+'\\Pack Capacity_tester.png')
    document.add_paragraph('In this session, the capacity expended by the pack, according to the tester measurement, is about %sAh ()' %tester_data['Capacity_cumulative'].max() )
    
    #%% BMS pack energy graph
    document.add_paragraph().add_run('Pack Energy (BMS):').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(range(len(energy_details)),energy_details['Pack_Energy_BMS'])
    plt.ylabel('Energy(Wh)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Pack Energy (BMS)',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Pack Energy_BMS.png',dpi=1200)
    
    document.add_picture(img_path+'\\Pack Energy_BMS.png')
    
    p_bms_d_engy=document.add_paragraph('The ')
    p_bms_d_engy.add_run('pack energy ').bold=True
    p_bms_d_engy.add_run('is about ')
    p_bms_d_engy.add_run('%s Wh.' % energy_details['Pack_Energy_BMS'].max()).bold=True
    
    #%% Tester energy data processing
    # Due to pulse charging, the tester data energy is calculated individually for each pulse
    # Hence, we accumulate the energy to get a proper estimate
    tester_data['Energy_diff'] = tester_data['Energy(Wh)'] - tester_data['Energy(Wh)'].shift()
    tester_data['Energy_diff'] = np.where(tester_data['Energy_diff']<0, 0, 
                                          tester_data['Energy_diff'])
    
    
    tester_data['Energy_inst'] = tester_data['Energy_diff'].groupby(tester_data['Step Type']).cumsum()
    
    tester_data['Energy_cumulative'] = np.where(tester_data['Step Type']=='CCCV Chg',
                                                  tester_data['Energy_inst'],
                                                  tester_data['Energy(Wh)'])
    
    
    tester_data = tester_data.drop(columns=['Energy_inst', 'Energy_diff'])
    
    #%% Tester pack energy graph
    document.add_paragraph().add_run('Pack Energy (Tester data):').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(range(len(tester_data)),tester_data['Energy(Wh)'])
    plt.ylabel('Energy(Wh)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Pack Energy (Tester)',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Pack Energy_tester.png',dpi=1200)
    
    document.add_picture(img_path+'\\Pack Energy_tester.png')
    
    p_bms_d_engy=document.add_paragraph('The ')
    p_bms_d_engy.add_run('pack energy, ').bold=True
    p_bms_d_engy.add_run('according to the tester observations, is about ')
    p_bms_d_engy.add_run('%s Wh.' % tester_data['Energy_cumulative'].max()).bold=True
    
    #%% EKF SoC graph
    document.add_paragraph().add_run('EKF SoC:').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(range(len(soc_data)),soc_data['EKF_SoC'])
    plt.ylabel('SoC(%)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('EKF SoC',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\EKF_SOC.png',dpi=1200)
    
    document.add_picture(img_path+'\\EKF_SOC.png')
    
    document.add_paragraph().add_run('CC SoC:').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(range(len(soc_data)),soc_data['CC_SoC'])
    plt.ylabel('SoC(%)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('CC SoC',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\CC_SOC.png',dpi=1200)
    
    document.add_picture(img_path+'\\CC_SOC.png')
    
    #%%  BMS Pack Voltage graph
    document.add_paragraph().add_run('Pack Voltage:').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(range(len(details)),details['Pack_Voltage'])
    plt.ylabel('Voltage(V)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Pack Voltage',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Pack Voltage.png',dpi=1200)
    
    document.add_picture(img_path+'\\Pack Voltage.png')
    pv=document.add_paragraph('At ')
    pv.add_run('discharge, ').bold=True
    pv.add_run('during the first cycle, the pack goes from to . The pack voltage at rest after discharge is . At ')
    pv.add_run('charge, ').bold=True
    pv.add_run('the pack goes from to . ')
    pv.add_run('In the end, the voltage stabilizes at V during the end rest.')
    
    #%% Tester Pack Voltage graph
    document.add_paragraph().add_run('Pack Voltage (Tester data):').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(range(len(tester_data)),tester_data['Voltage(V)'])
    plt.ylabel('Voltage(V)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Pack Voltage (Tester)',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Pack Voltage_tester.png',dpi=1200)
    
    document.add_picture(img_path+'\\Pack Voltage_tester.png')

    #%% Cell Voltage graph 
    """
    <*<*<*<*<*<*<*<*<*<*<*<Need to change this type of legend and color formatting>*>*>*>*>*>*>*>*>*>*>*>*>
    Need to make it readable and discernible
    need to default the saving process also
    """
    # This graph is not added in the word document
    # This is because the legend is too large to be
    # Make sure to plot the graph in the output "modified" excel file saved after running this code
    # Insert the graph from excel in the word document
    document.add_paragraph().add_run('Cell Voltage:').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(range(len(voltage_data)),voltage_data.loc[:,'Cell_1':'Cell_'+str(v_no)])
    plt.ylabel('Voltage(mV)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Cell Voltage',fontweight='bold')
    plt.legend(v_names_1)
    plt.tight_layout()
    plt.savefig(img_path+'\\Cell Voltage.png',dpi=1200)
    
    #%% Average voltage graph
    document.add_paragraph().add_run('Average Voltage:').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(range(len(voltage_data)),(voltage_data['Mean_V']))
    plt.ylabel('Voltage(V)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Average Voltage',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Average Voltage.png',dpi=1200)
    
    document.add_picture(img_path+'\\Average Voltage.png')
    
    av=document.add_paragraph('The average cell voltage during ')
    av.add_run('discharge ').bold=True
    av.add_run('goes from ')
    av.add_run('V to V. ').bold=True
    av.add_run('The average cell voltage at the rest after discharge is V. At ')
    av.add_run('charge').bold=True
    av.add_run(', the average cell voltage goes from ')
    av.add_run('V to V.').bold=True
    av.add_run('In the end rest, it stabilizes at %sV.' % (voltage_data['Mean_V'].iloc[-1]))
    
    #%% delV graph
    document.add_paragraph().add_run('Voltage Difference (delV):').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(range(len(max_min_V_data)),(max_min_V_data['deltaV(mV)_BMS']))
    plt.ylabel('Voltage(mV)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Voltage difference (delV)',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\delV_bms.png',dpi=1200)
    
    document.add_picture(img_path+'\\delV_bms.png')
    
    dv=document.add_paragraph('During discharge, the ')
    dv.add_run('maximum delV (difference between the highest and lowest cell voltage), ').bold=True
    dv.add_run('is about ')
    dv.add_run('mV, at the end of discharge. ').bold=True
    dv.add_run('At the rest after discharging, the delmV is mV. In the charging session, the maximum delV is mV. ')
    dv.add_run('The pack stabilizes at %smV delV at the end rest. '% (voltage_data['delV'].iloc[-1]))
    
    #%% Cell Temperature graph
    document.add_paragraph().add_run('Cell Temperature:').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(range(len(temperature_data)),temperature_data.loc[:,'T1':'T'+str(t_no)])
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.legend(t_names_1)
    plt.grid(linestyle='dotted')
    plt.title('Cell Temperature',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Cell Temperature.png',dpi=1200)
    
    ct=document.add_paragraph('The maximum temperature reached is around ')
    ct.add_run('%s'%(temperature_data.loc[:,'T1':'T'+str(t_no)].max(axis=1).max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True
    
    #%% Average Temperature graph
    document.add_paragraph().add_run('Average temperature:').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(range(len(temperature_data)),temperature_data['Mean_T'])
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Average Temperature',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Average Temperature.png',dpi=1200)
    
    document.add_picture(img_path+'\\Average Temperature.png')
    
    at=document.add_paragraph('As can be seen, the ')
    at.add_run('average cell temperature ').bold=True
    at.add_run('stays between ')
    at.add_run(u'\N{DEGREE SIGN}'+'C to '+u'\N{DEGREE SIGN}'+'C ').bold=True
    at.add_run('during ')
    at.add_run('discharging').bold=True
    at.add_run(' and ')
    at.add_run(u'\N{DEGREE SIGN}'+'C to '+u'\N{DEGREE SIGN}'+'C ').bold=True
    at.add_run('during ')
    at.add_run('charging.').bold=True
    
    #%% delT graph
    document.add_paragraph().add_run('Temperature Difference (delT):').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(range(len(temperature_data)),temperature_data['delT'])
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Temperature difference (delT)',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\delT.png',dpi=1200)
    
    document.add_picture(img_path+'\\delT.png')
    dt=document.add_paragraph('The ')
    dt.add_run('maximum delT ').bold=True
    dt.add_run('is about ')
    dt.add_run('%s' %temperature_data['delT'].max()).bold=True
    dt.add_run(u'\N{DEGREE SIGN}'+'C.').bold=True
    
    #%% Ambient Temperature graph
    document.add_paragraph().add_run('Ambient Temperature:').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(range(len(amb_T)),amb_T['Ambient_Temperature'])
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Ambient Temperature',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\AmbT.png',dpi=1200)
    
    document.add_picture(img_path+'\\AmbT.png')
    dt=document.add_paragraph('The ')
    dt.add_run('ambient temperature ').bold=True
    dt.add_run('is between ')
    dt.add_run(u'\N{DEGREE SIGN}'+'C ').bold=True
    dt.add_run('and ')
    dt.add_run(u'\N{DEGREE SIGN}'+'C.').bold=True
    
    #%% Relay Temperature graph
    document.add_paragraph().add_run('Relay Temperature:').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(range(len(amb_T)),amb_T['RT'])
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Relay Temperature',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\RT.png',dpi=1200)
    
    document.add_picture(img_path+'\\RT.png')
    dt=document.add_paragraph('The ')
    dt.add_run('relay temperature ').bold=True
    dt.add_run('is between ')
    dt.add_run(u'\N{DEGREE SIGN}'+'C ').bold=True
    dt.add_run('and ')
    dt.add_run(u'\N{DEGREE SIGN}'+'C.').bold=True
    
    #%% Fault status graph
    plt.figure(figsize=(l,b))
    plt.plot(range(len(misc_details)),misc_details['Fault_Status'])
    plt.ylabel('Fault Code',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Fault Status',fontweight='bold')
    plt.tight_layout()
    
    #%% Add final observation tabular column in the word document
    document.add_paragraph().add_run('Tabulation:').bold=True
    
    table2=document.add_table(rows=3,cols=22)
    
    table2.cell(0,0).text='Test Date'
    table2.cell(0,1).text='Step'
    table2.cell(0,2).text='Current'
    table2.cell(0,3).text='Pack Voltage'
    table2.cell(0,4).text='Capacity (BMS)'
    table2.cell(0,5).text='Capacity (Tester)'
    table2.cell(0,6).text='Energy (BMS)'
    table2.cell(0,7).text='Energy (Tester)'
    table2.cell(0,8).text='Tester SoC(%) range'
    table2.cell(0,9).text='BMS EKF SoC(%) range'
    table2.cell(0,10).text='BMS CC SoC(%) range'
    table2.cell(0,11).text='Average Cell Voltage'
    table2.cell(0,12).text='delV'
    table2.cell(0,14).text='Temperature'
    table2.cell(0,16).text='delT'
    table2.cell(0,18).text='Ambient Temperature'
    table2.cell(0,20).text='Relay Temperature'
    
    table2.cell(1,12).text='Start'
    table2.cell(1,13).text='End'
    table2.cell(1,14).text='Start'
    table2.cell(1,15).text='End'
    table2.cell(1,16).text='Start'
    table2.cell(1,17).text='End'
    table2.cell(1,18).text='Start'
    table2.cell(1,19).text='End'
    table2.cell(1,20).text='Start'
    table2.cell(1,21).text='End'

    table2.cell(2,1).text = step
    table2.cell(2,2).text = str((PV_data['Pack_Current'].max())*0.001)+'A/'+str((PV_data['Pack_Current'].min())*0.001) +'A'
    table2.cell(2,3).text = str(details['Pack_Voltage'].iloc[0]) + 'V to ' + str(details['Pack_Voltage'].iloc[-1]) +'V'
    table2.cell(2,4).text = str(energy_details['Pack_Ah_BMS'].max()) + 'Ah'
    table2.cell(2,5).text = str(tester_data['Capacity_cumulative'].max()) + 'Ah'
    table2.cell(2,6).text = str(energy_details['Pack_Energy_BMS'].max()) + 'Wh'
    table2.cell(2,7).text = str(tester_data['Energy_cumulative'].max()) + 'Wh'
    table2.cell(2,11).text = str(voltage_data['Mean_V'].iloc[0]) + 'V to ' + str(voltage_data['Mean_V'].iloc[-1]) +'V'
    table2.cell(2,12).text = str(voltage_data['delV'].iloc[0]) + 'mV'
    table2.cell(2,13).text = str(voltage_data['delV'].iloc[-1]) + 'mV'
    table2.cell(2,14).text = str(temperature_data['Min_T'].iloc[0]) + u'\N{DEGREE SIGN}'+'C'
    table2.cell(2,15).text = str(temperature_data['Max_T'].iloc[-1]) + u'\N{DEGREE SIGN}'+'C'
    table2.cell(2,16).text = str(temperature_data['delT'].iloc[0]) + u'\N{DEGREE SIGN}'+'C'
    table2.cell(2,17).text = str(temperature_data['delT'].iloc[-1]) + u'\N{DEGREE SIGN}'+'C'
    table2.cell(2,18).text = str(amb_T['Ambient_Temperature'].iloc[0]) + u'\N{DEGREE SIGN}'+'C'
    table2.cell(2,19).text = str(amb_T['Ambient_Temperature'].iloc[-1]) + u'\N{DEGREE SIGN}'+'C'
    table2.cell(2,20).text = str(amb_T['RT'].iloc[0]) + u'\N{DEGREE SIGN}'+'C'
    table2.cell(2,21).text = str(amb_T['RT'].iloc[-1]) + u'\N{DEGREE SIGN}'+'C'

    table2.style='Table Grid'
    
    #%% Save the word document in the destination folder
    path=img_path+'\\'+filename+'_modified.xlsx'
    document.save(img_path+'\\'+filename+'_observations_document.docx')
    
    #%% Save all the dataframes into an excel file as separate sheets
    writer = pd.ExcelWriter(path, engine = 'xlsxwriter')
    voltage_data.to_excel(writer, sheet_name = 'Cell Voltage',index=False)
    temperature_data.to_excel(writer, sheet_name = 'Cell Temperature',index=False)
    PV_data.to_excel(writer,sheet_name='Pack Voltage Details',index=False)
    relay_details.to_excel(writer,sheet_name='Relay Status',index=False)
    max_min_V_data.to_excel(writer,sheet_name='Voltage Max Min',index=False)
    max_min_temp_data.to_excel(writer,sheet_name='Temp Max Min',index=False)
    misc_details.to_excel(writer,sheet_name='Fault Status',index=False)
    energy_details.to_excel(writer,sheet_name='Energy Details',index=False)
    details.to_excel(writer,sheet_name='Pack Details',index=False) #DT
    amb_T.to_excel(writer, sheet_name = 'Ambient Temperature', index = False)
    soc_data.to_excel(writer, sheet_name = 'SoC', index = False)
    tester_data.to_excel(writer, sheet_name = 'Tester Data', index = False)
    writer.close()
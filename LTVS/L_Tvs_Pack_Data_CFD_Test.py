# -*- coding: utf-8 -*-
"""
Created on Fri May 31 17:27:37 2024

@author: IITM
"""

#%% Code Description

"""
USE THIS CODE IF:
    1.  
"""

"""
PRE-REQUISITES:
    1. 
"""

"""
NOTE:
    1. 

"""

#%% Import necessary libraries
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import docx
import matplotlib.colors as mcolors
from cycler import cycler
import os
import glob

#%% Get all the files with names that match the pattern of the conventional naming of the files that need processing
# Care must be taken that the files are saved with similar naming convention
files = glob.glob(r"D:\Benisha\LTVS\10kWh\CFD\20-4-24   tvs 10kwh pack drive testteraterm 1.xlsx")

# Define column names depending on the presence of date and time columns
# cols=['CV_CT_PC_tag','Cell 1','Cell 2','Cell 3','Cell 4','Cell 5','Cell 6','Cell 7','Cell 8','Cell 9','Cell 10','Cell 11','Cell 12','Cell 13','Cell 14','Cell 15','Cell 16','BR_ID_tag','BR_ID_th','tm','ts']
cols=['Date','Hours','Minutes','Seconds','CV_CT_PC_tag','Cell 1','Cell 2','Cell 3','Cell 4','Cell 5','Cell 6','Cell 7','Cell 8','Cell 9','Cell 10','Cell 11','Cell 12','Cell 13','Cell 14','Cell 15','Cell 16','BR_ID_tag','BR_ID_th','tm','ts']

# Specify sheet name of the input file, if necessary
# This is BMS data
# sheet="log2_raw" # Uncomment if needed

#%%
dvv=pd.DataFrame()
dtt=pd.DataFrame()
dr=pd.DataFrame()
de=pd.DataFrame()
dmv=pd.DataFrame()
dmt=pd.DataFrame()
dm=pd.DataFrame()
dd=pd.DataFrame()
dpv=pd.DataFrame()

# Specify number of cells in pack
v_no = 16
t_no = 16
#%% Loop through each of the files individually
for file in files:

    data = pd.read_excel(file,names=cols,index_col=False) # ,sheet_name=sheet # include this inside the paranthesis if needed
    
    # Get the destination folder and file name from the input file path
    img_path = file.rsplit('\\',1)[0]+'\\'+file.rsplit('\\',1)[1].rsplit('.',1)[0]
    filename = file.rsplit('\\',1)[1].rsplit('.',1)[0]
    
    # Create the destination folder if is not already there
    if not os.path.exists(img_path):
        os.makedirs(img_path)
    
    #%% Document object creation
    # The MS Word document object is created using the docx library
    document=docx.Document()
    document.styles['Normal'].font.name='Calibri'
    H1=document.add_heading()
    H1.alignment=1
    H1.add_run('LTVS Pack Data').font.size=docx.shared.Pt(18)
    
    # A table containing the pack specifications is added
    document.add_paragraph().add_run('\n\nPack Specifications:').bold=True
    
    table1=document.add_table(rows=5,cols=2)
    
    table1.cell(0,0).text='Chemistry'
    table1.cell(0,1).text='LFP'
    table1.cell(1,0).text='Pack Capacity'
    table1.cell(1,1).text='210Ah (2*105Ah)'
    table1.cell(2,0).text='Configuration'
    table1.cell(2,1).text='2P16S'
    table1.cell(3,0).text='UV cut-off'
    table1.cell(3,1).text='2.8V'
    table1.cell(4,0).text='OV cut-off'
    table1.cell(4,1).text='3.5V'
    
    table1.style='Table Grid'
    
    #%% Datetime column processing (if necessary)
    data[['S','mS']]=data['Seconds'].str.split('.',expand=True)
    data['Date']=data['Date'].str.replace('[','',regex=True)
    data['DateTime']=np.nan
    data['DateTime']=data['Date']+' '+data['Hours'].astype(str)+':'+data['Minutes'].astype(str)+':'+data['S']
    data=data.drop(['S','mS'],axis=1)
    
    # Since the data is usually huge, a portion of the data is sliced
    # This is to view the structure of the data logging, which will come in handy in the following section
    d1=data.head(1000)
    
    data['DateTime']=pd.to_datetime(data['DateTime'],format='%Y-%m-%d %H:%M:%S',errors='coerce')
    
    #%% Separate data into different dataframes
    
    # This is done as the data is in putty log format, which does not have data separated into proper columns
    # Note the values to which the "CV_CT_PC_tag" is checked against (eg, '|CV_1[1..16](mV)', '|Min_temp')
    # These values need to be checked if they are right, using the 'd1' dataframe in the previous section
    # Change the values if they are different in the input data file
    voltage_data=data[data['CV_CT_PC_tag']=='|CV_1[1..16](mV)'].dropna(how='all',axis=1).reset_index(drop=True)
    temperature_data=data[data['CV_CT_PC_tag']=='|Cell_Temperatures[1..16]'].dropna(how='all',axis=1).reset_index(drop=True)
    max_min_temp_data=data[data['CV_CT_PC_tag']=='|Min_temp'].dropna(how='all',axis=1).reset_index(drop=True) # CFD Test data
    max_min_V_data=data[data['CV_CT_PC_tag']=='|Max'].dropna(how='all',axis=1).reset_index(drop=True)
    PV_data=data[data['CV_CT_PC_tag']=='|PV(mV)'].dropna(how='all',axis=1).reset_index(drop=True)
    misc_details=data[data['CV_CT_PC_tag']=='|Fault'].dropna(how='all',axis=1).reset_index(drop=True)
    energy_details=data[data['CV_CT_PC_tag']=='|Pack_Energy'].dropna(how='all',axis=1).reset_index(drop=True)
    relay_details=data[data['CV_CT_PC_tag']=='|RLY'].dropna(how='all',axis=1).reset_index(drop=True)
    amb_T = data[data['CV_CT_PC_tag']=='AMB'].dropna(how='all',axis=1).reset_index(drop=True)
    
    #%%  Define Proper column names for each dataframe
    
    # The cell voltage and cell temperature values are large depending upon the number of cells
    # It is also quite repetitive, hence the naming process is automated here.
    v_names_1=[f'V{x}' for x in range(1,v_no + 1)] 
    t_names_1=[f'T{x}' for x in range(1,t_no + 1)]
    
    #%% If there is no datetime, uncomment this section
    
    # Change column names if necessary
    # Give the name "junk" to columns with useless data (like Nan or some symbol); It will be removed later
    # Changes often with each input dataset as the data logging process is usually irregular
    
    # mm_temp=['junk','Die_Temperature','junk','Min_T [Cell_No]','junk','junk','Max_T [Cell_No]','junk']
    # mm_v=['junk','junk','Max_Cell_V_BMS(mV)','Max_V_Cell_No_BMS','junk','junk','junk','Min_Cell_V_BMS(mV)','Min_V_Cell_No_BMS','junk','junk','deltaV(mV)_BMS','junk','Balancing_Status','junk','junk','Afe_1']
    # pv=['junk','Pack_Voltage_BMS(mV)','junk','junk','Battery_Voltage(mV)','junk','junk','Pack_Power(W)','junk','Flash_Usage(%)']
    # i_names=['mA_A_tag','Pack_Current','junk','junk','Raw_ADC','junk','Offset_Average','junk']
    # misc=['junk','junk','Fault_Status','junk','junk','junk','Latch_Counter','junk','Speed(kmph)','junk','junk','junk','junk','Speed_pulse_counter','junk']
    # engy=['junk','Pack_Energy_BMS','junk','junk','Pack_Ah_BMS','junk','junk','SoC(%)','junk','junk','SoH(%)','junk','junk','Cycles','junk','junk','Premature_Cycles','junk','BR_ID']
    # rly=['junk','junk','junk','Main_Relay_Status','junk','Precharge','junk','junk','Pre-discharge','junk']
    
    #%% If the date and time are present, this block needs to be uncommented
    
    # Change column names if necessary
    # Give the name "junk" to columns with useless data (like Nan or some symbol); It will be removed later
    # Changes often with each input dataset as the data logging process is usually irregular
    
    mm_temp=['junk','junk','junk','junk','junk','Min_T [Cell_No]','junk','junk','Max_T [Cell_No]','junk','Flash_Usage(%)','DateTime']
    mm_v=['junk','junk','junk','junk','junk','junk','Max_Cell_V_BMS(mV)','Max_V_Cell_No_BMS','junk','junk','junk','Min_Cell_V_BMS(mV)','Min_V_Cell_No_BMS','junk','deltaV(mV)_BMS','DateTime']# CFD Test Data
    pv=['junk','junk','junk','junk','junk','Pack_Voltage_PP','PP_PC','Pack_Current','DateTime']
    misc=['junk','junk','junk','junk','junk','junk','Fault_Status','junk','junk','junk','Latch_Counter','junk','DateTime']
    engy=['junk','junk','junk','junk','junk','Pack_Energy_BMS','junk','junk','Pack_Ah_BMS','junk','SoC(%)','junk','junk','SoH(%)','junk','Cycles','junk','junk','Premature_Cycles','DateTime']
    rly=['junk','junk','junk','junk','junk','junk','junk','Main_Relay_Status','junk','Precharge','junk','junk','Pre-discharge','junk','DateTime']
    amb = ['junk','junk','junk','junk','junk','amb_RT','RT','DateTime']
    
    #%% Uncomment if no datetime data is there
    
    # Column names for the voltage_data and temperature_data dataframes are properly defined
    # "junk" is added as the first element to represent the column 'CV_CT_PC_tag'
    # v_names/t_names is a copy as the list with only the cell number names will be used later to name the legend while plotting
  
    # v_names=v_names_1.copy()
    # v_names.insert(0,'junk')
    # t_names=t_names_1.copy()
    # t_names.insert(0,'junk')
    # t_names.append('junk')
    
    #%%  Uncomment only if date and timestamp is present in the input data
    # Column names for the voltage_data and temperature_data dataframes are properly defined
    # "junk" is added at first to represent the 'Date', 'Hours', 'minutes', 'seconds' and 'CV_CT_PC_tag'columns
    # v_names/t_names is a copy as the list with only the cell number names will be used later to name the legend while plotting
  
    v_names=v_names_1.copy()
    v_names.insert(0,'junk')
    v_names.insert(0,'junk')
    v_names.insert(0,'junk')
    v_names.insert(0,'junk')
    v_names.insert(0,'junk')
    v_names.append('DateTime')
    
    t_names=t_names_1.copy()
    t_names.append('junk')
    t_names.insert(0,'junk')
    t_names.insert(0,'junk')
    t_names.insert(0,'junk')
    t_names.insert(0,'junk')
    t_names.insert(0,'junk')
    t_names.append('DateTime')
    
    #%% Assign proper column names to the respective dataframes
    voltage_data.columns=v_names
    temperature_data.columns=t_names
    max_min_temp_data.columns=mm_temp
    max_min_V_data.columns=mm_v 
    PV_data.columns=pv
    misc_details.columns=misc
    energy_details.columns=engy
    relay_details.columns=rly
    amb_T.columns = amb
    
    #%% Data Cleaning and preparation
    amb_T['Ambient_Temperature'] = amb_T['amb_RT'].str.split('|',expand=True)[0]
    amb_T['junk'] = amb_T['amb_RT'].str.split('|',expand=True)[1]
    PV_data['Pack_Voltage'] = PV_data['Pack_Voltage_PP'].str.split('|', expand=True)[0]
    PV_data['junk'] = PV_data['Pack_Voltage_PP'].str.split('|', expand=True)[1]
    PV_data['Pack_Power(W)'] = PV_data['PP_PC'].str.split('|',expand=True)[0]
    PV_data['junk'] = PV_data['PP_PC'].str.split('|',expand=True)[1]
    
    voltage_data=voltage_data.drop([col for col in voltage_data.columns if "junk" in col], axis=1)
    temperature_data=temperature_data.drop([col for col in temperature_data.columns if "junk" in col], axis=1)
    
    max_min_temp_data=max_min_temp_data.drop([col for col in max_min_temp_data.columns if "junk" in col], axis=1)
    max_min_V_data=max_min_V_data.drop([col for col in max_min_V_data.columns if "junk" in col], axis=1)
    PV_data=PV_data.drop([col for col in PV_data.columns if "junk" in col], axis=1)
    misc_details=misc_details.drop([col for col in misc_details.columns if "junk" in col], axis=1)
    energy_details=energy_details.drop([col for col in energy_details.columns if "junk" in col], axis=1)
    relay_details=relay_details.drop([col for col in relay_details.columns if "junk" in col], axis=1)
    amb_T=amb_T.drop([col for col in amb_T.columns if "junk" in col], axis=1)
    
    voltage_data['V'+str(v_no)]=voltage_data['V'+str(v_no)].str.replace('|','',regex=True)
    PV_data['Pack_Current']=PV_data['Pack_Current'].str.replace('|','',regex=True)
    PV_data['Pack_Current']=PV_data['Pack_Current'].astype(str).str.extract(r'(\-{0,}\d{1,})')
    
    for x in v_names_1: 
        voltage_data[x]=voltage_data[x].astype(str).str.extract(r'(\d{1,})')
        voltage_data[x]=voltage_data[x].fillna(method='ffill')
        voltage_data[x]=voltage_data[x].astype(int)
        voltage_data=voltage_data[(voltage_data[x]>=2000) & (voltage_data[x]<=4300)]
    
    for x in t_names_1: 
        temperature_data[x]=temperature_data[x].astype(str).str.extract(r'(\d{1,})')
        
    #%% Calculate Maximum, Minimum, Average and delV from the cell voltage data
    voltage_data['Mean_V']=voltage_data.loc[:,'V1':'V'+str(v_no)].mean(axis=1)
    voltage_data['Max_V']=voltage_data.loc[:,'V1':'V'+str(v_no)].max(axis=1)
    voltage_data['Max_V_Cell']=voltage_data.loc[:,'V1':'V'+str(v_no)].astype(int).idxmax(axis=1)
    voltage_data['Min_V']=voltage_data.loc[:,'V1':'V'+str(v_no)].min(axis=1)
    voltage_data['Min_V_Cell']=voltage_data.loc[:,'V1':'V'+str(v_no)].astype(int).idxmin(axis=1)
    
    voltage_data['delV']=voltage_data['Max_V']-voltage_data['Min_V']
    
    #%% DT
    # voltage_data['ElapsedTime']=(voltage_data['DateTime']-voltage_data['DateTime'].iloc[0])
    
    #%% Calculate Maximum, Minimum, Average and delT from the cell temperature data
    temperature_data.loc[:,'T1':'T'+str(t_no)]=temperature_data.loc[:,'T1':'T'+str(t_no)].astype(float)
    temperature_data.loc[:,'T1':'T'+str(t_no)]=temperature_data.loc[:,'T1':'T'+str(t_no)].fillna(method='ffill')
    temperature_data['Mean_T']=temperature_data.loc[:,'T1':'T'+str(t_no)].mean(axis=1)
    temperature_data['Max_T']=temperature_data.loc[:,'T1':'T'+str(t_no)].max(axis=1)
    temperature_data['Max_TCell']=temperature_data.loc[:,'T1':'T'+str(t_no)].astype(float).idxmax(axis=1)
    temperature_data['Min_T']=temperature_data.loc[:,'T1':'T'+str(t_no)].min(axis=1)
    temperature_data['Min_TCell']=temperature_data.loc[:,'T1':'T'+str(t_no)].astype(float).idxmin(axis=1)
    
    temperature_data['delT']=temperature_data['Max_T']-temperature_data['Min_T']
    
    #%% DT
    # temperature_data['ElapsedTime']=(temperature_data['DateTime']-temperature_data['DateTime'].iloc[0])
    
    #%%Data cleaning and Preparation
    PV_data['Pack_Current'] = PV_data['Pack_Current'].astype(float)
    PV_data['Pack_Current'] = PV_data['Pack_Current']*0.001
    energy_details['Premature_Cycles'] = energy_details['Premature_Cycles'].str.replace('|','',regex=True)
    
    PV_data = PV_data.drop(columns=['Pack_Voltage_PP', 'PP_PC'])
    amb_T = amb_T.drop(columns=['amb_RT'])
    amb_T['RT'] = amb_T['RT'].str.replace('|','',regex=True)
    
    #%% DT
    # current_data['ElapsedTime']=(current_data['DateTime']-current_data['DateTime'].iloc[0])
    
    #%%
    relay_details['Precharge']=relay_details['Precharge'].astype(str).str.extract(r'(\d{1,})')
    relay_details['Main_Relay_Status']=relay_details['Main_Relay_Status'].astype(str).str.extract(r'(\d{1,})')
    relay_details['Pre-discharge']=relay_details['Pre-discharge'].astype(str).str.extract(r'(\d{1,})')
    
    relay_details['Main_Relay_Status']=relay_details['Main_Relay_Status'].fillna(method='ffill')
    relay_details['Main_Relay_Status']=relay_details['Main_Relay_Status'].fillna(method='bfill')
    
    relay_details['Precharge']=relay_details['Precharge'].fillna(method='ffill')
    relay_details['Precharge']=relay_details['Precharge'].fillna(method='bfill')
    
    relay_details['Pre-discharge']=relay_details['Pre-discharge'].fillna(method='ffill')
    relay_details['Pre-discharge']=relay_details['Pre-discharge'].fillna(method='bfill')
    
    energy_details['Pack_Ah_BMS']=energy_details['Pack_Ah_BMS'].astype(str).str.replace('0e','',regex=True).str.extract(r'(\-{0,}\d{1,}.\d{0,2})')
    energy_details['Pack_Energy_BMS']=energy_details['Pack_Energy_BMS'].astype(str).str.extract(r'(\-{0,}\d{1,}.\d{0,2})')
    
    energy_details['Pack_Ah_BMS']=energy_details['Pack_Ah_BMS'].fillna(method='ffill')
    energy_details['Pack_Energy_BMS']=energy_details['Pack_Energy_BMS'].fillna(method='ffill')
    
    # current_data['Pack_Current']=current_data['Pack_Current'].astype(str).str.extract(r'(\-{0,}\d{1,}.\d{1,})')
    PV_data['Pack_Voltage']=PV_data['Pack_Voltage'].astype(str).str.extract(r'(\d{1,}.\d{1,})')
    PV_data['Pack_Current']=PV_data['Pack_Current'].astype(float)
    PV_data['Pack_Voltage']=PV_data['Pack_Voltage'].astype(float)
    PV_data['Pack_Voltage']=PV_data['Pack_Voltage']*0.001
    
    #%% DT
    details=pd.DataFrame()
    details=pd.merge(left= PV_data, right= energy_details, on='DateTime') 
    p=list(PV_data.columns)
    r=list(energy_details.columns)
    r.remove('DateTime')

    p.extend(r)
    details.columns=p
    # Time difference between every datapoint
    details['Time_in_sec_s']=(details['DateTime']-details['DateTime'].shift(1))/np.timedelta64(1,'s')
    details['Time_in_sec']=(details['DateTime']-details['DateTime'].iloc[0])/np.timedelta64(1,'s')
    
    #%% Full
    complete_details = pd.DataFrame()
    complete_details = pd.merge(left = details, right = temperature_data, on = 'DateTime')
    complete_details = pd.merge(left = complete_details, right = amb_T, on = 'DateTime')
    q=list(details.columns)
    t=list(temperature_data.columns)
    t.remove('DateTime')
    q.extend(t)
    l=list(amb_T.columns)
    l.remove('DateTime')
    q.extend(l)
    
    complete_details.columns=q
    
    #%%
    # details=pd.DataFrame()
    # details=pd.concat([current_data,PV_data,energy_details],axis=1,ignore_index=True)
    # p=list(current_data.columns)
    # q=list(PV_data.columns)
    # r=list(energy_details.columns)
    # p.extend(q)
    # p.extend(r)
    # details.columns=p
    # details['Time_in_sec_s']=1
    #%%
    # details['Pack_Current']=details['Pack_Current'].astype(str).str.extract(r'(\-{0,}\d{1,}.\d{1,})')
    # details['Pack_Voltage']=details['Pack_Voltage_BMS(mV)'].astype(str).str.extract(r'(\d{1,}.\d{1,})')
    # details['Pack_Current']=details['Pack_Current'].astype(float)
    # details['Pack_Voltage']=details['Pack_Voltage'].astype(float)
    # details['Pack_Voltage']=details['Pack_Voltage']*0.001
    
    #%%
    details['State']=np.nan
    details['Time_in_sec_s_cap']=np.where((details['Time_in_sec_s']>300),np.nan,details['Time_in_sec_s'])
    details['State']=np.where(details['Pack_Current']>0,0,details['State']) # 1
    details['State']=np.where(details['Pack_Current']<0,1,details['State']) # 2
    details['State']=np.where(details['Pack_Current']==0,2,details['State'])
    
    #%%
    details['Cap_inst']=details['Time_in_sec_s_cap']*abs(details['Pack_Current'])/3600
    details['Capacity_calculated']=details['Cap_inst'].groupby(details['State']).cumsum()
    details['Capacity_calculated_chg']=(details[details['State']==0]['Cap_inst']).cumsum()
    details['Capacity_calculated_dchg']=(details[details['State']==1]['Cap_inst']).cumsum()
    
    details['Capacity_calculated_chg']=details['Capacity_calculated_chg'].fillna(method='bfill')
    details['Capacity_calculated_dchg']=details['Capacity_calculated_dchg'].fillna(method='bfill')
    
    details['Energy_calculated']=details['Capacity_calculated']*details['Pack_Voltage']
    details['Energy_calculated_chg']=details['Capacity_calculated_chg']*details['Pack_Voltage']
    details['Energy_calculated_dchg']=details['Capacity_calculated_dchg']*details['Pack_Voltage']
    
    #%%
    l=7.5
    b=4.5
    
    #%% Precharge graph
    document.add_page_break()
    document.add_paragraph().add_run('Precharge:').bold=True
    
    plt.figure(figsize=(l,b))
    
    plt.plot(range(len(relay_details)),relay_details['Precharge'])
    
    plt.grid(linestyle='dotted')
    plt.title('Precharge',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Precharge.png',dpi=1200)
    
    document.add_picture(img_path+'\\Precharge.png')
    
    #%% Pre-discharge graph
    document.add_paragraph().add_run('Pre-discharge:').bold=True
    
    plt.figure(figsize=(l,b))
    
    plt.plot(range(len(relay_details)),relay_details['Pre-discharge'])
    
    plt.grid(linestyle='dotted')
    plt.title('Pre-discharge',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Pre-discharge.png',dpi=1200)
    
    document.add_picture(img_path+'\\Pre-discharge.png')
    #%% Main Relay graph
    document.add_paragraph().add_run('Main Relay').bold=True
    
    plt.figure(figsize=(l,b))
    
    plt.plot(range(len(relay_details)),relay_details['Main_Relay_Status'])
    
    plt.grid(linestyle='dotted')
    plt.title('Main Relay',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Main Relay.png',dpi=1200)
    
    document.add_picture(img_path+'\\Main Relay.png')
    
    #%% Pack Current
    document.add_paragraph().add_run('Pack Current:').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(range(len(PV_data)),PV_data['Pack_Current'])
    
    plt.ylabel('Current(A)',fontweight='bold')
    
    plt.grid(linestyle='dotted')
    plt.title('Pack Current',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Pack Current.png',dpi=1200)
    
    document.add_picture(img_path+'\\Pack Current.png')
    
    #%% BMS Pack Capacity
    document.add_paragraph().add_run('Pack Capacity (BMS):').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(range(len(energy_details)),energy_details['Pack_Ah_BMS'])
    
    plt.ylabel('Capacity(Ah)',fontweight='bold')
    
    plt.grid(linestyle='dotted')
    plt.title('Pack Capacity (BMS)',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Pack Capacity_BMS.png',dpi=1200)
     
    document.add_picture(img_path+'\\Pack Capacity_BMS.png')
    document.add_paragraph('In this session, the capacity expended by the pack is about %sAh ()' % energy_details['Pack_Ah_BMS'].max() )
    
    #%% BMS Pack Energy
    document.add_paragraph().add_run('Pack Energy (BMS):').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(range(len(energy_details)),energy_details['Pack_Energy_BMS'])
    
    plt.ylabel('Energy(Wh)',fontweight='bold')
    
    plt.grid(linestyle='dotted')
    plt.title('Pack Energy (BMS)',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Pack Energy_BMS.png',dpi=1200)
    
    document.add_picture(img_path+'\\Pack Energy_BMS.png')
    p_bms_d_engy=document.add_paragraph('The ')
    p_bms_d_engy.add_run('discharging pack energy ').bold=True
    p_bms_d_engy.add_run('is about ')
    p_bms_d_engy.add_run('%s Wh.' % energy_details['Pack_Energy_BMS'].max()).bold=True
    
    #%% Pack Voltage
    document.add_paragraph().add_run('Pack Voltage:').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(range(len(details)),details['Pack_Voltage']) # DT
    
    plt.ylabel('Voltage(mV)',fontweight='bold')
    
    plt.grid(linestyle='dotted')
    plt.title('Pack Voltage',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Pack Voltage.png',dpi=1200)
    
    document.add_picture(img_path+'\\Pack Voltage.png')
    pv=document.add_paragraph('At ')
    pv.add_run('discharge, ').bold=True
    pv.add_run('during the first cycle, the pack goes from to . The pack voltage at rest after discharge is . At ')
    pv.add_run('charge, ').bold=True
    pv.add_run('the pack goes from to . ')
    pv.add_run('In the end, the voltage stabilizes at V during the end rest.')
    
    #%%
    document.add_paragraph().add_run('Cell Voltage:').bold=True
    
    plt.figure(figsize=(l,b))
    plt.rc('axes',prop_cycle=cycler('color',list(set(mcolors.XKCD_COLORS.values()))))
    plt.plot(range(len(voltage_data)),voltage_data.loc[:,'V1':'V'+str(v_no)])
    
    plt.ylabel('Voltage(mV)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Cell Voltage',fontweight='bold')
    plt.legend(v_names_1)
    plt.tight_layout()
    plt.savefig(img_path+'\\Cell Voltage.png',dpi=1200)
    
    #%%
    document.add_paragraph().add_run('Average Voltage:').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(range(len(voltage_data)),(voltage_data['Mean_V'])) #*0.0001
    
    plt.ylabel('Voltage(mV)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Average Voltage',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Average Voltage.png',dpi=1200)
    
    #%%
    document.add_picture(img_path+'\\Average Voltage.png')
    av=document.add_paragraph('The average cell voltage during ')
    av.add_run('discharge ').bold=True
    av.add_run('goes from ')
    av.add_run('mV to mV. ').bold=True
    av.add_run('The average cell voltage at the rest after discharge is mV. At ')
    av.add_run('charge').bold=True
    av.add_run(', the average cell voltage goes from ')
    av.add_run('mV to mV.').bold=True
    av.add_run('In the end rest, it  stabilizes at %smV.' % (voltage_data['Mean_V'].iloc[-1]))
    
    #%%
    # document.add_paragraph().add_run('Calculated Voltage Difference (delV):').bold=True
    # plt.figure(figsize=(l,b))
    # plt.plot(range(len(voltage_data)),(voltage_data['delV']))
    
    # plt.ylabel('Voltage(mV)',fontweight='bold')
    # plt.grid(linestyle='dotted')
    # plt.title('Voltage difference (delV)',fontweight='bold')
    # plt.tight_layout()
    # plt.savefig(img_path+'\\delV.png',dpi=1200)
    
    # document.add_picture(img_path+'\\delV.png')
    # dv=document.add_paragraph('During discharge, the ')
    # dv.add_run('maximum delV (difference between the highest and lowest cell voltage), ').bold=True
    # dv.add_run('is about ')
    # dv.add_run('mV, at the end of discharge. ').bold=True
    # dv.add_run('At the rest after discharging, the delmV is mV. In the charging session, the maximum delV is mV. ')
    # dv.add_run('The pack stabilizes at %smV delV at the end rest. '% (voltage_data['delV'].iloc[-1]))
    
    #%%
    max_min_V_data['deltaV(mV)_BMS']=max_min_V_data['deltaV(mV)_BMS'].astype(str).str.extract(r'(\d{1,})')
    max_min_V_data['deltaV(mV)_BMS']=max_min_V_data['deltaV(mV)_BMS'].fillna(method='ffill')
    document.add_paragraph().add_run('BMS Voltage Difference (delV):').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(range(len(max_min_V_data)),(max_min_V_data['deltaV(mV)_BMS']))
    
    plt.ylabel('Voltage(mV)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Voltage difference (delV)',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\delV_bms.png',dpi=1200)
    
    document.add_picture(img_path+'\\delV_bms.png')
    dv=document.add_paragraph('During discharge, the ')
    dv.add_run('maximum delV (difference between the highest and lowest cell voltage), ').bold=True
    dv.add_run('is about ')
    dv.add_run('mV, at the end of discharge. ').bold=True
    dv.add_run('At the rest after discharging, the delmV is mV. In the charging session, the maximum delV is mV. ')
    dv.add_run('The pack stabilizes at %smV delV at the end rest. '% (voltage_data['delV'].iloc[-1]))
    #%%
    # max_min_V_data['Balancing_Status']=max_min_V_data['Balancing_Status'].fillna(method='ffill')
    # document.add_paragraph().add_run('Balancing Status:').bold=True
    # plt.figure(figsize=(l,b))
    # plt.plot(range(len(max_min_V_data)),(max_min_V_data['Balancing_Status']))
    
    # plt.grid(linestyle='dotted')
    # plt.title('Balancing Status',fontweight='bold')
    # plt.tight_layout()
    # plt.savefig(img_path+'\\Balancing_Status.png',dpi=1200)
    
    # document.add_picture(img_path+'\\Balancing_Status.png')
    
    #%%
    document.add_paragraph().add_run('Cell Temperature:').bold=True
    
    plt.figure(figsize=(l,b))
    plt.rc('axes',prop_cycle=cycler('color',list(set(mcolors.XKCD_COLORS.values()))))
    plt.plot(range(len(temperature_data)),temperature_data.loc[:,'T1':'T'+str(t_no)])
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.legend(t_names_1)
    plt.grid(linestyle='dotted')
    plt.title('Cell Temperature',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Cell Temperature.png',dpi=1200)
    
    # document.add_picture(img_path+'\\Cell Temperature.png')
    ct=document.add_paragraph('The maximum temperature reached is around ')
    ct.add_run('%s'%(temperature_data.loc[:,'T1':'T'+str(t_no)].max(axis=1).max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True
    
    #%%
    document.add_paragraph().add_run('Average temperature:').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(range(len(temperature_data)),temperature_data['Mean_T'])
    
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Average Temperature',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Average Temperature.png',dpi=1200)
    
    document.add_picture(img_path+'\\Average Temperature.png')
    at=document.add_paragraph('As can be seen, the ')
    at.add_run('average cell temperature ').bold=True
    at.add_run('stays between ')
    at.add_run(u'\N{DEGREE SIGN}'+'C to '+u'\N{DEGREE SIGN}'+'C ').bold=True
    at.add_run('during ')
    at.add_run('discharging').bold=True
    at.add_run(' and ')
    at.add_run(u'\N{DEGREE SIGN}'+'C to '+u'\N{DEGREE SIGN}'+'C ').bold=True
    at.add_run('during ')
    at.add_run('charging.').bold=True
    
    #%%
    document.add_paragraph().add_run('Temperature Difference (delT):').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(range(len(temperature_data)),temperature_data['delT'])
    
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Temperature difference (delT)',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\delT.png',dpi=1200)
    
    document.add_picture(img_path+'\\delT.png')
    dt=document.add_paragraph('The ')
    dt.add_run('maximum delT ').bold=True
    dt.add_run('is about ')
    dt.add_run('%s' %temperature_data['delT'].max()).bold=True
    dt.add_run(u'\N{DEGREE SIGN}'+'C.').bold=True
    
    #%%
    document.add_paragraph().add_run('Max T:').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(range(len(temperature_data)),temperature_data['Max_T'])
    
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Maximum Temperature',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\MaxT.png',dpi=1200)
    
    document.add_picture(img_path+'\\MaxT.png')
    dt=document.add_paragraph('The ')
    dt.add_run('maximum Temperature ').bold=True
    dt.add_run('is about ')
    dt.add_run(u'\N{DEGREE SIGN}'+'C.').bold=True
    
    #%%
    # max_min_temp_data['Die_Temperature']=max_min_temp_data['Die_Temperature'].astype(str).str.extract(r'(\d{1,})')
    # max_min_temp_data['Die_Temperature']=max_min_temp_data['Die_Temperature'].astype(float)
    # document.add_paragraph().add_run('Die(Microcontroller Chip) Temperature:').bold=True
    # plt.figure(figsize=(l,b))
    # plt.plot(range(len(max_min_temp_data)),max_min_temp_data['Die_Temperature'])
    
    # plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    # plt.grid(linestyle='dotted')
    # plt.title('Die Temperature',fontweight='bold')
    # plt.tight_layout()
    # plt.savefig(img_path+'\\dieT.png',dpi=1200)
    
    # document.add_picture(img_path+'\\dieT.png')
    # dt=document.add_paragraph('The ')
    # dt.add_run('maximum die Temperature reached ').bold=True
    # dt.add_run('is about ')
    # dt.add_run('%s' %max_min_temp_data['Die_Temperature'].max()).bold=True
    # dt.add_run(u'\N{DEGREE SIGN}'+'C.').bold=True
    #%%
    path=img_path+'\\'+filename+'_modified.xlsx'
    document.save(img_path+'\\'+filename+'_observations_document.docx')
    
    #%%
    writer = pd.ExcelWriter(path, engine = 'xlsxwriter')
    complete_details.to_excel(writer, sheet_name = 'CFD Analysis Data',index=False)
    voltage_data.to_excel(writer, sheet_name = 'Cell Voltage',index=False)
    temperature_data.to_excel(writer, sheet_name = 'Cell Temperature',index=False)
    PV_data.to_excel(writer,sheet_name='Pack Voltage Details',index=False)
    relay_details.to_excel(writer,sheet_name='Relay Status',index=False)
    max_min_V_data.to_excel(writer,sheet_name='Voltage Max Min',index=False)
    max_min_temp_data.to_excel(writer,sheet_name='Temp Max Min',index=False)
    misc_details.to_excel(writer,sheet_name='Fault Status',index=False)
    energy_details.to_excel(writer,sheet_name='Energy Details',index=False)
    details.to_excel(writer,sheet_name='Pack Details',index=False) #DT
    amb_T.to_excel(writer, sheet_name = 'Ambient Temperature', index = False)
    writer.close()
    
    #%%
    dvv=pd.concat([dvv,voltage_data])
    dtt=pd.concat([dtt,temperature_data])
    dpv=pd.concat([dpv,PV_data])
    dr=pd.concat([dr,relay_details])
    de=pd.concat([de,energy_details])
    dmv=pd.concat([dmv,max_min_V_data])
    dmt=pd.concat([dmt,max_min_temp_data])
    dm=pd.concat([dm,misc_details])
    dd=pd.concat([dd,details])
    
    #%% 
    f=open(img_path+'\\'+filename+'_observations.txt',"w")
    vol_col_check=v_names_1
    temp_col_check=t_names_1
    
    print('======================================================================',file=f)
    print('Datapoints:%s' % (len(voltage_data)),file=f)
    print('Maximum Charging Current:%s ; Minimum Charging Current:%s' % (details[details['State']==0]['Pack_Current'].max(),details[details['State']==0]['Pack_Current'].min()),file=f)
    print('Maximum Discharging Current:%s ; Minimum Discharging Current:%s' % (details[details['State']==1]['Pack_Current'].max(),details[details['State']==1]['Pack_Current'].min()),file=f)
    print('Maximum Charging Capacity:%s ; Minimum Charging Capacity:%s' % (details['Capacity_calculated_chg'].max(),details['Capacity_calculated_chg'].min()),file=f)
    print('Maximum Discharging Capacity:%s ; Minimum Discharging Capacity:%s' % (details['Capacity_calculated_dchg'].max(),details['Capacity_calculated_dchg'].min()),file=f)
    print('Maximum Charging Energy:%s ; Minimum Charging Energy:%s' % (details['Energy_calculated_chg'].max(),details['Energy_calculated_chg'].min()),file=f)
    print('Maximum Discharging Energy:%s ; Minimum Discharging Energy:%s' % (details['Energy_calculated_dchg'].max(),details['Energy_calculated_dchg'].min()),file=f)
    print('Maximum delT:%s degC' %(temperature_data['delT'].max()),file=f)
    print('======================================================================',file=f)
    
    f.close()
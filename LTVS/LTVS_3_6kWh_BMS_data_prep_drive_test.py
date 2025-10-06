# -*- coding: utf-8 -*-
"""
Created on Wed May  8 18:36:06 2024

@author: IITM
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import docx
import os

import glob


v_no=17
t_no=17


file=r"D:\Benisha\LTVS\3.6kWh\Drive_Test\Phase_2\30-10-2023.xlsx"
data=pd.read_excel(file,index_col=False,sheet_name='bms_data') #,sheet_name=sheet ,sheet_name='chg and dchg 20-09-2023 (2)'
img_path=file.rsplit('\\',1)[0]+'\\'+file.rsplit('\\',1)[1].rsplit('.',1)[0]
filename=file.rsplit('\\',1)[1].rsplit('.',1)[0]

if not os.path.exists(img_path):
    os.makedirs(img_path)


# for i in data.columns:
#     data[i]=data[i].str.replace('|','')
#%% check
document=docx.Document()
document.styles['Normal'].font.name='Calibri'
H1=document.add_heading()
H1.alignment=1
H1.add_run('LTVS Pack Data').font.size=docx.shared.Pt(18)
# H1.font.color='black'
document.add_paragraph().add_run('\n\nPack Specifications:').bold=True

table1=document.add_table(rows=5,cols=2)

table1.cell(0,0).text='Chemistry'
table1.cell(0,1).text='LFP'
table1.cell(1,0).text='Pack Capacity'
table1.cell(1,1).text='72Ah (2*36Ah)'
table1.cell(2,0).text='Configuration'
table1.cell(2,1).text='2P16S'
table1.cell(3,0).text='UV cut-off'
table1.cell(3,1).text='2.8V'
table1.cell(4,0).text='OV cut-off'
table1.cell(4,1).text='3.5V'

table1.style='Table Grid'
d1=data.head(1000)
v_names_1=[f'V{x}' for x in range(1,v_no)] 
t_names_1=[f'T{x}' for x in range(1,t_no)]


#%%
data['End_Date']=data['End_Date'].str.replace('Date: ','')
data['End_Date']=data['End_Date'].str.replace('Time:','')
data['Start_Date']=data['Start_Date'].str.replace('Date: ','')
data['Start_Date']=data['Start_Date'].str.replace('Time:','')

#%%
data['Start_Date']=pd.to_datetime(data['Start_Date'],format='%d/%m/%Y %H:%M:%S')
data['End_Date']=pd.to_datetime(data['End_Date'],format='%d/%m/%Y %H:%M:%S')

#%%
data['DateTime']=data['DateTime'].str.replace('Date: ','')
data['DateTime']=data['DateTime'].str.replace('Time:','')

#%%
data['DateTime']=pd.to_datetime(data['DateTime'],format='%d/%m/%Y %H:%M:%S',errors='coerce')
data['DateTime']=data['DateTime'].fillna(method='ffill')
#%%
time=(data['End_Date'].iloc[0]-data['Start_Date'].iloc[0])/np.timedelta64(1,'s')
# data['Pack_Current']=data['Pack_Current'].str.replace('|','',regex=True)
data['Pack_Current']=data['PC(mA)']

for x in v_names_1: 
    data[x]=data[x].astype(str).str.extract(r'(\d{1,})')
    data[x]=data[x].fillna(method='ffill')
    data[x]=data[x].astype(int)
    data=data[(data[x]>=2000) & (data[x]<=4300)]

#%% 

# data.loc[:,'V1':'V'+str(v_no-1)]=data.loc[:,'V1':'V'+str(v_no-1)].fillna(method='ffill')
# data.loc[:,'V1':'V'+str(v_no-1)]=data.loc[:,'V1':'V'+str(v_no-1)].astype(int)
data['Mean_V']=data.loc[:,'V1':'V'+str(v_no-1)].mean(axis=1)
data['Max_V']=data.loc[:,'V1':'V'+str(v_no-1)].max(axis=1)
data['Max_V_Cell']=data.loc[:,'V1':'V'+str(v_no-1)].astype(int).idxmax(axis=1)
data['Min_V']=data.loc[:,'V1':'V'+str(v_no-1)].min(axis=1)
data['Min_V_Cell']=data.loc[:,'V1':'V'+str(v_no-1)].astype(int).idxmin(axis=1)

data['delV']=data['Max_V']-data['Min_V']

#%% DT
# data['ElapsedTime']=(data['DateTime']-data['DateTime'].iloc[0])


#%% 
data['Pack_Current']=data['Pack_Current'].astype(float)

#%%
# data['Pack_Current']=data['Pack_Current'].astype(str).str.extract(r'(\-{0,}\d{1,}.\d{1,})')
data['Pack_Voltage']=data['Pack_Voltage'].astype(str).str.extract(r'(\d{1,}.\d{1,})')
data['Pack_Current']=data['Pack_Current'].astype(float)
data['Pack_Current']=data['Pack_Current']*0.001

data['Pack_Voltage']=data['Pack_Voltage'].astype(float)

data['Time_in_sec_s']=time/len(data)
data['Time_in_sec_s']=(data['DateTime']-data['DateTime'].shift(1))/np.timedelta64(1,'s')
#%%
# data['Pack_Current']=data['Pack_Current'].astype(str).str.extract(r'(\-{0,}\d{1,}.\d{1,})')
# data['Pack_Voltage']=data['Pack_Voltage_BMS(mV)'].astype(str).str.extract(r'(\d{1,}.\d{1,})')
# data['Pack_Current']=data['Pack_Current'].astype(float)
# data['Pack_Voltage']=data['Pack_Voltage'].astype(float)
# data['Pack_Voltage']=data['Pack_Voltage']*0.001

#%%
data['State']=np.nan
data['Time_in_sec_s_cap']=np.where((data['Time_in_sec_s']>300),np.nan,data['Time_in_sec_s'])
data['State']=np.where(data['Pack_Current']>0,0,data['State']) # 1
data['State']=np.where(data['Pack_Current']<0,1,data['State']) # 2
data['State']=np.where(data['Pack_Current']==0,2,data['State'])

#%%
data['Cap_inst']=data['Time_in_sec_s_cap']*abs(data['Pack_Current'])/3600
data['Capacity_calculated']=data['Cap_inst'].groupby(data['State']).cumsum()
data['Capacity_calculated_chg']=(data[data['State']==0]['Cap_inst']).cumsum()
data['Capacity_calculated_dchg']=(data[data['State']==1]['Cap_inst']).cumsum()

data['Capacity_calculated_chg']=data['Capacity_calculated_chg'].fillna(method='bfill')
data['Capacity_calculated_dchg']=data['Capacity_calculated_dchg'].fillna(method='bfill')

data['Energy_calculated']=data['Capacity_calculated']*data['Pack_Voltage']
data['Energy_calculated_chg']=data['Capacity_calculated_chg']*data['Pack_Voltage']
data['Energy_calculated_dchg']=data['Capacity_calculated_dchg']*data['Pack_Voltage']

data['Capacity_calculated']=np.where(data['Capacity_calculated']<0.05,np.nan,data['Capacity_calculated'])
data['Capacity_calculated']=data['Capacity_calculated'].fillna(method='ffill')
data['Capacity_calclated']=data['Capacity_calculated'].fillna(method='bfill')

data['Energy_calculated']=np.where(data['Energy_calculated']<0.05,np.nan,data['Energy_calculated'])
data['Energy_calculated']=data['Energy_calculated'].fillna(method='ffill')
data['Energy_calculated']=data['Energy_calculated'].fillna(method='bfill')
#%%
l=7
b=4

#%%
document.add_page_break()

#%%
document.add_paragraph().add_run('Pack Current:').bold=True
plt.figure(figsize=(l,b))
plt.plot(range(len(data)),data['Pack_Current'])

plt.ylabel('Current(A)',fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('Pack Current',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Pack Current.png',dpi=1200)

document.add_picture(img_path+'\\Pack Current.png')


#%%
document.add_paragraph().add_run('Pack Voltage:').bold=True

plt.figure(figsize=(l,b))
plt.plot(range(len(data)),data['Pack_Voltage']) # DT

plt.ylabel('Voltage(mV)',fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('Pack Voltage',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Pack Voltage.png',dpi=1200)

#%%
document.add_picture(img_path+'\\Pack Voltage.png')
pv=document.add_paragraph('At ')
pv.add_run('discharge, ').bold=True
pv.add_run('during the first cycle, the pack goes from to . The pack voltage at rest after discharge is . At ')
pv.add_run('charge, ').bold=True
pv.add_run('the pack goes from to . ')
pv.add_run('In the end, the voltage stabilizes at V during the end rest.')

#%%
document.add_paragraph().add_run('Cell Voltage:').bold=True

#%%
plt.figure(figsize=(l,b))
plt.plot(range(len(data)),data.loc[:,'V1':'V'+str(v_no-1)])


plt.ylabel('Voltage(mV)',fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Cell Voltage',fontweight='bold')
plt.legend(v_names_1)
plt.tight_layout()
plt.savefig(img_path+'\\Cell Voltage.png',dpi=1200)

#%%
document.add_paragraph().add_run('Average Voltage:').bold=True

plt.figure(figsize=(l,b))
plt.plot(range(len(data)),(data['Mean_V'])) #*0.0001

plt.ylabel('Voltage(mV)',fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Average Voltage',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Average Voltage.png',dpi=1200)

#%%
document.add_picture(img_path+'\\Average Voltage.png')
av=document.add_paragraph('The average cell voltage during ')
av.add_run('discharge ').bold=True
av.add_run('goes from ')
av.add_run('mV to mV. ').bold=True
av.add_run('The average cell voltage at the rest after discharge is mV. At ')
av.add_run('charge').bold=True
av.add_run(', the average cell voltage goes from ')
av.add_run('mV to mV.').bold=True
av.add_run('In the end rest, it  stabilizes at %smV.' % (data['Mean_V'].iloc[-1]))

#%%
document.add_paragraph().add_run('Calculated Voltage Difference (delV):').bold=True
plt.figure(figsize=(l,b))
plt.plot(range(len(data)),(data['delV']))

plt.ylabel('Voltage(mV)',fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Voltage difference (delV)',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\delV.png',dpi=1200)

document.add_picture(img_path+'\\delV.png')
dv=document.add_paragraph('During discharge, the ')
dv.add_run('maximum delV (difference between the highest and lowest cell voltage), ').bold=True
dv.add_run('is about ')
dv.add_run('mV, at the end of discharge. ').bold=True
dv.add_run('At the rest after discharging, the delmV is mV. In the charging session, the maximum delV is mV. ')
dv.add_run('The pack stabilizes at %smV delV at the end rest. '% (data['delV'].iloc[-1]))





#%%
document.add_paragraph().add_run('Ambient temperature:').bold=True
plt.figure(figsize=(l,b))
plt.plot(range(len(data)),data['Ambient Temperature'])

plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Ambient Temperature',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Ambient Temperature.png',dpi=1200)

document.add_picture(img_path+'\\Ambient Temperature.png')
at=document.add_paragraph('As can be seen, the ')
at.add_run('ambient temperature ').bold=True
at.add_run('stays between ')
at.add_run(u'\N{DEGREE SIGN}'+'C to '+u'\N{DEGREE SIGN}'+'C ').bold=True





#%%
# data['Die_Temperature']=data['Die_Temperature'].astype(str).str.extract(r'(\d{1,})')
# data['Die_Temperature']=data['Die_Temperature'].astype(float)
document.add_paragraph().add_run('MOSFET Temperature:').bold=True
plt.figure(figsize=(l,b))
plt.plot(range(len(data)),data['Charge FET Temperature'],label='Chg_MOSFET')
plt.plot(range(len(data)),data['Discharge FET Temperature'],label='DChg_MOSFET')

plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('MOSFET Temperature',fontweight='bold')
plt.legend()
plt.tight_layout()
plt.savefig(img_path+'\\fetT.png',dpi=1200)

document.add_picture(img_path+'\\fetT.png')
dt=document.add_paragraph('The ')
dt.add_run('maximum charging MOSFET Temperature reached ').bold=True
dt.add_run('is about ')
dt.add_run('%s' %data['Charge FET Temperature'].max()).bold=True
dt.add_run(u'\N{DEGREE SIGN}'+'C.').bold=True
dt.add_run('The ')
dt.add_run('maximum discharging MOSFET Temperature reached ').bold=True
dt.add_run('is about ')
dt.add_run('%s' %data['Discharge FET Temperature'].max()).bold=True
dt.add_run(u'\N{DEGREE SIGN}'+'C.').bold=True
#%%
path=img_path+'\\'+filename+'_modified.xlsx'
document.save(img_path+'\\'+filename+'_observations_document.docx')

data.to_excel(img_path+'\\'+filename+'_observations.xlsx',index=False)


#%% 

f=open(img_path+'\\'+filename+'_observations.txt',"w")


print('======================================================================',file=f)
print('Datapoints:%s' % (len(data)),file=f)
print('Maximum Charging Current:%s ; Minimum Charging Current:%s' % (data[data['State']==0]['Pack_Current'].max(),data[data['State']==0]['Pack_Current'].min()),file=f)
print('Maximum Discharging Current:%s ; Minimum Discharging Current:%s' % (data[data['State']==1]['Pack_Current'].max(),data[data['State']==1]['Pack_Current'].min()),file=f)
print('Maximum Charging Capacity:%s ; Minimum Charging Capacity:%s' % (data['Capacity_calculated_chg'].max(),data['Capacity_calculated_chg'].min()),file=f)
print('Maximum Discharging Capacity:%s ; Minimum Discharging Capacity:%s' % (data['Capacity_calculated_dchg'].max(),data['Capacity_calculated_dchg'].min()),file=f)
print('Maximum Charging Energy:%s ; Minimum Charging Energy:%s' % (data['Energy_calculated_chg'].max(),data['Energy_calculated_chg'].min()),file=f)
print('Maximum Discharging Energy:%s ; Minimum Discharging Energy:%s' % (data['Energy_calculated_dchg'].max(),data['Energy_calculated_dchg'].min()),file=f)
print('======================================================================',file=f)

f.close()


# -*- coding: utf-8 -*-
"""
Created on Mon Apr 15 11:14:54 2024

@author: IITM
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import docx
import os
import glob
import datetime

v_no=103 # no+1
t_no=69 # no+1
mod_no=18 # no+1

# file=r"D:\Benisha\Andaman\In-house_testing_data\Bank1\07_march_Chag_ banks B1_parallel test.xlsx"
# file=r"D:\Benisha\Andaman\In-house_testing_data\Bank1\19_April_Bank1_series with deltal Discharge.xlsx"
# files=glob.glob(r"D:\Benisha\Andaman\In-house_testing_data\Bank1\*.xlsx")
# files=glob.glob(r"D:\Benisha\Andaman\In-house_testing_data\Bank2\*.xlsx")
files=glob.glob(r"D:\Benisha\Andaman\In-house_testing_data\Bank2\Raw Files\*.xlsx")

f=open(r"D:\Benisha\Andaman\Andaman_complete_observations_b2_one_file.txt","w")

for file in files:

    df=pd.read_excel(file)
#%%
    df.iloc[:,0]=df.iloc[:,0].str.replace('[','',regex=True)
    df.iloc[:,1]=df.iloc[:,1].str.replace(']','',regex=True)
    
    df.iloc[:,1]=df.iloc[:,1].astype(str)
    df[['Time_format','string']]=df.iloc[:,1].str.split('.',expand=True)
    
    df['DateTime']=df.iloc[:,0].astype(str)+' '+df['Time_format']
    df['DateTime']=pd.to_datetime(df['DateTime'],format='%Y-%m-%d %H:%M:%S',errors='coerce')
    
    df=df[['DateTime'] + [col for col in df.columns if col!='DateTime']]
    df=df.drop(['Time_format','string'],axis=1)
    
    date_time=df.iloc[:,:3]
    v_c=df.iloc[:,6:8]
    
    #%%
    vol_data=df.iloc[:,21:224]
    vol_data=pd.concat([date_time,vol_data,v_c],axis=1)
    #%%
    # vol_data=vol_data.drop(vol_data.iloc[:,6:8],axis=1)
    # vol_data=vol_data.drop([vol_data.columns[4],vol_data.columns[6],vol_data.columns[7]],axis=1)
    vol_data=vol_data[vol_data[vol_data.columns[10]].notna()]
    for col in vol_data.columns[3:-2]:
        vol_data[col]=vol_data[col].astype(str).str.extract(r'(\d{5})')
    #%% uncomment
    vol_data=vol_data[vol_data[vol_data.columns[-7]].notna()]
    vol_data=vol_data.dropna(how='all',axis=1)
    
    #%% one file; above cell
    # vol_data=vol_data[vol_data[vol_data.columns[-7]].notna()]
    # vol_data=vol_data[vol_data[vol_data.columns[4]].isna()]
    # vol_data=vol_data.dropna(how='all',axis=1)
    # vol_data=vol_data[vol_data[vol_data.columns[5]].isna()]
    # vol_data=vol_data.dropna(how='all',axis=1)
    # vol_data=vol_data[vol_data[vol_data.columns[-4]].isna()]
    # vol_data=vol_data.dropna(how='all',axis=1)
    # vol_data=vol_data[vol_data[vol_data.columns[4]].notna()]
    # vol_data=vol_data.dropna(how='all',axis=1)
    
    #%%
    temp_data=df.iloc[:,224:327].replace(0,np.nan).dropna(how='all',axis=1)
    
    temp_data=temp_data[temp_data[temp_data.columns[3]].notna()]
    temp_data=temp_data[temp_data[temp_data.columns[1]].isna()] # one file
    temp_data=temp_data[temp_data[temp_data.columns[-1]].isna()] # one file
    temp_data=temp_data.dropna(how='all',axis=0)    
    
    temp_data=temp_data.dropna(how='all',axis=1)
    
    temp_data=pd.concat([date_time,temp_data],axis=1)
    for col in temp_data.columns[3:]:
        temp_data[col]=temp_data[col].astype(str).str.extract(r'(\d{2})').astype(float)
        
    #%%
    # temp_data=temp_data[temp_data[temp_data.columns[6]].notna()]
    # temp_data=temp_data[temp_data[temp_data.columns[-5]].notna()]
    # temp_data=temp_data[temp_data[temp_data.columns[-5]].notna()] # one file
    temp_data=temp_data[temp_data[temp_data.columns[3]].isna()] # one file
    temp_data=temp_data[temp_data[temp_data.columns[-1]].isna()] # one file
    temp_data=temp_data.dropna(how='all',axis=0)
    # temp_data=temp_data[temp_data[temp_data.columns[-6]].notna()] # one file
    # temp_data=temp_data.replace(0,np.nan)
    temp_data=temp_data.replace(88,np.nan)
    temp_data=temp_data.dropna(how='all',axis=1)
    # temp_data=temp_data[temp_data[temp_data.columns[-1]].notna()]
    # temp_data=temp_data.dropna(how='all',axis=1)
    
    t=[f'MT_{x}' for x in range(1,t_no)]
    v=[f'MV_{x}' for x in range(1,v_no)]
    
    for i,n in enumerate(['DateTime','Date','Time']):
        v.insert(i,n)
        t.insert(i,n)
    
    v.extend(['Pack Voltage','Pack Current'])  
    temp_data.columns=t
    vol_data.columns=v
    
    
    #%%
    
    # files=glob.glob(r'D:\Benisha\Andaman\Andaman Battery Data_February\*.xlsx')
    
    
    relay_str=r'([A-Z]{2,5})' # r'([A-Z][a-z]{2,5})' # r'([A-Z][a-z]{4,})'
    
    #%%
    # for file in files:
    # cols=['Date','Time','CV_tag','status_V_T_tag','val1','junk1','val2','junk2','val_3','junk3','val_4','junk4','val_5','junk5','val_6','junk6','delmV']
    
    file_name=str('-'.join(file.rsplit('\\',1)[1].rsplit('.',1)[0].split('_')))
    # data=pd.read_excel(file,names=cols,index_col=False)
    img_path=file.rsplit('\\',1)[0]+'\\'+file.rsplit('\\',1)[1].rsplit('.',1)[0]
    
    if not os.path.exists(img_path):
        os.makedirs(img_path)
    #%%
    document=docx.Document()
    document.styles['Normal'].font.name='Calibri' # Calibri # Times New Roman
    H1=document.add_heading()
    H1.alignment=1
    H1.add_run('Andaman Data ('+ str('-'.join(file.rsplit('\\',1)[1].rsplit('.',1)[0].split('_'))) +' data)').font.size=docx.shared.Pt(18)
        
    #%%
    
    vol_data.loc[:,'MV_1':'MV_'+str(v_no-1)]=vol_data.loc[:,'MV_1':'MV_'+str(v_no-1)].fillna(method='ffill')
    vol_data.loc[:,'MV_1':'MV_'+str(v_no-1)]=vol_data.loc[:,'MV_1':'MV_'+str(v_no-1)].astype(int)
    
    temp_data.loc[:,'MT_1':'MT_'+str(t_no-1)]=temp_data.loc[:,'MT_1':'MT_'+str(t_no-1)].astype(float)
    temp_data.loc[:,'MT_1':'MT_'+str(t_no-1)]=temp_data.loc[:,'MT_1':'MT_'+str(t_no-1)].fillna(method='ffill')
    temp_data['Mean_T']=temp_data.loc[:,'MT_1':'MT_'+str(t_no-1)].mean(axis=1)
    
    temp_data['delT']=temp_data.loc[:,'MT_1':'MT_'+str(t_no-1)].max(axis=1)-temp_data.loc[:,'MT_1':'MT_'+str(t_no-1)].min(axis=1)
    
    #%% Calculating Capacity_calculated and energy
    
    vol_data['Pack Current']=vol_data['Pack Current'].astype(str).str.extract(r'(\-{0,}\d{1,})')
    vol_data['Pack Voltage']=vol_data['Pack Voltage'].astype(str).str.extract(r'(\d{1,})')
    vol_data['Pack Current']=vol_data['Pack Current'].fillna(0)
    vol_data['Pack Current']=vol_data['Pack Current'].astype(float)
    vol_data['Pack Voltage']=vol_data['Pack Voltage'].astype(float)
    
    vol_data['Pack Current']=vol_data['Pack Current'].where((vol_data['Pack Current']>-100) & (vol_data['Pack Current']<100),np.nan)
    vol_data['Pack Current']=vol_data['Pack Current'].fillna(method='ffill')
    vol_data['Pack Current']=vol_data['Pack Current'].fillna(method='bfill')
    
    #%%
    vol_data['Time_in_sec_s']=(vol_data['DateTime']-vol_data['DateTime'].shift(1))/np.timedelta64(1,'s')# time difference b/w every datapt
    vol_data['Time_in_sec']=(vol_data['DateTime']-vol_data['DateTime'].iloc[0])/np.timedelta64(1,'s')
    
    vol_data['State']=np.nan
    vol_data['Time_in_sec_s_cap']=np.where((vol_data['Time_in_sec_s']>300),np.nan,vol_data['Time_in_sec_s'])
    vol_data['State']=np.where(vol_data['Pack Current']>0,0,vol_data['State']) # 1 chg
    vol_data['State']=np.where(vol_data['Pack Current']<0,1,vol_data['State']) # 2 dchg
    vol_data['State']=np.where(vol_data['Pack Current']==0,2,vol_data['State'])
    
    #%%
    vol_data['Cap_inst']=vol_data['Time_in_sec_s_cap']*abs(vol_data['Pack Current'])/3600
    vol_data['Capacity_calculated']=vol_data['Cap_inst'].groupby(vol_data['State']).cumsum()
    vol_data['Capacity_calculated_chg']=(vol_data[vol_data['State']==0]['Cap_inst']).cumsum()
    vol_data['Capacity_calculated_dchg']=(vol_data[vol_data['State']==1]['Cap_inst']).cumsum()
    
    vol_data['Capacity_calculated_chg']=vol_data['Capacity_calculated_chg'].fillna(method='bfill')
    vol_data['Capacity_calculated_dchg']=vol_data['Capacity_calculated_dchg'].fillna(method='bfill')
    
    vol_data['Energy_calculated']=vol_data['Capacity_calculated']*vol_data['Pack Voltage']
    vol_data['Energy_calculated_chg']=vol_data['Capacity_calculated_chg']*vol_data['Pack Voltage']
    vol_data['Energy_calculated_dchg']=vol_data['Capacity_calculated_dchg']*vol_data['Pack Voltage']
    
    #%%
    vol_data['Mean_V']=vol_data.loc[:,'MV_1':'MV_'+str(v_no-1)].mean(axis=1)
    vol_data['delV']=vol_data.loc[:,'MV_1':'MV_'+str(v_no-1)].max(axis=1)-vol_data.loc[:,'MV_1':'MV_'+str(v_no-1)].min(axis=1)
    
    #%% Adding table in document
    
    document.add_paragraph().add_run('\n\nPack Specifications:').bold=True
    
    table1=document.add_table(rows=5,cols=2)
    
    table1.cell(0,0).text='Pack Energy'
    table1.cell(0,1).text='110 kWh'
    table1.cell(1,0).text='Pack Capacity'
    table1.cell(1,1).text='300Ah (2*150Ah)'
    table1.cell(2,0).text='Minimum Cell Voltage'
    table1.cell(2,1).text='3.0V'
    table1.cell(3,0).text='Nominal Cell Voltage'
    table1.cell(3,1).text='3.74V'
    table1.cell(4,0).text='Maximum Cell Voltage'
    table1.cell(4,1).text='4.2V'
    
    table1.style='Table Grid'
    
    
    document.add_paragraph().add_run('\nTest Specifications:').bold=True
    
    table2=document.add_table(rows=4,cols=2)
    table2.cell(0,0).text='Test Date and Time'
    table2.cell(0,1).text=str(vol_data['DateTime'].iloc[0])+' to '+str(vol_data['DateTime'].iloc[-1])+' ('+ str(round((vol_data['DateTime'].iloc[-1]-vol_data['DateTime'].iloc[0])/np.timedelta64(1,'h'),3)) +' hours)'
    table2.cell(1,0).text='Discharging Time Period'
    table2.cell(2,0).text='Charging Time Period'
    table2.cell(3,0).text='C-rate'
    
    table2.style='Table Grid'
    
    l=7
    b=4
    x_axis='DateTime'
    
    file_name='_'+file_name
    #%%
    document.add_paragraph().add_run('Pack Current:').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(vol_data[x_axis],vol_data['Pack Current'])
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    
    plt.xlabel(x_axis,fontweight='bold')
    plt.ylabel('Current(A)',fontweight='bold')
    
    plt.grid(linestyle='dotted') 
    plt.title('Pack Current'+file_name,fontweight='bold') # +file_name
    plt.tight_layout()
    # plt.savefig(img_path+'\\Pack Current.png',dpi=1200)
    
    document.add_picture(img_path+'\\Pack Current.png')
    
    #%%
    document.add_paragraph().add_run('Pack Capacity (calculated):').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(vol_data[x_axis],vol_data['Capacity_calculated_dchg'])
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    
    plt.xlabel(x_axis,fontweight='bold')
    plt.ylabel('Capacity(Ah)',fontweight='bold')
    
    plt.grid(linestyle='dotted')
    plt.title('Pack Discharging Capacity (Calculated)'+file_name,fontweight='bold')
    plt.tight_layout()
    # plt.savefig(img_path+'\\Pack Discharging Capacity_calculated.png',dpi=1200)
    
    document.add_picture(img_path+'\\Pack Discharging Capacity_calculated.png')
    document.add_paragraph('In this discharging session, the capacity expended by the pack is about %.3fAh ().' %vol_data['Capacity_calculated_dchg'].max())
    
    #%%
    plt.figure(figsize=(l,b))
    plt.plot(vol_data[x_axis],vol_data['Capacity_calculated_chg'])
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    
    plt.xlabel(x_axis,fontweight='bold')
    plt.ylabel('Capacity(Ah)',fontweight='bold')
    
    plt.grid(linestyle='dotted')
    plt.title('Pack Charging Capacity (Calculated)'+file_name,fontweight='bold')
    plt.tight_layout()
    
    # plt.savefig(img_path+'\\Pack Charging Capacity_calculated.png',dpi=1200)
    
    document.add_picture(img_path+'\\Pack Charging Capacity_calculated.png')
    document.add_paragraph('In this charging session, the capacity transferred to the pack is about %.3fAh (). ' %vol_data['Capacity_calculated_chg'].max())
    document.add_paragraph().add_run('Note: These values were calculated from the current (in amperes) values and time (in hours).').underline=True
    
    #%%
    document.add_paragraph().add_run('Pack Energy (calculated):').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(vol_data[x_axis],vol_data['Energy_calculated_dchg']*0.001)
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    
    plt.xlabel(x_axis,fontweight='bold')
    plt.ylabel('Energy(kWh)',fontweight='bold')
    
    plt.grid(linestyle='dotted')
    plt.title('Pack Discharging Energy (Calculated)'+file_name,fontweight='bold')
    plt.tight_layout()
    # plt.savefig(img_path+'\\Pack Discharging Energy_calculated.png',dpi=1200)
    
    document.add_picture(img_path+'\\Pack Discharging Energy_calculated.png')
    document.add_paragraph('The Pack Energy expended in this discharging session (approx.  minutes) is about %.3fkWh.' % (vol_data['Energy_calculated_dchg'].max()*0.001))
    
    #%%
    plt.figure(figsize=(l,b))
    plt.plot(vol_data[x_axis],vol_data['Energy_calculated_chg']*0.001)
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    
    plt.xlabel(x_axis,fontweight='bold')
    plt.ylabel('Energy(kWh)',fontweight='bold')
    
    plt.grid(linestyle='dotted')
    plt.title('Pack Charging Energy (Calculated)'+file_name,fontweight='bold')
    plt.tight_layout()
    # plt.savefig(img_path+'\\Pack Charging Energy_calculated.png',dpi=1200)
    
    # document.add_picture(img_path+'\\Pack Charging Energy_calculated.png')
    document.add_paragraph('The Pack Energy expended in this charging session () is about %.3fkWh.' % (vol_data['Energy_calculated_chg'].max()*0.001))
    document.add_paragraph().add_run('Note: These values were calculated from the calculated capacity (in Ah) values and pack voltage (in volts).').underline=True
    
    
    #%%
    # document.add_paragraph().add_run('Pack Energy (BMS):').bold=True
    
    # plt.figure(figsize=(l,b))
    # plt.plot(vol_data[x_axis],vol_data['Discharging Energy'])
    
    # # xmin, xmax = ax.get_xlim()
    # # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    # # ymin, ymax = ax.get_ylim()
    # # ax.set_yticks(np.round(np.linspace(ymin, ymax, 20), 2))
    
    # plt.xlabel(x_axis,fontweight='bold')
    # plt.ylabel('Energy(Wh)',fontweight='bold')
    
    # plt.grid(linestyle='dotted')
    # plt.title('Pack Discharging Energy (BMS)',fontweight='bold')
    # plt.tight_layout()
    # plt.savefig(img_path+'\\Pack Discharging Energy_BMS.png',dpi=1200)
    
    # document.add_picture(img_path+'\\Pack Discharging Energy_BMS.png')
    # p_bms_d_engy=document.add_paragraph('The ')
    # p_bms_d_engy.add_run('discharging pack energy ').bold=True
    # p_bms_d_engy.add_run('is about ')
    # p_bms_d_engy.add_run('%.3fWh.' %vol_data['Discharging Energy'].max()).bold=True
    
    #%%
    # plt.figure(figsize=(l,b))
    # plt.plot(vol_data[x_axis],vol_data['Charging Energy'])
    
    # # xmin, xmax = ax.get_xlim()
    # # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    
    # # ymin, ymax = ax.get_ylim()
    # # ax.set_yticks(np.round(np.linspace(ymin, ymax, 20), 2))
    
    # plt.xlabel(x_axis,fontweight='bold')
    # plt.ylabel('Energy(Wh)',fontweight='bold')
    
    # plt.grid(linestyle='dotted')
    # plt.title('Pack Charging Energy (BMS)',fontweight='bold')
    # plt.tight_layout()
    # plt.savefig(img_path+'\\Pack Charging Energy_BMS.png',dpi=1200)
    
    # document.add_picture(img_path+'\\Pack Charging Energy_BMS.png')
    # p_bms_c_engy=document.add_paragraph('The ')
    # p_bms_c_engy.add_run('charging pack energy ').bold=True
    # p_bms_c_engy.add_run('is about ')
    # p_bms_c_engy.add_run('%.3fWh.' %vol_data['Charging Energy'].max()).bold=True
    
    #%%
    document.add_paragraph().add_run('Pack Voltage:').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(vol_data[x_axis],vol_data['Pack Voltage'])
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    
    plt.xlabel(x_axis,fontweight='bold')
    plt.ylabel('Voltage(V)',fontweight='bold')
    
    plt.grid(linestyle='dotted')
    plt.title('Pack Voltage'+file_name,fontweight='bold')
    plt.tight_layout()
    # plt.savefig(img_path+'\\Pack Voltage.png',dpi=1200)
    
    document.add_picture(img_path+'\\Pack Voltage.png')
    pv=document.add_paragraph('During ')
    pv.add_run('discharge, ').bold=True
    pv.add_run('the pack goes from to . The pack voltage at rest after discharge is . At ')
    pv.add_run('charge, ').bold=True
    pv.add_run('the pack goes from to . ')
    pv.add_run('In the end, the voltage stabilizes at %sV during the end rest.' %vol_data['Pack Voltage'].iloc[-1])
    
    #%%
    document.add_paragraph().add_run('Cell Voltage:').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(vol_data[x_axis],vol_data.loc[:,'MV_1':'MV_'+str(v_no-1)]*0.0001,marker='o',markersize=3) #
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    plt.ylabel('Voltage(V)',fontweight='bold')
    plt.xlabel(x_axis,fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Cell Voltage'+file_name,fontweight='bold')
    plt.tight_layout()
    
    #%%
    document.add_paragraph().add_run('Average Voltage:').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(vol_data[x_axis],(vol_data['Mean_V'])*0.0001)
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    plt.ylabel('Voltage(V)',fontweight='bold')
    plt.xlabel(x_axis,fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Average Voltage'+file_name,fontweight='bold')
    plt.tight_layout()
    # plt.savefig(img_path+'\\Average Voltage.png',dpi=1200)
    
    document.add_picture(img_path+'\\Average Voltage.png')
    av=document.add_paragraph('The average cell voltage during ')
    av.add_run('discharge ').bold=True
    av.add_run('goes from ')
    av.add_run('V to V. ').bold=True
    av.add_run('The average cell voltage at the rest after discharge is V. At ')
    av.add_run('charge').bold=True
    av.add_run(', the average cell voltage goes from ')
    av.add_run('V to V.').bold=True
    av.add_run('In the end rest, it  stabilizes at %.3fV.' % (vol_data['Mean_V'].iloc[-1]*0.0001))
    
    #%%
    document.add_paragraph().add_run('Voltage Difference (delV):').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(vol_data[x_axis],(vol_data['delV'])*0.1)
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    plt.ylabel('Voltage(mV)',fontweight='bold')
    plt.xlabel(x_axis,fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Voltage difference (delV)'+file_name,fontweight='bold')
    plt.tight_layout()
    # plt.savefig(img_path+'\\delV.png',dpi=1200)
    
    document.add_picture(img_path+'\\delV.png')
    dv=document.add_paragraph('During discharge, the ')
    dv.add_run('maximum delV (difference between the highest and lowest cell voltage), ').bold=True
    dv.add_run('is about ')
    dv.add_run('mV, at the end of discharge. ').bold=True
    dv.add_run('At the rest after discharging, the delmV is mV. In the charging session, the maximum delV is mV. ')
    dv.add_run('The pack stabilizes at %.3fmV delV at the end rest. '% (vol_data['delV'].iloc[-1]*0.1))
    
    
    #%%
    document.add_paragraph().add_run('Cell Temperature:').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(temp_data[x_axis],temp_data.loc[:,'MT_1':'MT_'+str(t_no-1)])
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.xlabel(x_axis,fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Cell Temperature'+file_name,fontweight='bold')
    ct=document.add_paragraph('The maximum temperature reached is around ')
    ct.add_run('%.3f'%(temp_data.loc[:,'MT_1':'MT_'+str(t_no-1)].max(axis=1).max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True
    
    #%%
    document.add_paragraph().add_run('Average temperature:').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(temp_data[x_axis],temp_data['Mean_T'])
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.xlabel(x_axis,fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Average Temperature'+file_name,fontweight='bold')
    plt.tight_layout()
    # plt.savefig(img_path+'\\Average Temperature.png',dpi=1200)
    
    document.add_picture(img_path+'\\Average Temperature.png')
    at=document.add_paragraph('As can be seen, the ')
    at.add_run('average cell temperature ').bold=True
    at.add_run('stays between ')
    at.add_run(u'\N{DEGREE SIGN}'+'C to '+u'\N{DEGREE SIGN}'+'C ').bold=True
    at.add_run('during ')
    at.add_run('discharging').bold=True
    at.add_run(' and ')
    at.add_run(u'\N{DEGREE SIGN}'+'C to '+u'\N{DEGREE SIGN}'+'C ').bold=True
    at.add_run('during ')
    at.add_run('charging.').bold=True
    
    #%%
    document.add_paragraph().add_run('Temperature Difference (delT):').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(temp_data[x_axis],temp_data['delT'])
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.xlabel(x_axis,fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Temperature difference (delT)'+file_name,fontweight='bold')
    plt.tight_layout()
    # plt.savefig(img_path+'\\delT.png',dpi=1200)
    
    document.add_picture(img_path+'\\delT.png')
    dt=document.add_paragraph('The ')
    dt.add_run('maximum delT ').bold=True
    dt.add_run('is about ')
    dt.add_run('%.3f' %temp_data['delT'].max()).bold=True
    dt.add_run(u'\N{DEGREE SIGN}'+'C.').bold=True
    
    # #%%
    # plt.plot()
    # plt.figure(figsize=(l,b))
    # plt.plot(vol_data[x_axis],vol_data.loc[:,'MV_190':'MV_192'])
    
    # # xmin, xmax = ax.get_xlim()
    # # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    # plt.ylabel('Voltage(mV)',fontweight='bold')
    # plt.xlabel(x_axis,fontweight='bold')
    # plt.grid(linestyle='dotted')
    # plt.title('Cell Voltage',fontweight='bold')
    
    
    #%%
    # #%%
    # plt.plot()
    # plt.figure(figsize=(l,b))
    # plt.plot(temp_data[x_axis],temp_data.loc[:,'MT_1':'MT_48'])
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    # plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    # plt.xlabel('Time',fontweight='bold')
    # plt.legend(t_names)
    # plt.grid(linestyle='dotted')
    # plt.title('Cell Temperature',fontweight='bold')
    
    
    
    
    
    #%%
    # plt.plot()
    # plt.figure(figsize=(l,b))
    # plt.plot(vol_data[x_axis],vol_data['Capacity_calculated'])
    
    # # xmin, xmax = ax.get_xlim()
    # # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    
    # plt.xlabel('Time',fontweight='bold')
    # plt.ylabel('Capacity(Ah)',fontweight='bold')
    
    # plt.grid(linestyle='dotted')
    # plt.title('Pack Capacity (Calculated)',fontweight='bold')
    
    
    
    #%%
    # plt.plot()
    # plt.figure(figsize=(l,b))
    # plt.plot(vol_data[x_axis],vol_data['Energy_calculated']*0.001)
    
    # # xmin, xmax = ax.get_xlim()
    # # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    
    # plt.xlabel('Time',fontweight='bold')
    # plt.ylabel('Energy(kWh)',fontweight='bold')
    
    # plt.grid(linestyle='dotted')
    # plt.title('Pack Energy (Calculated)',fontweight='bold')

#%%

    
#%%
    vol_col_check=vol_data.columns[3:105]
    temp_col_check=temp_data.columns[3:-2]
    print('======================================================================',file=f)
    print('---------------------%s---------------------\n' %file_name,file=f)
    print('Datapoints:%s' % (len(vol_data)),file=f)
    print('Timeframe: %s to %s' % (vol_data['DateTime'].iloc[0],vol_data['DateTime'].iloc[-1]),file=f)
    print('Total Time Spent:%s' % (datetime.timedelta(seconds=(vol_data['DateTime'].iloc[-1]-vol_data['DateTime'].iloc[0])/np.timedelta64(1,'s'))),file=f)
    print('---------------------------',file=f)

    if len(vol_data[vol_data['State']==0])!=0:
        print('Chg Time Spent:%s' % (datetime.timedelta(seconds=(vol_data[vol_data['State']==0]['Time_in_sec_s'].sum()))),file=f)
        print('Maximum Charging Current:%.3fA ; Minimum Charging Current:%.3fA' % (vol_data[vol_data['State']==0]['Pack Current'].max(),vol_data[vol_data['State']==0]['Pack Current'].min()),file=f)
        print('Chg Pack Voltage range:%sV to %sV' %(vol_data[vol_data['State']==0]['Pack Voltage'].iloc[0],vol_data[vol_data['State']==0]['Pack Voltage'].iloc[-1]),file=f)
        print('---------------------------',file=f)
    
        print('Chg Min Vol:%.3fV' %((vol_data[vol_data['State']==0].loc[:,'MV_1':'MV_'+str(v_no-1)]*0.0001).min(axis=1).min()),file=f)
        print('Chg Max Vol:%.3fV' %((vol_data[vol_data['State']==0].loc[:,'MV_1':'MV_'+str(v_no-1)]*0.0001).max(axis=1).max()),file=f)
        print('Chg Average Vol:%.3fV' %((vol_data[vol_data['State']==0]['Mean_V'])*0.0001).mean(),file=f)
        
        print('Chg Minimum delV:%.3fmV' %((vol_data[vol_data['State']==0]['delV'])*0.1).min(),file=f)
        print('Chg Maximum delV:%.3fmV' %((vol_data[vol_data['State']==0]['delV'])*0.1).max(),file=f)
        print('Chg Average delV:%.3fmV' %((vol_data[vol_data['State']==0]['delV'])*0.1).mean(),file=f)
        print('---------------------------',file=f)
 
        c_dt=pd.to_datetime(vol_data[vol_data['State']==0]['DateTime'].tolist())
        chg_temp=temp_data[temp_data['DateTime'].isin(c_dt)]
        print('Chg Min Temp:%.3f degC' %(chg_temp.loc[:,'MT_1':'MT_'+str(t_no-1)].min(axis=1).min()),file=f)
        print('Chg Max Temp:%.3f degC' %(chg_temp.loc[:,'MT_1':'MT_'+str(t_no-1)].max(axis=1).max()),file=f)
        print('Chg Average Temp:%.3f degC' %(chg_temp['Mean_T'].mean()),file=f)
        
        print('Chg Minimum delT:%.3f degC' %(chg_temp['delT'].min()),file=f)
        print('Chg Maximum delT:%.3f degC' %(chg_temp['delT'].max()),file=f)
        print('Chg Average delT:%.3f degC' %(chg_temp['delT'].mean()),file=f)
        print('---------------------------',file=f)
        print('Cells with (lower)erroneous chg voltage values: %s' %([col for col in vol_col_check if vol_data[vol_data['State']==0][col].lt(23000).any()]),file=f)
        print('Cells with (higher)erroneous chg voltage values: %s' %([col for col in vol_col_check if vol_data[vol_data['State']==0][col].gt(43000).any()]),file=f)
        
        print('Cells with (lower)erroneous chg temp values: %s' %([col for col in temp_col_check if chg_temp[col].lt(17).any()]),file=f)
        print('Cells with (higher)erroneous chg temp values: %s' %([col for col in temp_col_check if chg_temp[col].gt(35).any()]),file=f)
    print('---------------------------',file=f)
    if len(vol_data[vol_data['State']==1])!=0:

        print('Dchg Time Spent:%s' % (datetime.timedelta(seconds=(vol_data[vol_data['State']==1]['Time_in_sec_s'].sum()))),file=f)
        print('Maximum Discharging Current:%.3fA ; Minimum Discharging Current:%.3fA' % (vol_data[vol_data['State']==1]['Pack Current'].max(),vol_data[vol_data['State']==1]['Pack Current'].min()),file=f)
        print('DChg Pack Voltage range:%sV to %sV' %(vol_data[vol_data['State']==1]['Pack Voltage'].iloc[0],vol_data[vol_data['State']==1]['Pack Voltage'].iloc[-1]),file=f)
        print('---------------------------',file=f)
    
        print('Dchg Min Vol:%.3fV' %((vol_data[vol_data['State']==1].loc[:,'MV_1':'MV_'+str(v_no-1)]*0.0001).min(axis=1).min()),file=f)
        print('Dchg Max Vol:%.3fV' %((vol_data[vol_data['State']==1].loc[:,'MV_1':'MV_'+str(v_no-1)]*0.0001).max(axis=1).max()),file=f)
        print('Dchg Average Vol:%.3fV' %((vol_data[vol_data['State']==1]['Mean_V'])*0.0001).mean(),file=f)
        
        print('Dchg Minimum delV:%.3fmV' %((vol_data[vol_data['State']==1]['delV'])*0.1).min(),file=f) 
        print('Dchg Maximum delV:%.3fmV' %((vol_data[vol_data['State']==1]['delV'])*0.1).max(),file=f) 
        print('Dchg Average delV:%.3fmV' %((vol_data[vol_data['State']==1]['delV'])*0.1).mean(),file=f) 
        print('---------------------------',file=f)
        
        d_dt=pd.to_datetime(vol_data[vol_data['State']==1]['DateTime'].tolist())
        dchg_temp=temp_data[temp_data['DateTime'].isin(d_dt)]  
        print('Dchg Min Temp:%.3f degC' %(dchg_temp.loc[:,'MT_1':'MT_'+str(t_no-1)].min(axis=1).min()),file=f)
        print('Dchg Max Temp:%.3f degC' %(dchg_temp.loc[:,'MT_1':'MT_'+str(t_no-1)].max(axis=1).max()),file=f)
        print('Dchg Average Temp:%.3f degC' %(dchg_temp['Mean_T'].mean()),file=f)
        print('Dchg Minimum delT:%.3f degC' %(dchg_temp['delT'].min()),file=f)
        print('Dchg Maximum delT:%.3f degC' %(dchg_temp['delT'].max()),file=f)
        print('Dchg Average delT:%.3f degC' %(dchg_temp['delT'].mean()),file=f)
        print('---------------------------',file=f)
        print('Cells with (lower)erroneous dchg voltage values: %s' %([col for col in vol_col_check if vol_data[vol_data['State']==1][col].lt(23000).any()]),file=f)
        print('Cells with (higher)erroneous dchg voltage values: %s' %([col for col in vol_col_check if vol_data[vol_data['State']==1][col].gt(43000).any()]),file=f)
        
        print('Cells with (lower)erroneous dchg temp values: %s' %([col for col in temp_col_check if dchg_temp[col].lt(17).any()]),file=f)
        print('Cells with (higher)erroneous dchg temp values: %s' %([col for col in temp_col_check if dchg_temp[col].gt(35).any()]),file=f)
    print('---------------------------',file=f)
    if len(vol_data[vol_data['State']==2])!=0:

        print('Rest Time Spent:%s' % (datetime.timedelta(seconds=(vol_data[vol_data['State']==2]['Time_in_sec_s'].sum()))),file=f)
        print('Rest Pack Voltage range:%sV to %sV' %(vol_data[vol_data['State']==2]['Pack Voltage'].iloc[0],vol_data[vol_data['State']==2]['Pack Voltage'].iloc[-1]),file=f)
        print('---------------------------',file=f)
        print('Rest Min Vol:%.3fV' %((vol_data[vol_data['State']==2].loc[:,'MV_1':'MV_'+str(v_no-1)]*0.0001).min(axis=1).min()),file=f)
        print('Rest Max Vol:%.3fV' %((vol_data[vol_data['State']==2].loc[:,'MV_1':'MV_'+str(v_no-1)]*0.0001).max(axis=1).max()),file=f)
        print('Rest Average Vol:%.3fV' %((vol_data[vol_data['State']==2]['Mean_V'])*0.0001).mean(),file=f)
        
        print('Rest Minimum delV:%.3fmV' %((vol_data[vol_data['State']==2]['delV'])*0.1).min(),file=f)
        print('Rest Maximum delV:%.3fmV' %((vol_data[vol_data['State']==2]['delV'])*0.1).max(),file=f) 
        print('Rest Average delV:%.3fmV' %((vol_data[vol_data['State']==2]['delV'])*0.1).mean(),file=f) 
        print('---------------------------',file=f)
        
        r_dt=pd.to_datetime(vol_data[vol_data['State']==2]['DateTime'].tolist())
        rst_temp=temp_data[temp_data['DateTime'].isin(r_dt)]
        print('Rest Min Temp:%.3f degC' %(rst_temp.loc[:,'MT_1':'MT_'+str(t_no-1)].min(axis=1).min()),file=f)
        print('Rest Max Temp:%.3f degC' %(rst_temp.loc[:,'MT_1':'MT_'+str(t_no-1)].max(axis=1).max()),file=f)
        print('Rest Average Temp:%.3f degC' %(rst_temp['Mean_T'].mean()),file=f)
        
        print('Rest Minimum delT:%.3f degC' %(rst_temp['delT'].min()),file=f)
        print('Rest Maximum delT:%.3f degC' %(rst_temp['delT'].max()),file=f)
        print('Rest Average delT:%.3f degC' %(rst_temp['delT'].mean()),file=f)
        print('---------------------------',file=f)
        print('Cells with (lower)erroneous Rest voltage values: %s' %([col for col in vol_col_check if vol_data[vol_data['State']==2][col].lt(23000).any()]),file=f)
        print('Cells with (higher)erroneous Rest voltage values: %s' %([col for col in vol_col_check if vol_data[vol_data['State']==2][col].gt(43000).any()]),file=f)
        
        print('Cells with (lower)erroneous Rest temp values: %s' %([col for col in temp_col_check if rst_temp[col].lt(17).any()]),file=f)
        print('Cells with (higher)erroneous Rest temp values: %s' %([col for col in temp_col_check if rst_temp[col].gt(35).any()]),file=f)
    print('---------------------------',file=f)
    print('Maximum Charging Capacity:%.3fAh ; Minimum Charging Capacity:%.3fAh' % (vol_data['Capacity_calculated_chg'].max(),vol_data['Capacity_calculated_chg'].min()),file=f)
    print('Maximum Discharging Capacity:%.3fAh ; Minimum Discharging Capacity:%.3fAh' % (vol_data['Capacity_calculated_dchg'].max(),vol_data['Capacity_calculated_dchg'].min()),file=f)
    print('Maximum Charging Energy:%.3fkWh ; Minimum Charging Energy:%.3fkWh' % ((vol_data['Energy_calculated_chg']*0.001).max(),(vol_data['Energy_calculated_chg']*0.001).min()),file=f)
    print('Maximum Discharging Energy:%.3fkWh ; Minimum Discharging Energy:%.3fkWh' % ((vol_data['Energy_calculated_dchg']*0.001).max(),(vol_data['Energy_calculated_dchg']*0.001).min()),file=f)
    
    print('Cells with (lower)erroneous voltage values: %s' %([col for col in vol_col_check if vol_data[col].lt(23000).any()]),file=f)
    print('Cells with (higher)erroneous voltage values: %s' %([col for col in vol_col_check if vol_data[col].gt(43000).any()]),file=f)
    
    print('Cells with (lower)erroneous temp values: %s' %([col for col in temp_col_check if temp_data[col].lt(17).any()]),file=f)
    print('Cells with (higher)erroneous temp values: %s' %([col for col in temp_col_check if temp_data[col].gt(35).any()]),file=f)
    
    print('======================================================================\n',file=f)
    print('\n',file=f)
    
    
    #%%
    # path=img_path+'\\'+file.rsplit('\\',1)[1].rsplit('.',1)[0]+'_modified.xlsx'
    # document.save(img_path+'\\'+file.rsplit('\\',1)[1].rsplit('.',1)[0]+'_observations.docx')
    
    # #%%
    # writer = pd.ExcelWriter(path, engine = 'xlsxwriter')
    # vol_data.to_excel(writer, sheet_name = 'Cell Voltage',index=False)
    # temp_data.to_excel(writer, sheet_name = 'Cell Temperature',index=False)
    # writer.close()
    
    # #%%
    # col_check=vol_data.columns[3:-2]
    # f=open(img_path+'\\'+file.rsplit('\\',1)[1].rsplit('.',1)[0]+'_observations.txt',"w")
    
    # print('======================================================================',file=f)
    # print('---------------------Current Data---------------------',file=f)
    # print('Datapoints:%s' % (len(vol_data)),file=f)
    # print('Total Time Spent:%s minutes' % ((vol_data['DateTime'].iloc[-1]-vol_data['DateTime'].iloc[0])/np.timedelta64(1,'m')),file=f)
    # print('Maximum Charging Current:%s ; Minimum Charging Current:%s' % (vol_data[vol_data['State']==0]['Pack Current'].max(),vol_data[vol_data['State']==0]['Pack Current'].min()),file=f)
    # print('Maximum Discharging Current:%s ; Minimum Discharging Current:%s' % (vol_data[vol_data['State']==1]['Pack Current'].max(),vol_data[vol_data['State']==1]['Pack Current'].min()),file=f)
    # print('Maximum Charging Capacity:%s ; Minimum Charging Capacity:%s' % (vol_data['Capacity_calculated_chg'].max(),vol_data['Capacity_calculated_chg'].min()),file=f)
    # print('Maximum Discharging Capacity:%s ; Minimum Discharging Capacity:%s' % (vol_data['Capacity_calculated_dchg'].max(),vol_data['Capacity_calculated_dchg'].min()),file=f)
    # print('Maximum Charging Energy:%s ; Minimum Charging Energy:%s' % (vol_data['Energy_calculated_chg'].max(),vol_data['Energy_calculated_chg'].min()),file=f)
    # print('Maximum Discharging Energy:%s ; Minimum Discharging Energy:%s' % (vol_data['Energy_calculated_dchg'].max(),vol_data['Energy_calculated_dchg'].min()),file=f)
    # print('Maximum delT:%s degC' %(temp_data['delT'].max()),file=f)
    # print('Cells with (lower)erroneous values: %s' %([col for col in col_check if vol_data[col].lt(23000).any()]),file=f)
    # print('Cells with (higher)erroneous values: %s' %([col for col in col_check if vol_data[col].gt(43000).any()]),file=f)
    # print('======================================================================',file=f)
    
    # f.close()

f.close()


#%%% Comment this section

# import pandas as pd
# import datetime
# import numpy as np

# f=r"D:\Benisha\Andaman\In-house_testing_data\Bank1\09_march_Dis Chag_ banks B1_parallel test\09_march_Dis Chag_ banks B1_parallel test_modified.xlsx"
# # f=r"D:\Benisha\Andaman\In-house_testing_data\Bank1\19_April_Bank1_series with deltal Discharge\19_April_Bank1_series with deltal Discharge_modified.xlsx"

# # f=r"D:\Benisha\Andaman\In-house_testing_data\Bank2\09_march_Dis Chag_ banks B2_parallel test\09_march_Dis Chag_ banks B2_parallel test_modified.xlsx"
# # f=r"D:\Benisha\Andaman\In-house_testing_data\Bank2\19_April_Bank2_series with deltal Discharge\19_April_Bank2_series with deltal Discharge_modified.xlsx"


# df=pd.read_excel(f,sheet_name='Cell Voltage')
# df_t=pd.read_excel(f,sheet_name='Cell Temperature')

# ##%%
# # f=r"D:\Benisha\Andaman\Andaman Battery Data_February\B1_014_02_24_log\B1_014_02_24_log_modified.xlsx"
# # f=r"D:\Benisha\Andaman\Andaman Battery Data_February\B2_014_02_24_log\B2_014_02_24_log_modified.xlsx"

# # df=pd.read_excel(f,sheet_name='Pack Details')
# # df_t=pd.read_excel(f,sheet_name='Cell Temperature')
# # df_v=pd.read_excel(f,sheet_name='Cell Voltage')
# # s=pd.read_excel(f,sheet_name='Pack Status')
# # m=pd.read_excel(f,sheet_name='Max Min Mean')


# df['Capacity_calculated']=df['Capacity_calculated'].fillna(method='ffill')
# df['Energy_calculated']=df['Energy_calculated'].fillna(method='ffill')

# df['ElapsedTime']=(df['DateTime']-df['DateTime'].iloc[0])
# df['ElapsedTime']=df['ElapsedTime'].astype(str)
# df[['days','junk','ElapsedTime']]=df['ElapsedTime'].str.split(' ',expand=True)
# df[['hours','mins','sec']]=df['ElapsedTime'].str.split(':',expand=True)
# df['hours']=df['hours'].astype(int)+df['days'].astype(int)*24
# df['ElapsedTime']=df['hours'].astype(str)+':'+df['mins']+':'+df['sec']
# df=df.drop(labels=['days','junk','hours','mins','sec'],axis=1)
# df=df.fillna(method='ffill')
# df=df.fillna(method='bfill')

# df_t['ElapsedTime']=(df_t['DateTime']-df_t['DateTime'].iloc[0])
# df_t['ElapsedTime']=df_t['ElapsedTime'].astype(str)
# df_t[['days','junk','ElapsedTime']]=df_t['ElapsedTime'].str.split(' ',expand=True)
# df_t[['hours','mins','sec']]=df_t['ElapsedTime'].str.split(':',expand=True)
# df_t['hours']=df_t['hours'].astype(int)+df_t['days'].astype(int)*24
# df_t['ElapsedTime']=df_t['hours'].astype(str)+':'+df_t['mins']+':'+df_t['sec']
# df_t=df_t.drop(labels=['days','junk','hours','mins','sec'],axis=1)
# df_t=df_t.fillna(method='ffill')
# df_t=df_t.fillna(method='bfill')


# # #%%
# # df_v['ElapsedTime']=(df_v['DateTime']-df_v['DateTime'].iloc[0])
# # df_v['ElapsedTime']=df_v['ElapsedTime'].astype(str)
# # df_v[['days','junk','ElapsedTime']]=df_v['ElapsedTime'].str.split(' ',expand=True)
# # df_v[['hours','mins','sec']]=df_v['ElapsedTime'].str.split(':',expand=True)
# # df_v['hours']=df_v['hours'].astype(int)+df_v['days'].astype(int)*24
# # df_v['ElapsedTime']=df_v['hours'].astype(str)+':'+df_v['mins']+':'+df_v['sec']
# # df_v=df_v.drop(labels=['days','junk','hours','mins','sec'],axis=1)
# # df_v=df_v.fillna(method='ffill')
# # df_v=df_v.fillna(method='bfill')

# # #%%
# # writer = pd.ExcelWriter(f, engine = 'xlsxwriter')
# # df_v.to_excel(writer, sheet_name = 'Cell Voltage',index=False)
# # df_t.to_excel(writer, sheet_name = 'Cell Temperature',index=False)
# # df.to_excel(writer,sheet_name='Pack Details',index=False)
# # s.to_excel(writer,sheet_name='Pack Status',index=False)
# # m.to_excel(writer,sheet_name='Max Min Mean',index=False)
# # writer.close()

# ##%%
# writer = pd.ExcelWriter(f, engine = 'xlsxwriter')
# df.to_excel(writer, sheet_name = 'Cell Voltage',index=False)
# df_t.to_excel(writer, sheet_name = 'Cell Temperature',index=False)
# writer.close()

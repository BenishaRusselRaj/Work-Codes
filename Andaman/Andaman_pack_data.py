# -*- coding: utf-8 -*-
"""
Created on Tue Apr  2 11:31:54 2024

@author: IITM
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import docx
import os
import glob
import datetime


#%%

# files=glob.glob(r"D:\Benisha\Andaman\Field_data\24_05_17\rs485 data bank 2 01-05-2024\*.xlsx")
# files=glob.glob(r"D:\Benisha\Andaman\Field_data\24_05_17\rs485 data bank 2 10-05-2024\*.xlsx")

files=glob.glob(r"D:\Benisha\Andaman\Field_data\24_05_17\rs485 data bank 2 10-05-2024\bank 2 10-05-2024.xlsx")

v_no=103 # no+1
t_no=69 # no+1
mod_no=18 # no+1

relay_str=r'([A-Z]{2,5})' # r'([A-Z][a-z]{2,5})' # r'([A-Z][a-z]{4,})'
# 
#%%
for file in files:
    cols=['Date','Time','CV_tag','status_V_T_tag','val1','junk1','val2','junk2','val_3','junk3','val_4','junk4','val_5','junk5','val_6','junk6','delmV']
    
    file_name=str('-'.join(file.rsplit('\\',1)[1].rsplit('.',1)[0].split('_')))
    data=pd.read_excel(file,names=cols,index_col=False)
    img_path=file.rsplit('\\',1)[0]+'\\'+file.rsplit('\\',1)[1].rsplit('.',1)[0]
    
    if not os.path.exists(img_path):
        os.makedirs(img_path)
    #%%
    document=docx.Document()
    document.styles['Normal'].font.name='Calibri' # Calibri # Times New Roman
    H1=document.add_heading()
    H1.alignment=1
    H1.add_run('Andaman Data ('+ str('-'.join(file.rsplit('\\',1)[1].rsplit('.',1)[0].split('_'))) +' data)').font.size=docx.shared.Pt(18)
    
    #%%
    data['Date']=data['Date'].str.replace('[','',regex=True)
    data['Time']=data['Time'].str.replace(']','',regex=True)
    
    data['Time']=data['Time'].astype(str)
    data[['Time_format','string']]=data['Time'].str.split('.',expand=True)
    
    data['DateTime']=data['Date'].astype(str)+' '+data['Time_format']
    data['DateTime']=pd.to_datetime(data['DateTime'],format='%Y-%m-%d %H:%M:%S',errors='coerce')
    
    data=data[['DateTime'] + [col for col in data.columns if col!='DateTime']]
    data=data.drop(['Time_format','string'],axis=1)
    
    #%%
    # data=data[data['DateTime']<='2023-10-10 17:30']
    
    #%%
    status_data=data[data['status_V_T_tag']=='STATUS:'].dropna(how='all',axis=1)
    max_min_data=data[data['status_V_T_tag']=='V:'].reset_index(drop=True)
    max_min_data=pd.concat([max_min_data,data[data['status_V_T_tag']=='T:'].drop(['DateTime','Date','Time'],axis=1).reset_index(drop=True)],axis=1).dropna(how='all',axis=1) # try to concatenate in same file
    details=data[data['status_V_T_tag']=='voltage:'].dropna(how='all',axis=1)
    
    #%% columns names
    v_names=[f'MV_{x}' for x in range(1,v_no)] 
    t_names=[f'MT_{x}' for x in range(1,t_no)]
    
    
    # detail_names=['DateTime','Date','Time','junk','junk','Pack Voltage','junk','junk','Pack Current','junk','SoH','junk','junk','Charging Energy','junk','junk','Discharging Energy','junk','junk','junk','Pack Available Energy']
    detail_names=['DateTime','Date','Time','junk','junk','Pack Voltage','junk','junk','Pack Current','junk','SoH','junk','Cycles_main']
    max_min_names=['DateTime','Date','Time','junk','junk','Min_V_BMS','','junk','junk','Max_V_BMS','','junk','junk','Average_CellV_BMS','junk','junk','junk','delV_BMS(mV)','junk','junk','Min_T_BMS','','junk','junk','Max_T_BMS','']
    
    try:
        status=['DateTime','Date','Time','junk','junk','Pack Status','junk','Shutdown','Normal','junk','junk','Precharge Relay','junk','junk','Main Relay','junk'] #,'junk','junk'
        status_data.columns=status
    except:
        status=['DateTime','Date','Time','junk','junk','Pack Status','junk','Shutdown','Normal','junk','junk','Precharge Relay','junk','junk','Main Relay'] #,'junk','junk'
        status_data.columns=status
    details.columns=detail_names
    max_min_data.columns=max_min_names
    #%%
    # status_data['Main Relay']=status_data['Main Relay'].astype(str).str.extract(r'([A-Z][a-z]{4,})') #[Open][close][open] 
    #%%
    status_data=status_data.drop([col for col in status_data.columns if "junk" in col], axis=1)
    details=details.drop([col for col in details.columns if "junk" in col], axis=1)
    max_min_data=max_min_data.drop([col for col in max_min_data.columns if "junk" in col], axis=1)
    
    #%%
    
    # details['Charging Energy']=details['Charging Energy'].astype(str).str.extract(r'(\d{1}.\d{1})')
    # details['Discharging Energy']=details['Discharging Energy'].astype(str).str.extract(r'(\d{1}.\d{1})')
    
    #%%
    details['Pack Current']=details['Pack Current'].str.replace('A','',regex=True)
    status_data['Precharge Relay']=status_data['Precharge Relay'].str.replace(',','',regex=True)
    # details['Charging Energy']=details['Charging Energy'].str.replace('engy:','',regex=True)
    # details['Discharging Energy']=details['Discharging Energy'].str.replace('engy:','',regex=True)
    # details['Charging Energy']=details['Charging Energy'].astype(float)
    # details['Discharging Energy']=details['Discharging Energy'].astype(float)
    
    #%%
    status_data['Main Relay']=status_data['Main Relay'].astype(str).str.extract(relay_str)
    status_data['Precharge Relay']=status_data['Precharge Relay'].astype(str).str.extract(relay_str)
    
    #%%
    # status_data['Precharge Relay']=status_data['Precharge Relay'].str.replace('R',np.nan,regex=True)
    
    #%%
    status_data['Main Relay']=status_data['Main Relay'].fillna(method='ffill')
    status_data['Main Relay']=status_data['Main Relay'].fillna(method='bfill')
    status_data['Main Relay']=status_data['Main Relay'].astype(str)
    
    status_data['Precharge Relay']=status_data['Precharge Relay'].fillna(method='ffill')
    status_data['Precharge Relay']=status_data['Precharge Relay'].fillna(method='bfill')
    status_data['Precharge Relay']=status_data['Precharge Relay'].astype(str)
    #%%
    voltage_data=data[data['CV_tag']=='MV1'].drop(['delmV'],axis=1).reset_index(drop=True)
    for i in range(2,mod_no):
        voltage_data=pd.concat([voltage_data,data[data['CV_tag']=='MV'+str(i)].drop(['DateTime','Date','Time','delmV'],axis=1).reset_index(drop=True)],axis=1)
    
    voltage_data=voltage_data.drop([col for col in voltage_data.columns if "junk" in col], axis=1)
    voltage_data=voltage_data.drop([col for col in voltage_data.columns if "tag" in col], axis=1)
    
    #%%
    v_names_1=v_names.copy()
    t_names_1=t_names.copy()
    for i,n in enumerate(['DateTime','Date','Time']):
        v_names.insert(i,n)
        t_names.insert(i,n)
    
    voltage_data.columns=v_names
    
    #%%
    temperature_data=data[data['CV_tag']=='MT1'].loc[:,:'junk2'].reset_index(drop=True)
    #%%
    
    for i in range(2,mod_no):
        temperature_data=pd.concat([temperature_data,data[data['CV_tag']=='MT'+str(i)].loc[:,:'junk2'].drop(['DateTime','Date','Time'],axis=1).reset_index(drop=True)],axis=1)
        
    
    temperature_data=temperature_data.dropna(how='all',axis=1)
    temperature_data=temperature_data.drop([col for col in temperature_data.columns if "tag" in col], axis=1)
    
    #%%
    temperature_data.columns=t_names
    
    for i in range(1,t_no):
        temperature_data['MT_'+str(i)]=temperature_data['MT_'+str(i)].astype(str).str.extract(r'(\d{2}.\d{5})')
        
    
    #%%
    voltage_data.loc[:,'MV_1':'MV_'+str(v_no-1)]=voltage_data.loc[:,'MV_1':'MV_'+str(v_no-1)].replace('_0',np.nan)
    voltage_data.loc[:,'MV_1':'MV_'+str(v_no-1)]=voltage_data.loc[:,'MV_1':'MV_'+str(v_no-1)].fillna(method='ffill')
    
    for i in range(1,v_no):
        voltage_data['MV_'+str(i)]=voltage_data['MV_'+str(i)].astype(str).str.extract(r'(\d{5})')
    
    voltage_data.loc[:,'MV_1':'MV_'+str(v_no-1)]=voltage_data.loc[:,'MV_1':'MV_'+str(v_no-1)].fillna(method='ffill')
    voltage_data.loc[:,'MV_1':'MV_'+str(v_no-1)]=voltage_data.loc[:,'MV_1':'MV_'+str(v_no-1)].astype(int)
    voltage_data.loc[:,'MV_1':'MV_'+str(v_no-1)]=voltage_data.loc[:,'MV_1':'MV_'+str(v_no-1)]*0.0001
    voltage_data['Mean_V']=voltage_data.loc[:,'MV_1':'MV_'+str(v_no-1)].mean(axis=1)
    
    #%%
    temperature_data.loc[:,'MT_1':'MT_'+str(t_no-1)]=temperature_data.loc[:,'MT_1':'MT_'+str(t_no-1)].astype(float)
    temperature_data.loc[:,'MT_1':'MT_'+str(t_no-1)]=temperature_data.loc[:,'MT_1':'MT_'+str(t_no-1)].fillna(method='ffill')
    temperature_data['Mean_T']=temperature_data.loc[:,'MT_1':'MT_'+str(t_no-1)].mean(axis=1)
    
    temperature_data['delT']=temperature_data.loc[:,'MT_1':'MT_'+str(t_no-1)].max(axis=1)-temperature_data.loc[:,'MT_1':'MT_'+str(t_no-1)].min(axis=1)
    
    #%% Calculating Capacity_calculated and energy
    
    details['Time_in_sec_s']=(details['DateTime']-details['DateTime'].shift(1))/np.timedelta64(1,'s')# time difference b/w every datapt
    details['Time_in_sec']=(details['DateTime']-details['DateTime'].iloc[0])/np.timedelta64(1,'s')
    
    details['Pack Current']=details['Pack Current'].astype(str).str.extract(r'(\-{0,}\d{1,}.\d{1,})')
    details['Pack Voltage']=details['Pack Voltage'].astype(str).str.extract(r'(\d{1,}.\d{1,})')
    details['Pack Current']=details['Pack Current'].astype(float)
    details['Pack Voltage']=details['Pack Voltage'].astype(float)
    # details.loc[details['Pack Voltage']>510,'Pack Voltage']=np.nan
    
    #%%
    details['State']=np.nan
    details['Time_in_sec_s_cap']=np.where((details['Time_in_sec_s']>300),np.nan,details['Time_in_sec_s'])
    details['State']=np.where(details['Pack Current']>0,0,details['State']) # 1
    details['State']=np.where(details['Pack Current']<0,1,details['State']) # 2
    details['State']=np.where(details['Pack Current']==0,2,details['State'])
    # details['State']=details['State'].astype(str)
    
    
    # details['State']=details['State'].fillna(method='ffill')
    # details['State']=details['State'].fillna(method='bfill')
    
    #%%
    details['Cap_inst']=details['Time_in_sec_s_cap']*abs(details['Pack Current'])/3600
    details['Capacity_calculated']=details['Cap_inst'].groupby(details['State']).cumsum()
    details['Capacity_calculated_chg']=(details[details['State']==0]['Cap_inst']).cumsum()
    details['Capacity_calculated_dchg']=(details[details['State']==1]['Cap_inst']).cumsum()
    
    details['Capacity_calculated_chg']=details['Capacity_calculated_chg'].fillna(method='bfill')
    details['Capacity_calculated_dchg']=details['Capacity_calculated_dchg'].fillna(method='bfill')
    
    details['Energy_calculated']=details['Capacity_calculated']*details['Pack Voltage']
    details['Energy_calculated_chg']=details['Capacity_calculated_chg']*details['Pack Voltage']
    details['Energy_calculated_dchg']=details['Capacity_calculated_dchg']*details['Pack Voltage']
    
    #%%
    voltage_data['delV']=(voltage_data.loc[:,'MV_1':'MV_'+str(v_no-1)].max(axis=1)-voltage_data.loc[:,'MV_1':'MV_'+str(v_no-1)].min(axis=1))*1000
    
    #%%
    complete_details=pd.DataFrame()
    complete_details=pd.merge(left=details,right=voltage_data,on='DateTime') # DT # left=PV_PC_data,right=PV_PC_data,
    complete_details=pd.merge(left=complete_details,right=temperature_data,on='DateTime')
    
    #%% Adding table in document
    
    document.add_paragraph().add_run('\n\nPack Specifications:').bold=True

    table1=document.add_table(rows=5,cols=2)
    
    table1.cell(0,0).text='Pack Energy'
    table1.cell(0,1).text='110 kWh'
    table1.cell(1,0).text='Pack Capacity'
    table1.cell(1,1).text='300Ah (2*150Ah)'
    table1.cell(2,0).text='Minimum Cell Voltage'
    table1.cell(2,1).text='3.0V'
    table1.cell(3,0).text='Nominal Cell Voltage'
    table1.cell(3,1).text='3.74V'
    table1.cell(4,0).text='Maximum Cell Voltage'
    table1.cell(4,1).text='4.2V'
    
    table1.style='Table Grid'
    
    
    document.add_paragraph().add_run('\nTest Specifications:').bold=True
    
    table2=document.add_table(rows=4,cols=2)
    table2.cell(0,0).text='Test Date and Time'
    table2.cell(0,1).text=str(status_data['DateTime'].iloc[0])+' to '+str(status_data['DateTime'].iloc[-1])+' ('+ str((status_data['DateTime'].iloc[-1]-status_data['DateTime'].iloc[0])/np.timedelta64(1,'m')) +' minutes)'
    table2.cell(1,0).text='Discharging Time Period'
    table2.cell(2,0).text='Charging Time Period'
    table2.cell(3,0).text='C-rate'
    
    table2.style='Table Grid'
    
    l=7
    b=4
    x_axis='DateTime'
    #%%
    # document.add_page_break()
    document.add_paragraph().add_run('Precharge Relay:').bold=True
    #%%
    # status_data['Precharge Relay']=status_data['Precharge Relay'].replace('Open','OFF')
    plt.figure(figsize=(l,b))
    plt.plot(status_data[x_axis],status_data['Precharge Relay'])
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    plt.xlabel(x_axis,fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Precharge Relay',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Precharge Relay.png',dpi=1200)
    
    document.add_picture(img_path+'\\Precharge Relay.png')
    document.add_paragraph().add_run('**Note: The “DateTime” axis in the graphs is read as DD/MM/YYYY h:m:s (i.e., // ::, // :: and so on. ').underline=True
    
    #%%
    document.add_paragraph().add_run('Main Relay:').bold=True
    # plt.plot()
    plt.figure(figsize=(l,b))
    plt.plot(status_data[x_axis],status_data['Main Relay'])
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    plt.xlabel(x_axis,fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Main Relay',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Main Relay.png',dpi=1200)
    
    document.add_picture(img_path+'\\Main Relay.png')
    
    #%%
    document.add_paragraph().add_run('Pack Status:').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(status_data[x_axis],status_data['Pack Status'])
    plt.xlabel(x_axis,fontweight='bold')
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    plt.grid(linestyle='dotted')
    plt.title('Pack Status',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Pack Status.png',dpi=1200)
    
    document.add_picture(img_path+'\\Pack Status.png')
    
    
    #%%
    document.add_paragraph().add_run('Pack Current:').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(details[x_axis],details['Pack Current'])
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    
    plt.xlabel(x_axis,fontweight='bold')
    plt.ylabel('Current(A)',fontweight='bold')
    
    plt.grid(linestyle='dotted')
    plt.title('Pack Current',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Pack Current.png',dpi=1200)
    
    document.add_picture(img_path+'\\Pack Current.png')
    
    #%%
    document.add_paragraph().add_run('Pack Capacity (calculated):').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(details[x_axis],details['Capacity_calculated_dchg'])
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    
    plt.xlabel(x_axis,fontweight='bold')
    plt.ylabel('Capacity(Ah)',fontweight='bold')
    
    plt.grid(linestyle='dotted')
    plt.title('Pack Discharging Capacity (Calculated)',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Pack Discharging Capacity_calculated.png',dpi=1200)
    
    document.add_picture(img_path+'\\Pack Discharging Capacity_calculated.png')
    document.add_paragraph('In this discharging session, the capacity expended by the pack is about %.3fAh ().' %details['Capacity_calculated_dchg'].max())
    
    #%%
    plt.figure(figsize=(l,b))
    plt.plot(details[x_axis],details['Capacity_calculated_chg'])
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    
    plt.xlabel(x_axis,fontweight='bold')
    plt.ylabel('Capacity(Ah)',fontweight='bold')
    
    plt.grid(linestyle='dotted')
    plt.title('Pack Charging Capacity (Calculated)',fontweight='bold')
    plt.tight_layout()
    
    plt.savefig(img_path+'\\Pack Charging Capacity_calculated.png',dpi=1200)
    
    document.add_picture(img_path+'\\Pack Charging Capacity_calculated.png')
    document.add_paragraph('In this charging session, the capacity transferred to the pack is about %.3fAh (). ' %details['Capacity_calculated_chg'].max())
    document.add_paragraph().add_run('Note: These values were calculated from the current (in amperes) values and time (in hours).').underline=True
    
    #%%
    document.add_paragraph().add_run('Pack Energy (calculated):').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(details[x_axis],details['Energy_calculated_dchg']*0.001)
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    
    plt.xlabel(x_axis,fontweight='bold')
    plt.ylabel('Energy(kWh)',fontweight='bold')
    
    plt.grid(linestyle='dotted')
    plt.title('Pack Discharging Energy (Calculated)',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Pack Discharging Energy_calculated.png',dpi=1200)
    
    document.add_picture(img_path+'\\Pack Discharging Energy_calculated.png')
    document.add_paragraph('The Pack Energy expended in this discharging session (approx.  minutes) is about %.3fkWh.' % (details['Energy_calculated_dchg'].max()*0.001))
    
    #%%
    plt.figure(figsize=(l,b))
    plt.plot(details[x_axis],details['Energy_calculated_chg']*0.001)
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    
    plt.xlabel(x_axis,fontweight='bold')
    plt.ylabel('Energy(kWh)',fontweight='bold')
    
    plt.grid(linestyle='dotted')
    plt.title('Pack Charging Energy (Calculated)',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Pack Charging Energy_calculated.png',dpi=1200)
    
    document.add_picture(img_path+'\\Pack Charging Energy_calculated.png')
    document.add_paragraph('The Pack Energy expended in this charging session () is about %.3fkWh.' % (details['Energy_calculated_chg'].max()*0.001))
    document.add_paragraph().add_run('Note: These values were calculated from the calculated capacity (in Ah) values and pack voltage (in volts).').underline=True
    
    
    #%%
    # document.add_paragraph().add_run('Pack Energy (BMS):').bold=True
    
    # plt.figure(figsize=(l,b))
    # plt.plot(details[x_axis],details['Discharging Energy'])
    
    # # xmin, xmax = ax.get_xlim()
    # # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    # # ymin, ymax = ax.get_ylim()
    # # ax.set_yticks(np.round(np.linspace(ymin, ymax, 20), 2))
    
    # plt.xlabel(x_axis,fontweight='bold')
    # plt.ylabel('Energy(Wh)',fontweight='bold')
    
    # plt.grid(linestyle='dotted')
    # plt.title('Pack Discharging Energy (BMS)',fontweight='bold')
    # plt.tight_layout()
    # plt.savefig(img_path+'\\Pack Discharging Energy_BMS.png',dpi=1200)
    
    # document.add_picture(img_path+'\\Pack Discharging Energy_BMS.png')
    # p_bms_d_engy=document.add_paragraph('The ')
    # p_bms_d_engy.add_run('discharging pack energy ').bold=True
    # p_bms_d_engy.add_run('is about ')
    # p_bms_d_engy.add_run('%.3fWh.' %details['Discharging Energy'].max()).bold=True
    
    #%%
    # plt.figure(figsize=(l,b))
    # plt.plot(details[x_axis],details['Charging Energy'])
    
    # # xmin, xmax = ax.get_xlim()
    # # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    
    # # ymin, ymax = ax.get_ylim()
    # # ax.set_yticks(np.round(np.linspace(ymin, ymax, 20), 2))
    
    # plt.xlabel(x_axis,fontweight='bold')
    # plt.ylabel('Energy(Wh)',fontweight='bold')
    
    # plt.grid(linestyle='dotted')
    # plt.title('Pack Charging Energy (BMS)',fontweight='bold')
    # plt.tight_layout()
    # plt.savefig(img_path+'\\Pack Charging Energy_BMS.png',dpi=1200)
    
    # document.add_picture(img_path+'\\Pack Charging Energy_BMS.png')
    # p_bms_c_engy=document.add_paragraph('The ')
    # p_bms_c_engy.add_run('charging pack energy ').bold=True
    # p_bms_c_engy.add_run('is about ')
    # p_bms_c_engy.add_run('%.3fWh.' %details['Charging Energy'].max()).bold=True
    
    #%%
    document.add_paragraph().add_run('Pack Voltage:').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(details[x_axis],details['Pack Voltage'])
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    
    plt.xlabel(x_axis,fontweight='bold')
    plt.ylabel('Voltage(V)',fontweight='bold')
    
    plt.grid(linestyle='dotted')
    plt.title('Pack Voltage',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Pack Voltage.png',dpi=1200)
    
    document.add_picture(img_path+'\\Pack Voltage.png')
    pv=document.add_paragraph('During ')
    pv.add_run('discharge, ').bold=True
    pv.add_run('the pack goes from to . The pack voltage at rest after discharge is . At ')
    pv.add_run('charge, ').bold=True
    pv.add_run('the pack goes from to . ')
    pv.add_run('In the end, the voltage stabilizes at %sV during the end rest.' %details['Pack Voltage'].iloc[-1])
    
    #%%
    document.add_paragraph().add_run('Cell Voltage:').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(voltage_data[x_axis],voltage_data.loc[:,'MV_1':'MV_'+str(v_no-1)],marker='o',markersize=3) #
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    plt.ylabel('Voltage(V)',fontweight='bold')
    plt.xlabel(x_axis,fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Cell Voltage',fontweight='bold')
    plt.tight_layout()
    
    #%%
    document.add_paragraph().add_run('Average Voltage:').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(voltage_data[x_axis],(voltage_data['Mean_V']))
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    plt.ylabel('Voltage(V)',fontweight='bold')
    plt.xlabel(x_axis,fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Average Voltage',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Average Voltage.png',dpi=1200)
    
    document.add_picture(img_path+'\\Average Voltage.png')
    av=document.add_paragraph('The average cell voltage during ')
    av.add_run('discharge ').bold=True
    av.add_run('goes from ')
    av.add_run('V to V. ').bold=True
    av.add_run('The average cell voltage at the rest after discharge is V. At ')
    av.add_run('charge').bold=True
    av.add_run(', the average cell voltage goes from ')
    av.add_run('V to V.').bold=True
    av.add_run('In the end rest, it  stabilizes at %.3fV.' % (voltage_data['Mean_V'].iloc[-1]*0.0001))
    
    #%%
    document.add_paragraph().add_run('Voltage Difference (delV):').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(voltage_data[x_axis],(voltage_data['delV']))
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    plt.ylabel('Voltage(mV)',fontweight='bold')
    plt.xlabel(x_axis,fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Voltage difference (delV)',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\delV.png',dpi=1200)
    
    document.add_picture(img_path+'\\delV.png')
    dv=document.add_paragraph('During discharge, the ')
    dv.add_run('maximum delV (difference between the highest and lowest cell voltage), ').bold=True
    dv.add_run('is about ')
    dv.add_run('mV, at the end of discharge. ').bold=True
    dv.add_run('At the rest after discharging, the delmV is mV. In the charging session, the maximum delV is mV. ')
    dv.add_run('The pack stabilizes at %.3fmV delV at the end rest. '% (voltage_data['delV'].iloc[-1]))
    
    
    #%%
    document.add_paragraph().add_run('Cell Temperature:').bold=True
    
    plt.figure(figsize=(l,b))
    plt.plot(temperature_data[x_axis],temperature_data.loc[:,'MT_1':'MT_'+str(t_no-1)])
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.xlabel(x_axis,fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Cell Temperature',fontweight='bold')
    ct=document.add_paragraph('The maximum temperature reached is around ')
    ct.add_run('%.3f'%(temperature_data.loc[:,'MT_1':'MT_'+str(t_no-1)].max(axis=1).max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True
    
    #%%
    document.add_paragraph().add_run('Average temperature:').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(temperature_data[x_axis],temperature_data['Mean_T'])
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.xlabel(x_axis,fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Average Temperature',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Average Temperature.png',dpi=1200)
    
    document.add_picture(img_path+'\\Average Temperature.png')
    at=document.add_paragraph('As can be seen, the ')
    at.add_run('average cell temperature ').bold=True
    at.add_run('stays between ')
    at.add_run(u'\N{DEGREE SIGN}'+'C to '+u'\N{DEGREE SIGN}'+'C ').bold=True
    at.add_run('during ')
    at.add_run('discharging').bold=True
    at.add_run(' and ')
    at.add_run(u'\N{DEGREE SIGN}'+'C to '+u'\N{DEGREE SIGN}'+'C ').bold=True
    at.add_run('during ')
    at.add_run('charging.').bold=True
    
    #%%
    document.add_paragraph().add_run('Temperature Difference (delT):').bold=True
    plt.figure(figsize=(l,b))
    plt.plot(temperature_data[x_axis],temperature_data['delT'])
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.xlabel(x_axis,fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Temperature difference (delT)',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\delT.png',dpi=1200)
    
    document.add_picture(img_path+'\\delT.png')
    dt=document.add_paragraph('The ')
    dt.add_run('maximum delT ').bold=True
    dt.add_run('is about ')
    dt.add_run('%.3f' %temperature_data['delT'].max()).bold=True
    dt.add_run(u'\N{DEGREE SIGN}'+'C.').bold=True
    
    # #%%
    # plt.plot()
    # plt.figure(figsize=(l,b))
    # plt.plot(voltage_data[x_axis],voltage_data.loc[:,'MV_190':'MV_192'])
    
    # # xmin, xmax = ax.get_xlim()
    # # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    # plt.ylabel('Voltage(mV)',fontweight='bold')
    # plt.xlabel(x_axis,fontweight='bold')
    # plt.grid(linestyle='dotted')
    # plt.title('Cell Voltage',fontweight='bold')
    
    
    #%%
    # #%%
    # plt.plot()
    # plt.figure(figsize=(l,b))
    # plt.plot(temperature_data[x_axis],temperature_data.loc[:,'MT_1':'MT_48'])
    
    # xmin, xmax = ax.get_xlim()
    # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    # plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    # plt.xlabel('Time',fontweight='bold')
    # plt.legend(t_names)
    # plt.grid(linestyle='dotted')
    # plt.title('Cell Temperature',fontweight='bold')
    
    
    
    
    
    #%%
    # plt.plot()
    # plt.figure(figsize=(l,b))
    # plt.plot(details[x_axis],details['Capacity_calculated'])
    
    # # xmin, xmax = ax.get_xlim()
    # # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    
    # plt.xlabel('Time',fontweight='bold')
    # plt.ylabel('Capacity(Ah)',fontweight='bold')
    
    # plt.grid(linestyle='dotted')
    # plt.title('Pack Capacity (Calculated)',fontweight='bold')
    
    
    
    #%%
    # plt.plot()
    # plt.figure(figsize=(l,b))
    # plt.plot(details[x_axis],details['Energy_calculated']*0.001)
    
    # # xmin, xmax = ax.get_xlim()
    # # ax.set_xticks(np.round(np.linspace(xmin, xmax, 18), 2))
    
    # plt.xlabel('Time',fontweight='bold')
    # plt.ylabel('Energy(kWh)',fontweight='bold')
    
    # plt.grid(linestyle='dotted')
    # plt.title('Pack Energy (Calculated)',fontweight='bold')
    
    
    
    
    #%%
    # path=img_path+'\\'+file.rsplit('\\',1)[1].rsplit('.',1)[0]+'_modified.xlsx'
    # document.save(img_path+'\\'+file.rsplit('\\',1)[1].rsplit('.',1)[0]+'_observations.docx')
    
    # #%%
    # writer = pd.ExcelWriter(path, engine = 'xlsxwriter')
    # voltage_data.to_excel(writer, sheet_name = 'Cell Voltage',index=False)
    # temperature_data.to_excel(writer, sheet_name = 'Cell Temperature',index=False)
    # details.to_excel(writer,sheet_name='Pack Details',index=False)
    # status_data.to_excel(writer,sheet_name='Pack Status',index=False)
    # max_min_data.to_excel(writer,sheet_name='Max Min Mean',index=False)
    # writer.close()
    
    # %%
#     f=open(img_path+'\\'+file.rsplit('\\',1)[1].rsplit('.',1)[0]+'_observations.txt',"w")
#     print('======================================================================',file=f)
#     print('---------------------%s---------------------\n' %file_name,file=f)
#     print('Datapoints:%s' % (len(voltage_data)),file=f)
#     print('Timeframe: %s to %s' % (voltage_data['DateTime'].iloc[0],voltage_data['DateTime'].iloc[-1]),file=f)
#     print('Total Time Spent:%s' % (datetime.timedelta(seconds=(voltage_data['DateTime'].iloc[-1]-voltage_data['DateTime'].iloc[0])/np.timedelta64(1,'s'))),file=f)
#     print('---------------------------',file=f)
    

# # #%%

#     vol_col_check=v_names_1
#     temp_col_check=t_names_1


#     if len(complete_details[complete_details['State']==0])!=0:
#         c_dt=pd.to_datetime(complete_details[complete_details['State']==0]['DateTime'].tolist())
        
#         chg_vol=complete_details[complete_details['DateTime'].isin(c_dt)]
#         print('Timeframe: %s to %s' % (chg_vol['DateTime'].iloc[0],chg_vol['DateTime'].iloc[-1]),file=f)
#         print('Chg Time Spent:%s' % (datetime.timedelta(seconds=(chg_vol['Time_in_sec_s'].sum()))),file=f)
#         print('Maximum Charging Current:%.3fA ; Minimum Charging Current:%.3fA' % (chg_vol['Pack Current'].max(),chg_vol['Pack Current'].min()),file=f)
#         print('Chg Pack Voltage range:%sV to %sV' %(chg_vol['Pack Voltage'].iloc[0],chg_vol['Pack Voltage'].iloc[-1]),file=f)
#         print('---------------------------',file=f)
    
#         print('Chg Min Vol:%.3fV' %((chg_vol.loc[:,'MV_1':'MV_'+str(v_no-1)]).min(axis=1).min()),file=f)
#         print('Chg Max Vol:%.3fV' %((chg_vol.loc[:,'MV_1':'MV_'+str(v_no-1)]).max(axis=1).max()),file=f)
#         print('Chg Average Vol:%.3fV' %((chg_vol['Mean_V'])).mean(),file=f)
#         print('Chg Start Average Vol:%.3fV' %((chg_vol['Mean_V']).iloc[0]),file=f)
#         print('Chg End Average Vol:%.3fV' %((chg_vol['Mean_V'].iloc[-1])),file=f)
        
#         print('Chg Minimum delV:%.3fmV' %((chg_vol['delV'])).min(),file=f)
#         print('Chg Maximum delV:%.3fmV' %((chg_vol['delV'])).max(),file=f)
#         print('Chg Average delV:%.3fmV' %((chg_vol['delV'])).mean(),file=f)
#         print('Chg start delV:%.3fmV' %((chg_vol['delV'].iloc[0])),file=f)
#         print('Chg end delV:%.3fmV' %((chg_vol['delV'].iloc[-1])),file=f)
#         print('---------------------------',file=f)
 
        
#         chg_temp=temperature_data[temperature_data['DateTime'].isin(c_dt)]
#         print('Chg Min Temp:%.3f degC' %(chg_temp.loc[:,'MT_1':'MT_'+str(t_no-1)].min(axis=1).min()),file=f)
#         print('Chg Max Temp:%.3f degC' %(chg_temp.loc[:,'MT_1':'MT_'+str(t_no-1)].max(axis=1).max()),file=f)
#         print('Chg Average Temp:%.3f degC' %(chg_temp['Mean_T'].mean()),file=f)
#         print('Chg Start Average Temp:%.3f degC' %(chg_temp['Mean_T'].iloc[0]),file=f)
#         print('Chg End Average Temp:%.3f degC' %(chg_temp['Mean_T'].iloc[-1]),file=f)
        
#         print('Chg Minimum delT:%.3f degC' %(chg_temp['delT'].min()),file=f)
#         print('Chg Maximum delT:%.3f degC' %(chg_temp['delT'].max()),file=f)
#         print('Chg Average delT:%.3f degC' %(chg_temp['delT'].mean()),file=f)
#         print('Chg Start delT:%.3f degC' %(chg_temp['delT'].iloc[0]),file=f)
#         print('Chg End delT:%.3f degC' %(chg_temp['delT'].iloc[-1]),file=f)
#         print('---------------------------',file=f)
#         print('Cells with (lower)erroneous chg voltage values: %s' %([col for col in vol_col_check if complete_details[complete_details['State']==0][col].lt(2.3000).any()]),file=f)
#         print('Cells with (higher)erroneous chg voltage values: %s' %([col for col in vol_col_check if complete_details[complete_details['State']==0][col].gt(4.3000).any()]),file=f)
        
#         print('Cells with (lower)erroneous chg temp values: %s' %([col for col in temp_col_check if chg_temp[col].lt(17).any()]),file=f)
#         print('Cells with (higher)erroneous chg temp values: %s' %([col for col in temp_col_check if chg_temp[col].gt(35).any()]),file=f)
#     print('---------------------------',file=f)
#     if len(complete_details[complete_details['State']==1])!=0:
#         d_dt=pd.to_datetime(complete_details[complete_details['State']==1]['DateTime'].tolist())

#         dchg_vol=complete_details[complete_details['DateTime'].isin(d_dt)]
#         print('Timeframe: %s to %s' % (dchg_vol['DateTime'].iloc[0],dchg_vol['DateTime'].iloc[-1]),file=f)
#         print('Dchg Time Spent:%s' % (datetime.timedelta(seconds=(dchg_vol['Time_in_sec_s'].sum()))),file=f)
#         print('Maximum Discharging Current:%.3fA ; Minimum Discharging Current:%.3fA' % (dchg_vol['Pack Current'].max(),dchg_vol['Pack Current'].min()),file=f)
#         print('DChg Pack Voltage range:%sV to %sV' %(dchg_vol['Pack Voltage'].iloc[0],dchg_vol['Pack Voltage'].iloc[-1]),file=f)
#         print('---------------------------',file=f)
    
#         print('Dchg Min Vol:%.3fV' %((dchg_vol.loc[:,'MV_1':'MV_'+str(v_no-1)]).min(axis=1).min()),file=f)
#         print('Dchg Max Vol:%.3fV' %((dchg_vol.loc[:,'MV_1':'MV_'+str(v_no-1)]).max(axis=1).max()),file=f)
#         print('Dchg Average Vol:%.3fV' %((dchg_vol['Mean_V'])).mean(),file=f)
#         print('Dchg Start Average Vol:%.3fV' %((dchg_vol['Mean_V'].iloc[0])),file=f)
#         print('Dchg End Average Vol:%.3fV' %((dchg_vol['Mean_V'].iloc[-1])),file=f)
        
#         print('Dchg Minimum delV:%.3fmV' %((dchg_vol['delV'])).min(),file=f) 
#         print('Dchg Maximum delV:%.3fmV' %((dchg_vol['delV'])).max(),file=f) 
#         print('Dchg Average delV:%.3fmV' %((dchg_vol['delV'])).mean(),file=f) 
#         print('Dchg Average delV:%.3fmV' %((dchg_vol['delV'].iloc[0])),file=f) 
#         print('Dchg Average delV:%.3fmV' %((dchg_vol['delV'].iloc[-1])),file=f) 
#         print('---------------------------',file=f)
        
        
#         dchg_temp=temperature_data[temperature_data['DateTime'].isin(d_dt)]  
#         print('Dchg Min Temp:%.3f degC' %(dchg_temp.loc[:,'MT_1':'MT_'+str(t_no-1)].min(axis=1).min()),file=f)
#         print('Dchg Max Temp:%.3f degC' %(dchg_temp.loc[:,'MT_1':'MT_'+str(t_no-1)].max(axis=1).max()),file=f)
#         print('Dchg Average Temp:%.3f degC' %(dchg_temp['Mean_T'].mean()),file=f)
#         print('Dchg Start Average Temp:%.3f degC' %(dchg_temp['Mean_T'].iloc[0]),file=f)
#         print('Dchg End Average Temp:%.3f degC' %(dchg_temp['Mean_T'].iloc[-1]),file=f)
        
#         print('Dchg Minimum delT:%.3f degC' %(dchg_temp['delT'].min()),file=f)
#         print('Dchg Maximum delT:%.3f degC' %(dchg_temp['delT'].max()),file=f)
#         print('Dchg Average delT:%.3f degC' %(dchg_temp['delT'].mean()),file=f)
#         print('Dchg Start delT:%.3f degC' %(dchg_temp['delT'].iloc[0]),file=f)
#         print('Dchg end delT:%.3f degC' %(dchg_temp['delT'].iloc[-1]),file=f)
#         print('---------------------------',file=f)
#         print('Cells with (lower)erroneous dchg voltage values: %s' %([col for col in vol_col_check if complete_details[complete_details['State']==1][col].lt(2.3000).any()]),file=f)
#         print('Cells with (higher)erroneous dchg voltage values: %s' %([col for col in vol_col_check if complete_details[complete_details['State']==1][col].gt(4.3000).any()]),file=f)
        
#         print('Cells with (lower)erroneous dchg temp values: %s' %([col for col in temp_col_check if dchg_temp[col].lt(17).any()]),file=f)
#         print('Cells with (higher)erroneous dchg temp values: %s' %([col for col in temp_col_check if dchg_temp[col].gt(35).any()]),file=f)
#     print('---------------------------',file=f)
#     if len(complete_details[complete_details['State']==2])!=0:

#         r_dt=pd.to_datetime(complete_details[complete_details['State']==2]['DateTime'].tolist())
#         rst_vol=complete_details[complete_details['DateTime'].isin(r_dt)]
#         print('Timeframe: %s to %s' % (rst_vol['DateTime'].iloc[0],rst_vol['DateTime'].iloc[-1]),file=f)
#         print('Rest Time Spent:%s' % (datetime.timedelta(seconds=(rst_vol['Time_in_sec_s'].sum()))),file=f)
#         print('Rest Pack Voltage range:%sV to %sV' %(rst_vol['Pack Voltage'].iloc[0],rst_vol['Pack Voltage'].iloc[-1]),file=f)
#         print('---------------------------',file=f)
#         print('Rest Min Vol:%.3fV' %((rst_vol.loc[:,'MV_1':'MV_'+str(v_no-1)]).min(axis=1).min()),file=f)
#         print('Rest Max Vol:%.3fV' %((rst_vol.loc[:,'MV_1':'MV_'+str(v_no-1)]).max(axis=1).max()),file=f)
#         print('Rest Average Vol:%.3fV' %((rst_vol['Mean_V'])).mean(),file=f)
#         print('Rest Start Average Vol:%.3fV' %((rst_vol['Mean_V'].iloc[0])),file=f)
#         print('Rest End Average Vol:%.3fV' %((rst_vol['Mean_V'].iloc[-1])),file=f)
        
#         print('Rest Minimum delV:%.3fmV' %((rst_vol['delV'])).min(),file=f)
#         print('Rest Maximum delV:%.3fmV' %((rst_vol['delV'])).max(),file=f) 
#         print('Rest Average delV:%.3fmV' %((rst_vol['delV'])).mean(),file=f) 
#         print('Rest Start delV:%.3fmV' %((rst_vol['delV'].iloc[0])),file=f) 
#         print('Rest End delV:%.3fmV' %((rst_vol['delV'].iloc[-1])),file=f) 
#         print('---------------------------',file=f)
        
        
#         rst_temp=temperature_data[temperature_data['DateTime'].isin(r_dt)]
#         print('Rest Min Temp:%.3f degC' %(rst_temp.loc[:,'MT_1':'MT_'+str(t_no-1)].min(axis=1).min()),file=f)
#         print('Rest Max Temp:%.3f degC' %(rst_temp.loc[:,'MT_1':'MT_'+str(t_no-1)].max(axis=1).max()),file=f)
#         print('Rest Average Temp:%.3f degC' %(rst_temp['Mean_T'].mean()),file=f)
#         print('Rest Start Average Temp:%.3f degC' %(rst_temp['Mean_T'].iloc[0]),file=f)
#         print('Rest End  Average Temp:%.3f degC' %(rst_temp['Mean_T'].iloc[-1]),file=f)
        
#         print('Rest Minimum delT:%.3f degC' %(rst_temp['delT'].min()),file=f)
#         print('Rest Maximum delT:%.3f degC' %(rst_temp['delT'].max()),file=f)
#         print('Rest Average delT:%.3f degC' %(rst_temp['delT'].mean()),file=f)
#         print('Rest Start delT:%.3f degC' %(rst_temp['delT'].iloc[0]),file=f)
#         print('Rest End delT:%.3f degC' %(rst_temp['delT'].iloc[-1]),file=f)
        
#         print('---------------------------',file=f)
#         print('Cells with (lower)erroneous Rest voltage values: %s' %([col for col in vol_col_check if complete_details[complete_details['State']==2][col].lt(2.3000).any()]),file=f)
#         print('Cells with (higher)erroneous Rest voltage values: %s' %([col for col in vol_col_check if complete_details[complete_details['State']==2][col].gt(4.3000).any()]),file=f)
        
#         print('Cells with (lower)erroneous Rest temp values: %s' %([col for col in temp_col_check if rst_temp[col].lt(17).any()]),file=f)
#         print('Cells with (higher)erroneous Rest temp values: %s' %([col for col in temp_col_check if rst_temp[col].gt(35).any()]),file=f)
#     print('---------------------------',file=f)
#     print('Maximum Charging Capacity:%.3fAh ; Minimum Charging Capacity:%.3fAh' % (complete_details['Capacity_calculated_chg'].max(),complete_details['Capacity_calculated_chg'].min()),file=f)
#     print('Maximum Discharging Capacity:%.3fAh ; Minimum Discharging Capacity:%.3fAh' % (complete_details['Capacity_calculated_dchg'].max(),complete_details['Capacity_calculated_dchg'].min()),file=f)
#     print('Maximum Charging Energy:%.3fkWh ; Minimum Charging Energy:%.3fkWh' % ((complete_details['Energy_calculated_chg']*0.001).max(),(complete_details['Energy_calculated_chg']*0.001).min()),file=f)
#     print('Maximum Discharging Energy:%.3fkWh ; Minimum Discharging Energy:%.3fkWh' % ((complete_details['Energy_calculated_dchg']*0.001).max(),(complete_details['Energy_calculated_dchg']*0.001).min()),file=f)
    
#     print('Cells with (lower)erroneous voltage values: %s' %([col for col in vol_col_check if complete_details[col].lt(2.3000).any()]),file=f)
#     print('Cells with (higher)erroneous voltage values: %s' %([col for col in vol_col_check if complete_details[col].gt(4.3000).any()]),file=f)
    
#     print('Cells with (lower)erroneous temp values: %s' %([col for col in temp_col_check if temperature_data[col].lt(17).any()]),file=f)
#     print('Cells with (higher)erroneous temp values: %s' %([col for col in temp_col_check if temperature_data[col].gt(35).any()]),file=f)
    
#     print('======================================================================\n',file=f)
#     print('\n',file=f)
    
#     f.close()

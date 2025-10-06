# -*- coding: utf-8 -*-
"""
Created on Fri Dec 29 15:04:24 2023

@author: IITM
"""

#%% Code Description

"""
1. This code is used to process the CALB cell data
2. 
"""

#%% Import necessary libraries
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import glob
import docx

#%% Load all the necessary files using 'glob'
dlogger_chg_dchg_files=glob.glob("D:\\Benisha\\CALB\\AL Calb Cell Test\\AL Test\\Data Loger data\\Charge & Discharge Test\\*\\*.csv")
dlogger_hppc_files=glob.glob("D:\\Benisha\\CALB\\AL Calb Cell Test\\AL Test\\Data Loger data\\HPPC TEST\\*.csv")
tester_chg_dchg_files=glob.glob("D:\\Benisha\\CALB\\AL Calb Cell Test\\AL Test\\Tester Data\\Charge and Discharge Test\\*\\*.xlsx")
tester_hppc_files=glob.glob("D:\\Benisha\\CALB\\AL Calb Cell Test\\AL Test\\Tester Data\\HPPC\\*\\*.xlsx")
tester_temp_chg_dchg_files=glob.glob("D:\\Benisha\\CALB\\AL Calb Cell Test\\AL Test\\Tester Data\\Charge and Discharge Test\\*\\Temperature_data\\*.xlsx")
tester_temp_hppc_files=glob.glob("D:\\Benisha\\CALB\\AL Calb Cell Test\\AL Test\\Tester Data\\HPPC\\*\\Temperature_data\\*.xlsx")

#%% 
# #%% dlog chg/dchg
# f="D:\\Benisha\\CALB\\AL Calb Cell Test\\AL Test\\Data Loger data\\Charge & Discharge Test\\0.5C_0.5C_0.3C_0.3C\\Data 8199 2391 25-Dec-23 13_42_32_2_combined_file.csv"

# # for f in dlogger_chg_dchg_files:
# cols=['Scan','Time','101 <C1_HF_Negative Terminal side> (VDC)','Alarm 101','102 <C1_TC_Negative Terminal side> (C)','Alarm 102','103 <C1_HF_Positive Terminal side> (VDC)','Alarm 103','104 <C1_TC_Positive Terminal side> (C)','Alarm 104','105 <C1_HF_Cell Center> (VDC)','Alarm 105','106 <C1_TC_Cell Center> (C)','Alarm 106','107 <C1_HF_Cell Side> (VDC)','Alarm 107','108 <C1_TC_Cell Side> (C)','Alarm 108','109 <C1_HF_Cell Bottom> (VDC)','Alarm 109','110 <C1_TC_Cell Bottom> (C)','Alarm 110','111 <C2_TC_Positive Terminal side> (C)','Alarm 111','112 <C2_HF_Positive Terminal side> (VDC)','Alarm 112','113 <C2_HF_Negative Terminal side> (C)','Alarm 113','114 <C2_TC_Negative Terminal side> (VDC)','Alarm 114','115 <C2_HF_Cell Center> (C)','Alarm 115','116 <C2_TC_Cell Center> (VDC)','Alarm 116','117 <C2_HF_Cell Side> (C)','Alarm 117','118 <C2_TC_Cell Side> (VDC)','Alarm 118','119 <C2_HF_Cell Bottom> (C)','Alarm 119','120 <C2_TC_Cell Bottom> (VDC)','Alarm 120']
# data=pd.read_csv(f,encoding='utf-8',sep=',',header=5)

# #%% dlog hppc
# f="D:\\Benisha\\CALB\\AL Calb Cell Test\\AL Test\\Data Loger data\\HPPC TEST\\Data 8199 2391 26-Dec-23 16_22_29.csv"

# data=pd.read_csv(f,encoding='utf-16',header=28)

# data['Time']=data['Time'].astype(str)
# data[['DateTime','junk']]=data['Time'].str.rsplit(':',1,expand=True)

# data['DateTime']=pd.to_datetime(data['DateTime'],errors='coerce')
# data=data[['DateTime'] + [col for col in data.columns if col!='DateTime']]
# #%% tester chg/dchg and temp
# f="D:\\Benisha\\CALB\\AL Calb Cell Test\\AL Test\\Tester Data\\Charge and Discharge Test\\0.5C_0.5C&0.3C_0.3C\\25deg data.xlsx"

# data=pd.read_excel(f,sheet_name='record')

# #%% tester hppc and temp
# f="D:\\Benisha\\CALB\\AL Calb Cell Test\\AL Test\\Tester Data\\HPPC\\45 Deg\\45 deg ambient temperature.xlsx"

# data=pd.read_excel(f,sheet_name='record')

#%% Creating a Word document

document=docx.Document()
document.styles['Normal'].font.name='Calibri'
H1=document.add_heading()
H1.alignment=1
H1.add_run('Calb Cell Observations').font.size=docx.shared.Pt(18)

document.add_paragraph().add_run('\n\nCell Specifications:').bold=True

table1=document.add_table(rows=3,cols=2)

table1.cell(0,0).text='Chemistry'
table1.cell(0,1).text='LFP'
table1.cell(1,0).text='Capacity'
table1.cell(1,1).text='230Ah'
table1.cell(2,0).text='No. of cells tested'
table1.cell(2,1).text='2'

table1.style='Table Grid'

# Setting the length and bradth of the plot window
# Values set are the optimal ones for displaying in a word document
l=7
b=4

#%% Tester Cycling data

document.add_paragraph().add_run('\n\nTester Charge Discharge data:').bold=True

# Loop through the files
for f in tester_chg_dchg_files:
        
    data=pd.read_excel(f,sheet_name='record') # Read tester file
    img_path=f.rsplit('\\',1)[0]
    name=f.rsplit('\\',2)[1]+'_'+f.rsplit('\\',2)[2].rsplit('.',1)[0]
    
    data['Date']=pd.to_datetime(data['Date'],errors='coerce')
    data=data.reset_index(drop=True)

    data['Time_in_sec_s']=(data['Date']-data['Date'].shift(1))/np.timedelta64(1,'s') # Time difference b/w every datapt
    data['Time_in_sec']=(data['Date']-data['Date'].iloc[0])/np.timedelta64(1,'s')

    
    data['Current(A)']=data['Current(A)'].astype(float)
    data['State']=np.nan
    # data['State']=np.where(data['Current(A)']<0,'DChg','Chg')
    data['State']=np.where(data['Current(A)']<0,1,2)
    
    data['Cycle_No']=((data['State'].shift()-data['State'])<0).cumsum()
    # data['Cycle_No']=(data['State']=='DChg').cumsum()
    
    data['Time']=np.nan
    x_axis='Time'

    # document.add_paragraph().add_run('Current ('+name+'):').bold=True
    # plt.figure(figsize=(l,b))
    # print('Current')
    # plt.plot(data[x_axis],data['Current(A)'])
    
    # plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    # plt.ylabel('Current(A)',fontweight='bold')
    
    # plt.grid(linestyle='dotted')
    # plt.title('Tester Current',fontweight='bold')
    # plt.tight_layout()
    # plt.savefig(img_path+'\\Current'+name+'.png',dpi=1200)
    # plt.close()
    
    # document.add_picture(img_path+'\\Current'+name+'.png')
  
    document.add_paragraph().add_run('Capacity ('+name+'):').bold=True
    
    plt.figure(figsize=(l,b))
    print('Capacity(Ah)')
    plt.plot(data[x_axis],data['Capacity(Ah)'])
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Capacity(Ah)',fontweight='bold')
    
    plt.grid(linestyle='dotted')
    plt.title('Cell Capacity',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Capacity'+name+'.png',dpi=1200)
    plt.close()
    
    # document.add_picture(img_path+'\\Capacity'+name+'.png')
    document.add_paragraph('In this discharging session, the capacity expended by the pack is about %sAh ()' %data['Capacity(Ah)'].max())

    # document.add_paragraph().add_run('Voltage ('+name+'):').bold=True
    # plt.figure(figsize=(l,b))
    # print('Pack Voltage')
    # plt.plot(data[x_axis],data['Voltage(V)'])
    
    # plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    # plt.ylabel('Voltage(V)',fontweight='bold')
    
    # plt.grid(linestyle='dotted')
    # plt.title('Cell Voltage',fontweight='bold')
    # plt.tight_layout()
    # # plt.savefig(img_path+'\\Voltage'+name+'.png',dpi=1200)
    # # plt.close()
    
    # document.add_picture(img_path+'\\Voltage'+name+'.png')


    # document.add_paragraph().add_run('Cell Temperature ('+name+'):').bold=True

    # plt.figure(figsize=(l,b))
    # print('Cell Temperature')
    # plt.plot(data[x_axis],data['T1(℃)'])
    
    # plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    # plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    # plt.grid(linestyle='dotted')
    # plt.title('Cell Temperature',fontweight='bold')
    # plt.tight_layout()
    # # plt.savefig(img_path+'\\Cell Temperature'+name+'.png',dpi=1200)
    # # plt.close()
    
    # document.add_picture(img_path+'\\Cell Temperature'+name+'.png')

    df=pd.DataFrame()
    
    grouped=[f[1] for f in data.groupby(['Cycle_No'])]
    for g in grouped:
        g['Time']=(g['Date']-g['Date'].iloc[0])/np.timedelta64(1,'h')
        df=pd.concat([df,g])

    for i,n in enumerate(df['Cycle_No'].unique()):
    
        fig,(ax1,ax2,ax3)=plt.subplots(3,sharex=True)
        
        ax1.plot(df[df['Cycle_No']==n][x_axis],df[df['Cycle_No']==n]['Current(A)'])
        
        ax1.set_ylabel('Current(A)',fontweight='bold')
        
        ax1.grid(linestyle='dotted')
        ax1.set_title('Current, Cell Voltage and Cell Temperature',fontweight='bold')
        
    
        ax2.plot(df[df['Cycle_No']==n][x_axis],df[df['Cycle_No']==n]['Voltage(V)'])
        
        ax2.set_ylabel('Voltage(V)',fontweight='bold')
        
        ax2.grid(linestyle='dotted')
        
        # ax2.set_title('Cell Voltage',fontweight='bold')
    
    
        ax3.plot(df[df['Cycle_No']==n][x_axis],df[df['Cycle_No']==n]['T1(℃)'])
        
        ax3.set_ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
        ax3.set_xlabel('Time Elapsed (in hours)',fontweight='bold')
        ax3.grid(linestyle='dotted')
        # ax3.set_title('Cell Temperature',fontweight='bold')
        plt.savefig(img_path+'\\Current, Cell Voltage and Cell Temperature_'+name+'.png',dpi=1200)
        plt.close()
        
        document.add_picture(img_path+'\\Current, Cell Voltage and Cell Temperature_'+name+'.png')
        pv=document.add_paragraph('During ')
        pv.add_run('discharge, ').bold=True
        pv.add_run('the pack goes from to . The pack voltage at rest after discharge is . At ')
        pv.add_run('charge, ').bold=True
        pv.add_run('the pack goes from to . ')
        pv.add_run('In the end, the voltage stabilizes at %sV during the end rest.' %data['Voltage(V)'].iloc[-1])
        
        ct=document.add_paragraph('The maximum temperature reached is around ')
        ct.add_run('%s'%(data['T1(℃)'].max())).bold=True
        ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True
    
#%% 
document.add_paragraph().add_run('Tester Charge Discharge Temperature data:').bold=True

for f1 in tester_temp_chg_dchg_files:
    data=pd.read_excel(f1,sheet_name='record')
    img_path=f1.rsplit('\\',1)[0]
    name=f1.rsplit('\\',3)[1]+'_'+f1.rsplit('\\',3)[3].rsplit('.',1)[0]
    
    data['Date']=pd.to_datetime(data['Date'],errors='coerce')
    data['Time']=(data['Date']-data['Date'].iloc[0])/np.timedelta64(1,'h')
    x_axis='Time'    

    document.add_paragraph().add_run('Temperature ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print('Temperature')
    plt.plot(data['Time' ],data['T1(℃)'])
    
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Temperature ('+name+')',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Cell Temperature_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\Cell Temperature_'+name+'.png')
    ct=document.add_paragraph('The maximum temperature reached is around ')
    ct.add_run('%s'%(data['T1(℃)'].max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True
    
    del data
    
#%% 
document.add_paragraph().add_run('DataLogger Charge Discharge Cycle Temperature data:').bold=True

#%% Loop through the dlogger Chg/Dchg files
for f1 in dlogger_chg_dchg_files:
    cols=['Scan','Time','101 <C1_HF_Negative Terminal side> (VDC)','Alarm 101','102 <C1_TC_Negative Terminal side> (C)','Alarm 102','103 <C1_HF_Positive Terminal side> (VDC)','Alarm 103','104 <C1_TC_Positive Terminal side> (C)','Alarm 104','105 <C1_HF_Cell Center> (VDC)','Alarm 105','106 <C1_TC_Cell Center> (C)','Alarm 106','107 <C1_HF_Cell Side> (VDC)','Alarm 107','108 <C1_TC_Cell Side> (C)','Alarm 108','109 <C1_HF_Cell Bottom> (VDC)','Alarm 109','110 <C1_TC_Cell Bottom> (C)','Alarm 110','111 <C2_TC_Positive Terminal side> (C)','Alarm 111','112 <C2_HF_Positive Terminal side> (VDC)','Alarm 112','113 <C2_HF_Negative Terminal side> (C)','Alarm 113','114 <C2_TC_Negative Terminal side> (VDC)','Alarm 114','115 <C2_HF_Cell Center> (C)','Alarm 115','116 <C2_TC_Cell Center> (VDC)','Alarm 116','117 <C2_HF_Cell Side> (C)','Alarm 117','118 <C2_TC_Cell Side> (VDC)','Alarm 118','119 <C2_HF_Cell Bottom> (C)','Alarm 119','120 <C2_TC_Cell Bottom> (VDC)','Alarm 120']
    try:
        data=pd.read_csv(f1,encoding='utf-8',sep=',',names=cols)
    except:
        data=pd.read_csv(f1,encoding='utf-16',sep=',',names=cols)
    data['Time']=data['Time'].astype(str)
    data[['DateTime','junk']]=data['Time'].str.rsplit(':',1,expand=True)
    
    data['DateTime']=pd.to_datetime(data['DateTime'],errors='coerce')
    data=data[['DateTime'] + [col for col in data.columns if col!='DateTime']]
    data=data.dropna(thresh=4)
    data=data.reset_index(drop=True)
    data=data.loc[2:,:]
    
    data['Time']=(data['DateTime']-data['DateTime'].iloc[0])/np.timedelta64(1,'h')
    x_axis='Time' 
        
    img_path=f1.rsplit('\\',1)[0]
    name=f1.rsplit('\\',2)[1]+'_'+f1.rsplit('\\',2)[2].rsplit('.',1)[0]
    
    C1=[col for col in cols if "HF" in col]
    for c in C1:
        data[c]=data[c].astype(float)
        data[c]=((data[c])/(3.77))*1000


    title= 'C1_HF_Negative Terminal side'  
    document.add_paragraph().add_run(title+'('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data[x_axis],data['101 <C1_HF_Negative Terminal side> (VDC)']*1000)
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Heat Flux (W/m^2)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')


   
    title= 'C1_TC_Negative Terminal side'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data[x_axis],data['102 <C1_TC_Negative Terminal side> (C)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    ct=document.add_paragraph('The '+title+' has a maximum temperature value of about ')
    ct.add_run('%s'%(data['102 <C1_TC_Negative Terminal side> (C)'].astype(float).max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True
 

    title='C1_HF_Positive Terminal side'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data[x_axis],data['103 <C1_HF_Positive Terminal side> (VDC)']*1000)
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Heat Flux (W/m^2)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')


    title='C1_TC_Positive Terminal side'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data[x_axis],data['104 <C1_TC_Positive Terminal side> (C)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    ct=document.add_paragraph('The '+title+' has a maximum temperature value of about ')
    ct.add_run('%s'%(data['104 <C1_TC_Positive Terminal side> (C)'].astype(float).max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True

    title='C1_HF_Cell Center'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data[x_axis],data['105 <C1_HF_Cell Center> (VDC)']*1000)
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Heat Flux (W/m^2)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    

    title='C1_TC_Cell Center'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data[x_axis],data['106 <C1_TC_Cell Center> (C)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    ct=document.add_paragraph('The '+title+' has a maximum temperature value of about ')
    ct.add_run('%s'%(data['106 <C1_TC_Cell Center> (C)'].astype(float).max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True    

    title='C1_HF_Cell Side'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data[x_axis],data['107 <C1_HF_Cell Side> (VDC)']*1000)
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Heat Flux (W/m^2)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    

    title='C1_TC_Cell Side'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data[x_axis],data['108 <C1_TC_Cell Side> (C)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    ct=document.add_paragraph('The '+title+' has a maximum temperature value of about ')
    ct.add_run('%s'%(data['108 <C1_TC_Cell Side> (C)'].astype(float).max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True     

    title='C1_HF_Cell Bottom'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data[x_axis],data['109 <C1_HF_Cell Bottom> (VDC)']*1000)
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Heat Flux (W/m^2)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    

    title='C1_TC_Cell Bottom'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data[x_axis],data['110 <C1_TC_Cell Bottom> (C)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    ct=document.add_paragraph('The '+title+' has a maximum temperature value of about ')
    ct.add_run('%s'%(data['110 <C1_TC_Cell Bottom> (C)'].astype(float).max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True    

    title='C2_TC_Positive Terminal side'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data[x_axis],data['111 <C2_TC_Positive Terminal side> (C)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    ct=document.add_paragraph('The '+title+' has a maximum temperature value of about ')
    ct.add_run('%s'%(data['111 <C2_TC_Positive Terminal side> (C)'].astype(float).max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True    

    title='C2_HF_Positive Terminal side'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data[x_axis],data['112 <C2_HF_Positive Terminal side> (VDC)']*1000)
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Heat Flux (W/m^2)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')


    title='C2_HF_Negative Terminal side'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data[x_axis],data['113 <C2_HF_Negative Terminal side> (C)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    

    title='C2_TC_Negative Terminal side'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data[x_axis],data['114 <C2_TC_Negative Terminal side> (VDC)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    ct=document.add_paragraph('The '+title+' has a maximum temperature value of about ')
    ct.add_run('%s'%(data['114 <C2_TC_Negative Terminal side> (VDC)'].astype(float).max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True    

    title='C2_HF_Cell Center'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data[x_axis],data['115 <C2_HF_Cell Center> (C)']*1000)
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Heat Flux (W/m^2)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    

    title='C2_TC_Cell Center'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data[x_axis],data['116 <C2_TC_Cell Center> (VDC)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    ct=document.add_paragraph('The '+title+' has a maximum temperature value of about ')
    ct.add_run('%s'%(data['116 <C2_TC_Cell Center> (VDC)'].astype(float).max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True    

    title='C2_HF_Cell Side'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data[x_axis],data['117 <C2_HF_Cell Side> (C)']*1000)
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Heat Flux (W/m^2)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    

    title='C2_TC_Cell Side'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data[x_axis],data['118 <C2_TC_Cell Side> (VDC)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    ct=document.add_paragraph('The '+title+' has a maximum temperature value of about ')
    ct.add_run('%s'%(data['118 <C2_TC_Cell Side> (VDC)'].astype(float).max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True    

    title='C2_HF_Cell Bottom'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data[x_axis],data['119 <C2_HF_Cell Bottom> (C)']*1000)
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Heat Flux (W/m^2)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    

    title='C2_TC_Cell Bottom'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data[x_axis],data['120 <C2_TC_Cell Bottom> (VDC)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    ct=document.add_paragraph('The '+title+' has a maximum temperature value of about ')
    ct.add_run('%s'%(data['120 <C2_TC_Cell Bottom> (VDC)'].astype(float).max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True    
    
    del data
#%%
document.add_paragraph().add_run('Tester HPPC data:').bold=True

for f1 in tester_hppc_files:
    data=pd.read_excel(f1,sheet_name='record')
    img_path=f1.rsplit('\\',1)[0]
    name=f1.rsplit('\\',2)[1]+'_'+f1.rsplit('\\',2)[2].rsplit('.',1)[0]
    
    data['Date']=pd.to_datetime(data['Date'],errors='coerce')
    data=data[data['Date'].dt.year>=2020]

    data['Time_in_sec_s']=(data['Date']-data['Date'].shift(1))/np.timedelta64(1,'s')# time difference b/w every datapt
    data['Time_in_sec']=(data['Date']-data['Date'].iloc[0])/np.timedelta64(1,'s')

    data['State']=np.nan
    data['State']=np.where(data['Current(A)']<0,1,2) 
    data['Cycle_No']=((data['State'].shift()-data['State'])<0).cumsum()
    
    data['Time']=np.nan
    x_axis='Time'

    data['Current(A)']=data['Current(A)'].astype(float)

    document.add_paragraph().add_run('Current ('+name+'):').bold=True
    plt.figure(figsize=(l,b))
    print('Cell Current ('+name+'):')
    plt.plot(data['Date'],data['Current(A)'])
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Current(A)',fontweight='bold')
    
    plt.grid(linestyle='dotted')
    plt.title('Tester Current',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Current.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\Current.png')
    
    
    document.add_paragraph().add_run('Capacity ('+name+'):').bold=True
    
    plt.figure(figsize=(l,b))
    print('Capacity ('+name+'):')
    plt.plot(data['Date'],data['Capacity(Ah)'])
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Capacity(Ah)',fontweight='bold')
    
    plt.grid(linestyle='dotted')
    plt.title('Cell Capacity',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Capacity.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\Capacity.png')
    document.add_paragraph('In this discharging session, the capacity expended by the pack is about %sAh ()' %data['Capacity(Ah)'].max())

    document.add_paragraph().add_run('Voltage ('+name+'):').bold=True
    plt.figure(figsize=(l,b))
    print('Pack Voltage ('+name+'):')
    plt.plot(data['Date'],data['Voltage(V)'])
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Voltage(V)',fontweight='bold')
    
    plt.grid(linestyle='dotted')
    plt.title('Cell Voltage',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Voltage.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\Voltage.png')
    pv=document.add_paragraph('During ')
    pv.add_run('discharge, ').bold=True
    pv.add_run('the pack goes from to . The pack voltage at rest after discharge is . At ')
    pv.add_run('charge, ').bold=True
    pv.add_run('the pack goes from to . ')
    pv.add_run('In the end, the voltage stabilizes at V during the end rest.')


    document.add_paragraph().add_run('Cell Temperature ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print('Cell Temperature ('+name+'):')
    plt.plot(data['Date'],data['T1(℃)'])
    
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Cell Temperature',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Cell Temperature.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\Cell Temperature.png')
    ct=document.add_paragraph('The maximum temperature reached is around ')
    ct.add_run('%s'%(data['T1(℃)'].max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True
##----------------------------------------------------------------------------------
    # df=pd.DataFrame()
    
    # grouped=[f[1] for f in data.groupby(['Cycle_No'])]
    # for g in grouped:
    #     g['Time']=(g['Date']-g['Date'].iloc[0])/np.timedelta64(1,'h')
    #     df=pd.concat([df,g])

    # for i,n in enumerate(df['Cycle_No'].unique()):
    
    #     fig,(ax1,ax2,ax3)=plt.subplots(3,sharex=True)
        
    #     ax1.plot(df[df['Cycle_No']==n][x_axis],df[df['Cycle_No']==n]['Current(A)'])
        
    #     ax1.set_ylabel('Current(A)',fontweight='bold')
        
    #     ax1.grid(linestyle='dotted')
    #     ax1.set_title('Current, Cell Voltage and Cell Temperature',fontweight='bold')
        
    
    #     ax2.plot(df[df['Cycle_No']==n][x_axis],df[df['Cycle_No']==n]['Voltage(V)'])
        
    #     ax2.set_ylabel('Voltage(V)',fontweight='bold')
        
    #     ax2.grid(linestyle='dotted')
        
    #     # ax2.set_title('Cell Voltage',fontweight='bold')
    
    
    #     ax3.plot(df[df['Cycle_No']==n][x_axis],df[df['Cycle_No']==n]['T1(℃)'])
        
    #     ax3.set_ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    #     ax3.set_xlabel('Time Elapsed (in hours)',fontweight='bold')
    #     ax3.grid(linestyle='dotted')
    
    pv=document.add_paragraph('During ')
    pv.add_run('discharge, ').bold=True
    pv.add_run('the pack goes from to . The pack voltage at rest after discharge is . At ')
    pv.add_run('charge, ').bold=True
    pv.add_run('the pack goes from to . ')
    pv.add_run('In the end, the voltage stabilizes at V during the end rest.')


    
    ct=document.add_paragraph('The maximum temperature reached is around ')
    ct.add_run('%s'%(df['T1(℃)'].max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True
    
    del data

#%%
document.add_paragraph().add_run('Tester HPPC Temperature data:').bold=True

for f1 in tester_temp_hppc_files:
    data=pd.read_excel(f1,sheet_name='record')
    img_path=f1.rsplit('\\',1)[0]
    name=f1.rsplit('\\',3)[1]+'_'+f1.rsplit('\\',3)[3].rsplit('.',1)[0]
    
    data['Date']=pd.to_datetime(data['Date'],errors='coerce')

    
    document.add_paragraph().add_run('Temperature ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print('Temperature ('+name+'):')
    plt.plot(data['Date'],data['T1(℃)'])
    
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title('Temperature ('+name+')',fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\Cell Temperature_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\Cell Temperature_'+name+'.png')
    ct=document.add_paragraph('The maximum temperature reached is around ')
    ct.add_run('%s'%(data['T1(℃)'].max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True
    
    del data

#%%
document.add_paragraph().add_run('DataLogger HPPC Temperature data:').bold=True

#%%
for f1 in dlogger_hppc_files:
    cols=['Scan','Time','101 <C1_HF_Negative Terminal side> (VDC)','Alarm 101','102 <C1_TC_Negative Terminal side> (C)','Alarm 102','103 <C1_HF_Positive Terminal side> (VDC)','Alarm 103','104 <C1_TC_Positive Terminal side> (C)','Alarm 104','105 <C1_HF_Cell Center> (VDC)','Alarm 105','106 <C1_TC_Cell Center> (C)','Alarm 106','107 <C1_HF_Cell Side> (VDC)','Alarm 107','108 <C1_TC_Cell Side> (C)','Alarm 108','109 <C1_HF_Cell Bottom> (VDC)','Alarm 109','110 <C1_TC_Cell Bottom> (C)','Alarm 110','111 <C2_TC_Positive Terminal side> (C)','Alarm 111','112 <C2_HF_Positive Terminal side> (VDC)','Alarm 112','113 <C2_HF_Negative Terminal side> (C)','Alarm 113','114 <C2_TC_Negative Terminal side> (VDC)','Alarm 114','115 <C2_HF_Cell Center> (C)','Alarm 115','116 <C2_TC_Cell Center> (VDC)','Alarm 116','117 <C2_HF_Cell Side> (C)','Alarm 117','118 <C2_TC_Cell Side> (VDC)','Alarm 118','119 <C2_HF_Cell Bottom> (C)','Alarm 119','120 <C2_TC_Cell Bottom> (VDC)','Alarm 120']
    try:
        data=pd.read_csv(f1,encoding='utf-8',sep=',',names=cols)
    except:
        data=pd.read_csv(f1,encoding='utf-16',sep=',',names=cols)
    
    data['Time']=data['Time'].astype(str)
    data[['DateTime','junk']]=data['Time'].str.rsplit(':',1,expand=True)
    
    data['DateTime']=pd.to_datetime(data['DateTime'],errors='coerce')
    data=data[['DateTime'] + [col for col in data.columns if col!='DateTime']]
    
    data=data.dropna(thresh=4)
    data=data.reset_index(drop=True)
    data=data.loc[2:,:]
    
    data['Time']=(data['DateTime']-data['DateTime'].iloc[0])/np.timedelta64(1,'h')
    x_axis='Time'
        
    img_path=f1.rsplit('\\',1)[0]
    name=f1.rsplit('\\',2)[1]+'_'+f1.rsplit('\\',2)[2].rsplit('.',1)[0]
    
    C1=[col for col in cols if "HF" in col]
    for c in C1:
        data[c]=data[c].astype(float)
        data[c]=((data[c])/(3.77))*1000
        data[c]=data[c]*1000


    title= 'C1_HF_Negative Terminal side'  
    document.add_paragraph().add_run(title+'('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data['DateTime'],data['101 <C1_HF_Negative Terminal side> (VDC)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Heat Flux (W/m^2)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')

    
    title= 'C1_TC_Negative Terminal side'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data['DateTime'],data['102 <C1_TC_Negative Terminal side> (C)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    ct=document.add_paragraph('The '+title+' has a maximum temperature value of about ')
    ct.add_run('%s'%(data['102 <C1_TC_Negative Terminal side> (C)'].astype(float).max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True  

    title='C1_HF_Positive Terminal side'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data['DateTime'],data['103 <C1_HF_Positive Terminal side> (VDC)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Heat Flux (W/m^2)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')


    title='C1_TC_Positive Terminal side'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data['DateTime'],data['104 <C1_TC_Positive Terminal side> (C)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    ct=document.add_paragraph('The '+title+' has a maximum temperature value of about ')
    ct.add_run('%s'%(data['104 <C1_TC_Positive Terminal side> (C)'].astype(float).max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True  

    title='C1_HF_Cell Center'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data['DateTime'],data['105 <C1_HF_Cell Center> (VDC)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Heat Flux (W/m^2)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')


    title='C1_TC_Cell Center'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data['DateTime'],data['106 <C1_TC_Cell Center> (C)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    ct=document.add_paragraph('The '+title+' has a maximum temperature value of about ')
    ct.add_run('%s'%(data['106 <C1_TC_Cell Center> (C)'].astype(float).max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True      

    title='C1_HF_Cell Side'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data['DateTime'],data['107 <C1_HF_Cell Side> (VDC)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Heat Flux (W/m^2)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    

    title='C1_TC_Cell Side'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data['DateTime'],data['108 <C1_TC_Cell Side> (C)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    ct=document.add_paragraph('The '+title+' has a maximum temperature value of about ')
    ct.add_run('%s'%(data['108 <C1_TC_Cell Side> (C)'].astype(float).max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True  

    title='C1_HF_Cell Bottom'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data['DateTime'],data['109 <C1_HF_Cell Bottom> (VDC)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Heat Flux (W/m^2)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    

    title='C1_TC_Cell Bottom'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data['DateTime'],data['110 <C1_TC_Cell Bottom> (C)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    ct=document.add_paragraph('The '+title+' has a maximum temperature value of about ')
    ct.add_run('%s'%(data['110 <C1_TC_Cell Bottom> (C)'].astype(float).max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True      

    title='C2_TC_Positive Terminal side'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data['DateTime'],data['111 <C2_TC_Positive Terminal side> (C)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    ct=document.add_paragraph('The '+title+' has a maximum temperature value of about ')
    ct.add_run('%s'%(data['111 <C2_TC_Positive Terminal side> (C)'].astype(float).max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True      

    title='C2_HF_Positive Terminal side'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data['DateTime'],data['112 <C2_HF_Positive Terminal side> (VDC)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Heat Flux (W/m^2)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')


    title='C2_HF_Negative Terminal side'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data['DateTime'],data['113 <C2_HF_Negative Terminal side> (C)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Heat Flux (W/m^2)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    

    title='C2_TC_Negative Terminal side'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data['DateTime'],data['114 <C2_TC_Negative Terminal side> (VDC)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    ct=document.add_paragraph('The '+title+' has a maximum temperature value of about ')
    ct.add_run('%s'%(data['114 <C2_TC_Negative Terminal side> (VDC)'].astype(float).max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True      

    title='C2_HF_Cell Center'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data['DateTime'],data['115 <C2_HF_Cell Center> (C)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Heat Flux (W/m^2)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    

    title='C2_TC_Cell Center'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data['DateTime'],data['116 <C2_TC_Cell Center> (VDC)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    ct=document.add_paragraph('The '+title+' has a maximum temperature value of about ')
    ct.add_run('%s'%(data['116 <C2_TC_Cell Center> (VDC)'].astype(float).max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True      

    title='C2_HF_Cell Side'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data['DateTime'],data['117 <C2_HF_Cell Side> (C)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Heat Flux (W/m^2)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    

    title='C2_TC_Cell Side'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data['DateTime'],data['118 <C2_TC_Cell Side> (VDC)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    ct=document.add_paragraph('The '+title+' has a maximum temperature value of about ')
    ct.add_run('%s'%(data['118 <C2_TC_Cell Side> (VDC)'].astype(float).max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True      

    title='C2_HF_Cell Bottom'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data['DateTime'],data['119 <C2_HF_Cell Bottom> (C)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Heat Flux (W/m^2)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    

    title='C2_TC_Cell Bottom'
    document.add_paragraph().add_run(title+' ('+name+'):').bold=True

    plt.figure(figsize=(l,b))
    print(title+'('+name+')')
    plt.plot(data['DateTime'],data['120 <C2_TC_Cell Bottom> (VDC)'].astype(float))
    
    plt.xlabel('Time Elapsed (in hours)',fontweight='bold')
    plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
    plt.grid(linestyle='dotted')
    plt.title(title,fontweight='bold')
    plt.tight_layout()
    plt.savefig(img_path+'\\'+title+'_'+name+'.png',dpi=1200)
    plt.close()
    
    document.add_picture(img_path+'\\'+title+'_'+name+'.png')
    ct=document.add_paragraph('The '+title+' has a maximum temperature value of about ')
    ct.add_run('%s'%(data['120 <C2_TC_Cell Bottom> (VDC)'].astype(float).max())).bold=True
    ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True      
    
    # del data
#%%

document.save('D:\\Benisha\\CALB\\AL Calb Cell Test\\AL Test\\Calb_230Ah_cell_combined_Observations.docx')



# #%%
# f=open(file.rsplit('.',1)[0]+'_raw_observations.txt',"w")

# print('======================================================================',file=f)
# print('---------------------Current Data---------------------',file=f)
# print('Datapoints:%s' % (len(voltage_data)),file=f)
# print('Total Time Spent:%s minutes' % ((voltage_data[x_axis].iloc[-1]-voltage_data[x_axis].iloc[0])/np.timedelta64(1,'m')),file=f)
# print('Maximum Charging Current:%s ; Minimum Charging Current:%s' % (details[details['State']==0]['Pack Current'].max(),details[details['State']==0]['Pack Current'].min()),file=f)
# print('Maximum Discharging Current:%s ; Minimum Discharging Current:%s' % (details[details['State']==1]['Pack Current'].max(),details[details['State']==1]['Pack Current'].min()),file=f)
# print('Maximum Charging Capacity:%s ; Minimum Charging Capacity:%s' % (details['Capacity_calculated_chg'].max(),details['Capacity_calculated_chg'].min()),file=f)
# print('Maximum Discharging Capacity:%s ; Minimum Discharging Capacity:%s' % (details['Capacity_calculated_dchg'].max(),details['Capacity_calculated_dchg'].min()),file=f)
# print('Maximum Charging Energy:%s ; Minimum Charging Energy:%s' % (details['Energy_calculated_chg'].max(),details['Energy_calculated_chg'].min()),file=f)
# print('Maximum Discharging Energy:%s ; Minimum Discharging Energy:%s' % (details['Energy_calculated_dchg'].max(),details['Energy_calculated_dchg'].min()),file=f)
# print('Maximum Charging Energy(BMS):%s ; Minimum Charging Energy(BMS):%s' % (details['Charging Energy'].max(),details['Charging Energy'].min()),file=f)
# print('Maximum Discharging Energy(BMS):%s ; Minimum Discharging Energy(BMS):%s' % (details['Discharging Energy'].max(),details['Discharging Energy'].min()),file=f)
# print('Maximum delT:%s degC' %(temperature_data['delT'].max()),file=f)
# print('Maximum Temperature:%s degC' %((temperature_data.loc[:,'MT_1':'MT_48'].max(axis=1)).max()),file=f)
# print('======================================================================',file=f)

# f.close()

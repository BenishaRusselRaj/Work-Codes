# -*- coding: utf-8 -*-

#%% Code Description

"""
USE THIS CODE WHEN:
    1. Putty log data for the monthly data of individual strings needs to be processed and report generated
    
"""

"""
PRE-REQUISITES:
    1. The putty log need to be imported into MS Excel first
    2. While opening in Excel, open it by mentioning the "comma", "space" and in "other" - ":", as the delimiters
    3. Save the file as an Excel file (i.e., with .xlsx extension)
"""

"""
NOTE:
    1. This code generally needs multiple changes for it to work without errors
    2. It is because the data collection process is highly irregular
    3. Proceed with caution at every cell, otherwise useful data may be lost.
    4. Check and double-check the data after executing every cell
"""
#%% Import the necessary libraries
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import docx


#%% Include the path to the file

file=r"D:\Benisha\IOE\24_03_04_report\03-04-14_string1_c2.xlsx"

# Define the column names 
cols=['Date','Time','CV_tag','status_V_T_tag','val1','junk1','val2','junk2','val_3','junk3','val_4','junk4','val_5','junk5','val_6','junk6','delmV','Available_Energy']

# cols=['CV_tag','status_V_T_tag','val1','junk1','val2','junk2','val_3','junk3','val_4','junk4','val_5','junk5','val_6','junk6','delmV','Available_Energy']

# Read the excel file
data=pd.read_excel(file,names=cols,index_col=False)
img_path=file.rsplit('\\',1)[0]

#%% Initialize a word document
document=docx.Document()
document.styles['Normal'].font.name='Calibri'
H1=document.add_heading()
H1.alignment=1
H1.add_run('IoE String-0x Data ('+ str('-'.join(file.split('\\',4)[3].split('_')[::-1])) +' data)').font.size=docx.shared.Pt(18)

#%% DT
data['Date']=data['Date'].str.replace('[','',regex=True)
data['Time']=data['Time'].str.replace(']','',regex=True)

data['Time']=data['Time'].astype(str)
data[['Time_format','string']]=data['Time'].str.split('.',expand=True)

data['DateTime']=data['Date'].astype(str)+' '+data['Time_format']
data['DateTime']=pd.to_datetime(data['DateTime'],format='%Y-%m-%d %H:%M:%S',errors='coerce')

data=data[['DateTime'] + [col for col in data.columns if col!='DateTime']]
data=data.drop(['Time_format','string'],axis=1)


#%%
status_data=data[data['status_V_T_tag']=='STATUS:'].dropna(how='all',axis=1)
max_min_data=data[data['status_V_T_tag']=='V:'].reset_index(drop=True)
max_min_data=pd.concat([max_min_data,data[data['status_V_T_tag']=='T:'].drop(['DateTime','Date','Time'],axis=1).reset_index(drop=True)],axis=1).dropna(how='all',axis=1) # try to concatenate in same file
details=data[data['status_V_T_tag']=='voltage:'].dropna(how='all',axis=1)
soc_details = data[data['CV_tag']=='OCV'].dropna(how = 'all', axis = 1)

#%% columns names
v_names=[f'MV_{x}' for x in range(1,193)] 
t_names=[f'MT_{x}' for x in range(1,129)]

# Column names for the status dataframe are defined here
status=['DateTime','Date','Time','junk','junk','Pack Status','junk','Shutdown','Normal','junk','junk','Precharge Relay','junk','junk','Main Relay'] #,'junk','junk'

# Another column that might be there in detail_names dataframe is: 'Pack Available Energy'
detail_names=['DateTime','Date','Time','junk','junk','Pack Voltage','junk','junk','Pack Current','junk','SoH','junk','junk','Charging Energy','junk','junk','Discharging Energy','junk','junk'] 
max_min_names=['DateTime','Date','Time','junk','junk','Min_V_BMS','','junk','junk','Max_V_BMS','','junk','junk','Average_CellV_BMS','junk','junk','junk','delV_BMS(mV)','junk','junk','Min_T_BMS','','junk','junk','Max_T_BMS','']
soc_names = ['DateTime','Date','Time', 'junk', 'junk', 'OCV_SoC', 'junk', 'junk', 'Pack_SoC', 'junk', 'junk',
             'Usable_SoC', 'junk', 'junk', 'Weighted_SoC']

# status=['junk','junk','Pack Status','junk','Shutdown','Normal','junk','junk','Precharge Relay','junk','junk','Main Relay','junk'] #,'junk','junk'
# # detail_names=['junk','junk','Pack Voltage','junk','junk','Pack Current','junk','SoH','junk','junk','Charging Energy','junk','junk','Discharging Energy','junk','junk','junk','Pack Available Energy']
# detail_names=['junk','junk','Pack Voltage','junk','junk','Pack Current','junk','SoH','junk','junk','Charging Energy','junk','junk','junk','Discharging Energy','junk']
# max_min_names=['junk','junk','Min_V_BMS','','junk','junk','Max_V_BMS','','junk','junk','Average_CellV_BMS','junk','junk','junk','delV_BMS(mV)','junk','junk','Min_T_BMS','','junk','junk','Max_T_BMS','']

# Assigning proper column names to the corresponding dataframes
status_data.columns=status
details.columns=detail_names
max_min_data.columns=max_min_names
soc_details.columns = soc_names

#%%
# status_data['Main Relay']=status_data['Main Relay'].astype(str).str.extract(r'([A-Z][a-z]{4,})') #[Open][close][open] 

#%% Drop junk columns from the dataframes

status_data=status_data.drop([col for col in status_data.columns if "junk" in col], axis=1)
details=details.drop([col for col in details.columns if "junk" in col], axis=1)
max_min_data=max_min_data.drop([col for col in max_min_data.columns if "junk" in col], axis=1)

#%% Extract energy data in proper format
details['Charging Energy']=details['Charging Energy'].astype(str).str.extract(r'(\d{1,}.\d{1,})')
details['Discharging Energy']=details['Discharging Energy'].astype(str).str.extract(r'(\d{1,}.\d{1,})')

#%% Extract data in proper format
details['Pack Current']=details['Pack Current'].str.replace('A','',regex=True)
status_data['Precharge Relay']=status_data['Precharge Relay'].str.replace(',','',regex=True)
details['Charging Energy']=details['Charging Energy'].str.replace('engy:','',regex=True)
details['Discharging Energy']=details['Discharging Energy'].str.replace('engy:','',regex=True)
details['Charging Energy']=details['Charging Energy'].astype(float)
details['Discharging Energy']=details['Discharging Energy'].astype(float)

# status_data['Main Relay']=status_data['Main Relay'].astype(str).str.extract(r'([A-Z]{2,5})')
# status_data['Precharge Relay']=status_data['Precharge Relay'].astype(str).str.extract(r'([A-Z]{2,5})')

status_data['Main Relay']=status_data['Main Relay'].astype(str).str.extract(r'([A-Z][a-z]{2,5})')
status_data['Precharge Relay']=status_data['Precharge Relay'].astype(str).str.extract(r'([A-Z][a-z]{2,5})')

#%%
# status_data['Precharge Relay']=status_data['Precharge Relay'].str.replace('R',np.nan,regex=True)

#%% Data pre-processing
status_data['Main Relay']=status_data['Main Relay'].replace('Relay',np.nan)
status_data['Main Relay']=status_data['Main Relay'].fillna(method='ffill')
status_data['Main Relay']=status_data['Main Relay'].fillna(method='bfill')
status_data['Main Relay']=status_data['Main Relay'].astype(str)

status_data['Precharge Relay']=status_data['Precharge Relay'].fillna(method='ffill')
status_data['Precharge Relay']=status_data['Precharge Relay'].fillna(method='bfill')
status_data['Precharge Relay']=status_data['Precharge Relay'].astype(str)

#%% Filter cell voltage data from the log
voltage_data=data[data['CV_tag']=='MV1'].drop(['delmV','Available_Energy'],axis=1).reset_index(drop=True)

for i in range(2,33):
    voltage_data=pd.concat([voltage_data,data[data['CV_tag']=='MV'+str(i)].drop(['DateTime','Date','Time','delmV','Available_Energy'],axis=1).reset_index(drop=True)],axis=1)
    # If no datetime columns are there, use the following command instead of the top one:
    # voltage_data=pd.concat([voltage_data,data[data['CV_tag']=='MV'+str(i)].drop(['delmV','Available_Energy'],axis=1).reset_index(drop=True)],axis=1)

# Remove junk columns in the voltage dataframe
voltage_data=voltage_data.drop([col for col in voltage_data.columns if "junk" in col], axis=1)
voltage_data=voltage_data.drop([col for col in voltage_data.columns if "tag" in col], axis=1)

#%% Use this block if datetime is present in the original log
for i,n in enumerate(['DateTime','Date','Time']):
    v_names.insert(i,n)
    t_names.insert(i,n)

voltage_data.columns=v_names

#%% Filter the temperature data from the log
temperature_data=data[data['CV_tag']=='MT1'].loc[:,:'junk2'].reset_index(drop=True)

#%% Clean temperature data entries
for i in range(2,33):
    temperature_data=pd.concat([temperature_data,data[data['CV_tag']=='MT'+str(i)].loc[:,:'junk2'].drop(['DateTime','Date','Time'],axis=1).reset_index(drop=True)],axis=1)
    # temperature_data=pd.concat([temperature_data,data[data['CV_tag']=='MT'+str(i)].loc[:,:'junk2'].reset_index(drop=True)],axis=1)
    
temperature_data=temperature_data.dropna(how='all',axis=1)
temperature_data=temperature_data.drop([col for col in temperature_data.columns if "tag" in col], axis=1)

#%% Filter the temperature data entries
temperature_data.columns=t_names

for i in range(1,129):
    temperature_data['MT_'+str(i)]=temperature_data['MT_'+str(i)].astype(str).str.extract(r'(\d{2}.\d{5})')

#%% Clean the cell voltage data entries
voltage_data.loc[:,'MV_1':'MV_192']=voltage_data.loc[:,'MV_1':'MV_192'].replace('_0',np.nan)
voltage_data.loc[:,'MV_1':'MV_192']=voltage_data.loc[:,'MV_1':'MV_192'].fillna(method='ffill')

for i in range(1,193):
    voltage_data['MV_'+str(i)]=voltage_data['MV_'+str(i)].astype(str).str.extract(r'(\d{5})')

voltage_data.loc[:,'MV_1':'MV_192']=voltage_data.loc[:,'MV_1':'MV_192'].fillna(method='ffill')
voltage_data.loc[:,'MV_1':'MV_192']=voltage_data.loc[:,'MV_1':'MV_192'].astype(int)
voltage_data['Mean_V']=voltage_data.loc[:,'MV_1':'MV_192'].mean(axis=1)

#%% Temperature data is filtered from the complete dataset
temperature_data.loc[:,'MT_1':'MT_128']=temperature_data.loc[:,'MT_1':'MT_128'].astype(float)
temperature_data.loc[:,'MT_1':'MT_128']=temperature_data.loc[:,'MT_1':'MT_128'].fillna(method='ffill')
temperature_data['Mean_T']=temperature_data.loc[:,'MT_1':'MT_128'].mean(axis=1)

temperature_data['delT']=temperature_data.loc[:,'MT_1':'MT_128'].max(axis=1)-temperature_data.loc[:,'MT_1':'MT_128'].min(axis=1)

#%% Calculating Capacity_calculated and energy

# Time difference for every datapoints, to perform coulomb counting
details['Time_in_sec_s']=(details['DateTime']-details['DateTime'].shift(1))/np.timedelta64(1,'s')# time difference b/w every datapt
details['Time_in_sec']=(details['DateTime']-details['DateTime'].iloc[0])/np.timedelta64(1,'s')

# Data entries are cleaned and data in proper format is retrieved
details['Pack Current']=details['Pack Current'].astype(str).str.extract(r'(\-{0,}\d{1,}.\d{1,})')
details['Pack Voltage']=details['Pack Voltage'].astype(str).str.extract(r'(\d{1,}.\d{1,})')
details['Pack Current']=details['Pack Current'].astype(float)
details['Pack Voltage']=details['Pack Voltage'].astype(float)

#%% The pack state is set as 0 for charging and 1 for discharging
details['State']=np.nan
details['State']=np.where(details['Pack Current']>0,0,details['State'])
details['State']=np.where(details['Pack Current']<0,1,details['State'])

# The time difference is filtered to set 0 for any time difference greater than 300 seconds.
# This is done so as to remove wrong calculations due to the time spent by the pack in rest.
details['Time_in_sec_s_cap']=np.where((details['Time_in_sec_s']>300),np.nan,details['Time_in_sec_s'])


#%% 'details' dataframe is processed to include the capacity and energy values
# These values are calculated using the 'Coulomb-counting method', which makes use of time and current
details['Cap_inst']=details['Time_in_sec_s_cap']*abs(details['Pack Current'])/3600
details['Capacity_calculated']=details['Cap_inst'].groupby(details['State']).cumsum()
details['Capacity_calculated_chg']=(details[details['State']==0]['Cap_inst']).cumsum()
details['Capacity_calculated_dchg']=(details[details['State']==1]['Cap_inst']).cumsum()

details['Capacity_calculated_chg']=details['Capacity_calculated_chg'].fillna(method='bfill')
details['Capacity_calculated_dchg']=details['Capacity_calculated_dchg'].fillna(method='bfill')

details['Energy_calculated']=details['Capacity_calculated']*details['Pack Voltage']
details['Energy_calculated_chg']=details['Capacity_calculated_chg']*details['Pack Voltage']
details['Energy_calculated_dchg']=details['Capacity_calculated_dchg']*details['Pack Voltage']

#%% calculate delV from the cell voltage values
voltage_data['delV']=voltage_data.loc[:,'MV_1':'MV_192'].max(axis=1)-voltage_data.loc[:,'MV_1':'MV_192'].min(axis=1)

#%% Adding datapoints to check the frequency of data collection
for x in [voltage_data, temperature_data, details, status_data, max_min_data]:
    x['Datapoints'] = range(1,len(x)+1)
#%% Adding table in document 
document.add_paragraph().add_run('\nTest Specifications:').bold=True

table1=document.add_table(rows=4,cols=2)
table1.cell(0,0).text='Test Date and Time'
table1.cell(0,1).text=str(status_data['DateTime'].iloc[0])+' to '+str(status_data['DateTime'].iloc[-1])+' ('+ str(status_data['DateTime'].iloc[-1]-status_data['DateTime'].iloc[0]) +' )'
table1.cell(1,0).text='Discharging Time Period'
table1.cell(2,0).text='Charging Time Period'
table1.cell(3,0).text='C-rate'

table1.style='Table Grid'

# Set the dimensions of the plots in the word document; selected through trial-and-error
l=7
b=4

# Set the x-axis column name
x_axis='DateTime'

#%% Plot Precharge relay 
document.add_paragraph().add_run('Precharge Relay:').bold=True

plt.figure(figsize=(l,b))
plt.plot(status_data[x_axis],status_data['Precharge Relay'])

plt.xlabel(x_axis,fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Precharge Relay',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Precharge Relay.png',dpi=1200)

document.add_picture(img_path+'\\Precharge Relay.png')
document.add_paragraph().add_run('**Note: The “DateTime” axis in the graphs is read as DD/MM/YYYY h:m:s (i.e., // ::, // :: and so on. ').underline=True

#%% Plot Main Relay
document.add_paragraph().add_run('Main Relay:').bold=True
plt.figure(figsize=(l,b))
plt.plot(status_data[x_axis],status_data['Main Relay'])

plt.xlabel(x_axis,fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Main Relay',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Main Relay.png',dpi=1200)

document.add_picture(img_path+'\\Main Relay.png')

#%% Plot Pack Status
document.add_paragraph().add_run('Pack Status:').bold=True
plt.figure(figsize=(l,b))
plt.plot(status_data[x_axis],status_data['Pack Status'])
plt.xlabel(x_axis,fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('Pack Status',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Pack Status.png',dpi=1200)

document.add_picture(img_path+'\\Pack Status.png')


#%% Plot Pack Current
document.add_paragraph().add_run('Pack Current:').bold=True
plt.figure(figsize=(l,b))
plt.plot(details[x_axis],details['Pack Current'])

plt.xlabel(x_axis,fontweight='bold')
plt.ylabel('Current(A)',fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('String Current',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Pack Current.png',dpi=1200)

document.add_picture(img_path+'\\Pack Current.png')

#%% Plot Calculated dicharging Pack Capacity
document.add_paragraph().add_run('Pack Capacity (calculated):').bold=True

plt.figure(figsize=(l,b))
plt.plot(details[x_axis],details['Capacity_calculated_dchg'])

plt.xlabel(x_axis,fontweight='bold')
plt.ylabel('Capacity(Ah)',fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('Pack Discharging Capacity',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Pack Discharging Capacity_calculated.png',dpi=1200)

document.add_picture(img_path+'\\Pack Discharging Capacity_calculated.png')
document.add_paragraph('In this discharging session, the capacity expended by the pack is about %.3fAh ().' %details['Capacity_calculated_dchg'].max())

#%% Plot calculated charging pack capacity
plt.figure(figsize=(l,b))
plt.plot(details[x_axis],details['Capacity_calculated_chg'])

plt.xlabel(x_axis,fontweight='bold')
plt.ylabel('Capacity(Ah)',fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('Pack Charging Capacity',fontweight='bold')
plt.tight_layout()

plt.savefig(img_path+'\\Pack Charging Capacity_calculated.png',dpi=1200)

document.add_picture(img_path+'\\Pack Charging Capacity_calculated.png')
document.add_paragraph('In this charging session, the capacity transferred to the pack is about %.3fAh (). ' %details['Capacity_calculated_chg'].max())
document.add_paragraph().add_run('Note: These values were calculated from the current (in amperes) values and time (in hours).').underline=True

#%% Plot calculated discharging pack energy
document.add_paragraph().add_run('Pack Energy (calculated):').bold=True
plt.figure(figsize=(l,b))
plt.plot(details[x_axis],details['Energy_calculated_dchg']*0.001)

plt.xlabel(x_axis,fontweight='bold')
plt.ylabel('Energy(kWh)',fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('Pack Discharging Energy',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Pack Discharging Energy_calculated.png',dpi=1200)

document.add_picture(img_path+'\\Pack Discharging Energy_calculated.png')
document.add_paragraph('The Pack Energy expended in this discharging session (approx.  minutes) is about %.3fkWh.' % (details['Energy_calculated_dchg'].max()*0.001))

#%% Plot calculated charging energy
plt.figure(figsize=(l,b))
plt.plot(details[x_axis],details['Energy_calculated_chg']*0.001)

plt.xlabel(x_axis,fontweight='bold')
plt.ylabel('Energy(kWh)',fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('Pack Charging Energy',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Pack Charging Energy_calculated.png',dpi=1200)

document.add_picture(img_path+'\\Pack Charging Energy_calculated.png')
document.add_paragraph('The Pack Energy expended in this charging session () is about %.3fkWh.' % (details['Energy_calculated_chg'].max()*0.001))
document.add_paragraph().add_run('Note: These values were calculated from the calculated capacity (in Ah) values and pack voltage (in volts).').underline=True


#%% Plot BMS Charging energy
document.add_paragraph().add_run('Pack Energy (BMS):').bold=True

plt.figure(figsize=(l,b))
plt.plot(details[x_axis],details['Discharging Energy'])

plt.xlabel(x_axis,fontweight='bold')
plt.ylabel('Energy(kWh)',fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('String Discharging Energy',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Pack Discharging Energy_BMS.png',dpi=1200)

document.add_picture(img_path+'\\Pack Discharging Energy_BMS.png')
p_bms_d_engy=document.add_paragraph('The ')
p_bms_d_engy.add_run('discharging pack energy ').bold=True
p_bms_d_engy.add_run('is about ')
p_bms_d_engy.add_run('%.3fWh.' %details['Discharging Energy'].max()).bold=True

#%% Plot Charging energy
plt.figure(figsize=(l,b))
plt.plot(details[x_axis],details['Charging Energy'])

plt.xlabel(x_axis,fontweight='bold')
plt.ylabel('Energy(kWh)',fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('String Charging Energy',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Pack Charging Energy_BMS.png',dpi=1200)

document.add_picture(img_path+'\\Pack Charging Energy_BMS.png')
p_bms_c_engy=document.add_paragraph('The ')
p_bms_c_engy.add_run('charging pack energy ').bold=True
p_bms_c_engy.add_run('is about ')
p_bms_c_engy.add_run('%.3fWh.' %details['Charging Energy'].max()).bold=True

#%% Plot Pack Voltage
document.add_paragraph().add_run('Pack Voltage:').bold=True
plt.figure(figsize=(l,b))
plt.plot(details[x_axis],details['Pack Voltage'])

plt.xlabel(x_axis,fontweight='bold')
plt.ylabel('Voltage(V)',fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('String Voltage',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Pack Voltage.png',dpi=1200)

document.add_picture(img_path+'\\Pack Voltage.png')
pv=document.add_paragraph('During ')
pv.add_run('discharge, ').bold=True
pv.add_run('the pack goes from to . The pack voltage at rest after discharge is . At ')
pv.add_run('charge, ').bold=True
pv.add_run('the pack goes from to . ')
pv.add_run('In the end, the voltage stabilizes at %sV during the end rest.' %details['Pack Voltage'].iloc[-1])

#%% Plot Cell Voltage
document.add_paragraph().add_run('Cell Voltage:').bold=True

plt.figure(figsize=(l,b))
plt.plot(voltage_data[x_axis],voltage_data.loc[:,'MV_1':'MV_192']*0.0001,marker='o',markersize=3) #

plt.ylabel('Voltage(V)',fontweight='bold')
plt.xlabel(x_axis,fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Cell Voltage',fontweight='bold')
plt.tight_layout()

#%% Plot Average Voltage
document.add_paragraph().add_run('Average Voltage:').bold=True
plt.figure(figsize=(l,b))
plt.plot(voltage_data[x_axis],(voltage_data['Mean_V'])*0.0001)

plt.ylabel('Voltage(V)',fontweight='bold')
plt.xlabel(x_axis,fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Average Voltage',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Average Voltage.png',dpi=1200)

document.add_picture(img_path+'\\Average Voltage.png')
av=document.add_paragraph('The average cell voltage during ')
av.add_run('discharge ').bold=True
av.add_run('goes from ')
av.add_run('V to V. ').bold=True
av.add_run('The average cell voltage at the rest after discharge is V. At ')
av.add_run('charge').bold=True
av.add_run(', the average cell voltage goes from ')
av.add_run('V to V.').bold=True
av.add_run('In the end rest, it stabilizes at %.3fV.' % (voltage_data['Mean_V'].iloc[-1]*0.0001))

#%% Plot delV 
document.add_paragraph().add_run('Voltage Difference (delV):').bold=True
plt.figure(figsize=(l,b))
plt.plot(voltage_data[x_axis],(voltage_data['delV'])*0.1)

plt.ylabel('Voltage(mV)',fontweight='bold')
plt.xlabel(x_axis,fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Voltage difference (delV)',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\delV.png',dpi=1200)

document.add_picture(img_path+'\\delV.png')
dv=document.add_paragraph('During discharge, the ')
dv.add_run('maximum delV (difference between the highest and lowest cell voltage), ').bold=True
dv.add_run('is about ')
dv.add_run('mV, at the end of discharge. ').bold=True
dv.add_run('At the rest after discharging, the delmV is mV. In the charging session, the maximum delV is mV. ')
dv.add_run('The pack stabilizes at %.3fmV delV at the end rest. '% (voltage_data['delV'].iloc[-1]*0.1))

#%% Plot Cell Temperature
document.add_paragraph().add_run('Cell Temperature:').bold=True

plt.figure(figsize=(l,b))
plt.plot(temperature_data[x_axis],temperature_data.loc[:,'MT_1':'MT_128'])

plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
plt.xlabel(x_axis,fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Cell Temperature',fontweight='bold')
ct=document.add_paragraph('The maximum temperature reached is around ')
ct.add_run('%.3f'%(temperature_data.loc[:,'MT_1':'MT_48'].max(axis=1).max())).bold=True
ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True

#%% Plot Average Temperature
document.add_paragraph().add_run('Average temperature:').bold=True
plt.figure(figsize=(l,b))
plt.plot(temperature_data[x_axis],temperature_data['Mean_T'])

plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
plt.xlabel(x_axis,fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Average Temperature',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Average Temperature.png',dpi=1200)

document.add_picture(img_path+'\\Average Temperature.png')
at=document.add_paragraph('As can be seen, the ')
at.add_run('average cell temperature ').bold=True
at.add_run('stays between ')
at.add_run(u'\N{DEGREE SIGN}'+'C to '+u'\N{DEGREE SIGN}'+'C ').bold=True
at.add_run('during ')
at.add_run('discharging').bold=True
at.add_run(' and ')
at.add_run(u'\N{DEGREE SIGN}'+'C to '+u'\N{DEGREE SIGN}'+'C ').bold=True
at.add_run('during ')
at.add_run('charging.').bold=True

#%% Plot delT
document.add_paragraph().add_run('Temperature Difference (delT):').bold=True
plt.figure(figsize=(l,b))
plt.plot(temperature_data[x_axis],temperature_data['delT'])

plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
plt.xlabel(x_axis,fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Temperature difference (delT)',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\delT.png',dpi=1200)

document.add_picture(img_path+'\\delT.png')
dt=document.add_paragraph('The ')
dt.add_run('maximum delT ').bold=True
dt.add_run('is about ')
dt.add_run('%.3f' %temperature_data['delT'].max()).bold=True
dt.add_run(u'\N{DEGREE SIGN}'+'C.').bold=True

#%% Plot OCV SoC 
document.add_paragraph().add_run('OCV SoC:').bold=True
plt.figure(figsize=(l,b))
plt.plot(soc_details[x_axis], soc_details['OCV_SoC']*100)
plt.xlabel(x_axis,fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('State of Charge',fontweight='bold')
plt.ylabel('SoC (%)',fontweight='bold')
plt.xlabel(x_axis,fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\OCV_SoC.png',dpi=1200)

document.add_picture(img_path+'\\OCV_SoC.png')

#%% Plot Pack SoC
document.add_paragraph().add_run('Pack SoC:').bold=True
plt.figure(figsize=(l,b))
plt.plot(soc_details[x_axis], soc_details['Pack_SoC']*100)
plt.xlabel(x_axis,fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('State of Charge',fontweight='bold')
plt.ylabel('SoC (%)',fontweight='bold')
plt.xlabel(x_axis,fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Pack_SoC.png',dpi=1200)

document.add_picture(img_path+'\\Pack_SoC.png')

#%% Plot usable SoC
document.add_paragraph().add_run('Usable SoC:').bold=True
plt.figure(figsize=(l,b))
plt.plot(soc_details[x_axis], soc_details['Usable_SoC']*100)
plt.xlabel(x_axis,fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('State of Charge',fontweight='bold')
plt.ylabel('SoC (%)',fontweight='bold')
plt.xlabel(x_axis,fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Usable_SoC.png',dpi=1200)

document.add_picture(img_path+'\\Usable_SoC.png')

#%% Plot Weighted SoC
document.add_paragraph().add_run('Weighted SoC:').bold=True
plt.figure(figsize=(l,b))
plt.plot(soc_details[x_axis], soc_details['Weighted_SoC'] * 100)
plt.xlabel(x_axis,fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('State of Charge',fontweight='bold')
plt.ylabel('SoC (%)',fontweight='bold')
plt.xlabel(x_axis,fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Weighted_SoC.png',dpi=1200)

document.add_picture(img_path+'\\Weighted_SoC.png')

#%% Change the units from mV to V
voltage_data.loc[:,'MV_1':'MV_192']=voltage_data.loc[:,'MV_1':'MV_192']*0.0001

#%% 
path=file.rsplit('.',1)[0]+'_modified.xlsx'
document.save(file.rsplit('.',1)[0]+'_observations.docx')

#%%
# "writer" opens the excelwriter module of pandas library
writer = pd.ExcelWriter(path, engine = 'xlsxwriter')
voltage_data.to_excel(writer, sheet_name = 'Cell Voltage',index=False)
temperature_data.to_excel(writer, sheet_name = 'Cell Temperature',index=False)
details.to_excel(writer,sheet_name='Pack Details',index=False)
status_data.to_excel(writer,sheet_name='Pack Status',index=False)
max_min_data.to_excel(writer,sheet_name='Max Min Mean',index=False)
writer.close()

#%% Save text file containing observations
f=open(file.rsplit('.',1)[0]+'_observations.txt',"w")

print('======================================================================',file=f)
print('---------------------Current Data---------------------',file=f)
print('Datapoints:%s' % (len(voltage_data)),file=f)
print('Total Time Spent:%s minutes' % ((voltage_data['DateTime'].iloc[-1]-voltage_data['DateTime'].iloc[0])/np.timedelta64(1,'m')),file=f)
print('Maximum Charging Current:%s ; Minimum Charging Current:%s' % (details[details['State']==0]['Pack Current'].max(),details[details['State']==0]['Pack Current'].min()),file=f)
print('Maximum Discharging Current:%s ; Minimum Discharging Current:%s' % (details[details['State']==1]['Pack Current'].max(),details[details['State']==1]['Pack Current'].min()),file=f)
print('Maximum Charging Capacity:%s ; Minimum Charging Capacity:%s' % (details['Capacity_calculated_chg'].max(),details['Capacity_calculated_chg'].min()),file=f)
print('Maximum Discharging Capacity:%s ; Minimum Discharging Capacity:%s' % (details['Capacity_calculated_dchg'].max(),details['Capacity_calculated_dchg'].min()),file=f)
print('Maximum Charging Energy:%s ; Minimum Charging Energy:%s' % (details['Energy_calculated_chg'].max(),details['Energy_calculated_chg'].min()),file=f)
print('Maximum Discharging Energy:%s ; Minimum Discharging Energy:%s' % (details['Energy_calculated_dchg'].max(),details['Energy_calculated_dchg'].min()),file=f)
print('Maximum Charging Energy(BMS):%s ; Minimum Charging Energy(BMS):%s' % (details['Charging Energy'].max(),details['Charging Energy'].min()),file=f)
print('Maximum Discharging Energy(BMS):%s ; Minimum Discharging Energy(BMS):%s' % (details['Discharging Energy'].max(),details['Discharging Energy'].min()),file=f)
print('Maximum delT:%s degC' %(temperature_data['delT'].max()),file=f)
print('Maximum Temperature:%s degC' %((temperature_data.loc[:,'MT_1':'MT_48'].max(axis=1)).max()),file=f)
print('======================================================================',file=f)

f.close()

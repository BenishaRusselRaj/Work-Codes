# -*- coding: utf-8 -*-

#%% Code Description

"""
USE THIS CODE WHEN:
    1. Putty log data needs to be processed and report generated
 
"""

"""
PRE-REQUISITES:
    1. The putty log need to be imported into MS Excel first
    2. While opening in Excel, open it by mentioning "comma", "space" and in "other" - ":", as the delimiters
    3. Save the file as an Excel file (i.e., with .xlsx extension)
"""

"""
NOTE:
    1. This code generally needs multiple changes for it to work without errors
    2. It is because the data collection process is highly irregular
    3. Proceed with caution at every cell, otherwise useful data may be lost.
    4. Check and double-check the data after executing every cell
"""

#%% Import necessary libraries
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import docx

#%% Column names
data=pd.DataFrame()
cols=['Date','Time','MV_MT_tag','PV_V_T_status_tag','val1','val_t2','val2','val_t3','val_3','junk','val_4','junk','val_5','junk','val_6','junk','val_7','junk','val_8','junk','val_9','junk','val_10','junk','val_11','junk','val_12','junk','val_13','junk','val_14','junk']

#%% Read excel file using pandas
tag = '_1'
file=r"D:\Benisha\15kW_LTO_Pack\25_03_10\Log_1\log 10-03-2025.xlsx"

data=pd.read_excel(file,names=cols,index_col=False)
img_path=file.rsplit('\\',1)[0]

#%% Create a .docx document
document=docx.Document()
document.styles['Normal'].font.name='Calibri'
H1=document.add_heading()
H1.alignment=1
H1.add_run('LTO 15kW Pack Data ('+ str('-'.join(file.split('\\',5)[4].split('_')[::-1])) +' data)').font.size=docx.shared.Pt(18)
document.add_paragraph().add_run('\n\nPack Specifications:').bold=True

# Add a table containing the pack details
table1=document.add_table(rows=6,cols=2)

table1.cell(0,0).text='Chemistry'
table1.cell(0,1).text='LTO'
table1.cell(1,0).text='Cell Type'
table1.cell(1,1).text='Cylindrical'
table1.cell(2,0).text='Capacity'
table1.cell(2,1).text='40Ah'
table1.cell(3,0).text='Configuration'
table1.cell(3,1).text='1P168S'
table1.cell(4,0).text='UV cut-off'
table1.cell(4,1).text='1.85V'
table1.cell(5,0).text='OV cut-off'
table1.cell(5,1).text='2.6V'

table1.style='Table Grid'

#%% Clean the date and time entries in the log and create a datetime column

data['Date']=data['Date'].str.replace('[','',regex=True)
data['Time']=data['Time'].str.replace(']','',regex=True)

data['Time']=data['Time'].astype(str)
data[['Time_format','string']]=data['Time'].str.split('.',expand=True)

data['DateTime']=data['Date'].astype(str)+' '+data['Time_format']
data['DateTime']=pd.to_datetime(data['DateTime'],format='%Y-%m-%d %H:%M:%S',errors='coerce')

data=data[['DateTime'] + [col for col in data.columns if col!='DateTime']]
data=data.drop(['Time_format','string'],axis=1)
data=data.sort_values(['DateTime'])

#%% Filter the different data values based on tags in the input log
status_data=data[data['PV_V_T_status_tag']=='STATUS:'].dropna(how='all',axis=1)
max_min_data=data[data['PV_V_T_status_tag']=='V:'].dropna(how='all',axis=1).reset_index(drop=True)
max_min_data=pd.concat([max_min_data,data[data['PV_V_T_status_tag']=='T:'].drop(['DateTime','Date','Time'],axis=1).reset_index(drop=True)],axis=1).dropna(how='all',axis=1) # try to concatenate in same file
details=data[data['PV_V_T_status_tag']=='voltage:'].dropna(how='all',axis=1)

#%% The column names for the voltage and temperature data is alloted iteratively in a for loop
v_names=[f'MV_{x}' for x in range(1,169)] 
t_names=[f'MT_{x}' for x in range(1,49)]


cellVColumns = [f'MV_{x}' for x in range(1,169)]
cellTColumns = [f'MT_{x}' for x in range(1,49)]

# Set the column names for the dataframe named "status", which has the details regarding the status of various relays
status=['DateTime','Date','Time','junk','junk','Pack Status','junk','Shutdown','Normal','junk','junk','Precharge Relay','junk','junk','Main Relay']
status_data.columns=status

# details file has the pack voltage, pack current, charging and discharging energy
detail_names=['DateTime','Date','Time','junk','junk','Pack Voltage','junk','junk','Pre Command','junk','junk','Pack Current','junk','SoH','junk','junk','Charging Energy','junk','junk','junk','Discharging Energy','junk']
details.columns=detail_names

# The max_min_names dataframe has maximum and minimum voltage and temperature data
max_min_names=['DateTime','Date','Time','junk','junk','Min_V_BMS','','junk','junk','Max_V_BMS','','junk','junk','Average_CellV_BMS','junk','junk','junk','delV_BMS(mV)','junk','junk','Min_T_BMS','','junk','junk','Max_T_BMS','']
max_min_data.columns=max_min_names

# The data entries are cleaned into their proper format
status_data['Main Relay']=status_data['Main Relay'].astype(str).str.extract(r'([A-Z][a-z]{3,})')
status_data['Precharge Relay']=status_data['Precharge Relay'].astype(str).str.extract(r'([A-Z][a-z]{3,})')
status_data['Precharge Relay']=status_data['Precharge Relay'].fillna(method='ffill')

# Drop junk columns
status_data=status_data.drop([col for col in status_data.columns if "junk" in col], axis=1)
details=details.drop([col for col in details.columns if "junk" in col], axis=1)
max_min_data=max_min_data.drop([col for col in max_min_data.columns if "junk" in col], axis=1)

#%%
details['Pack Current']=details['Pack Current'].str.replace('A','',regex=True)
details['Pack Voltage']=details['Pack Voltage'].astype(int)

# Clean data entries to their proper format
status_data['Precharge Relay']=status_data['Precharge Relay'].str.replace(',','',regex=True)
details['Charging Energy']=details['Charging Energy'].str.replace('engy:','',regex=True)
details['Discharging Energy']=details['Discharging Energy'].str.replace('engy:','',regex=True)
details['Charging Energy']=details['Charging Energy'].astype(float)
details['Discharging Energy']=details['Discharging Energy'].astype(float)

#%% Filter the voltage data based on "MV1" tag
voltage_data=data[data['MV_MT_tag']=='MV1'].reset_index(drop=True)
for i in range(2,13):
    voltage_data=pd.concat([voltage_data,data[data['MV_MT_tag']=='MV'+str(i)].drop(['DateTime','Date','Time'],axis=1).reset_index(drop=True)],axis=1)

# Drop junk columns from the voltage data
voltage_data=voltage_data.drop([col for col in voltage_data.columns if "junk" in col], axis=1)
voltage_data=voltage_data.drop([col for col in voltage_data.columns if "tag" in col], axis=1)
voltage_data=voltage_data.drop(['val_t2','val_t3'],axis=1)

#%% Add datetime related column names are added
for i,n in enumerate(['DateTime','Date','Time']):
    v_names.insert(i,n)
    t_names.insert(i,n)

voltage_data.columns=v_names

#%% The temperature data is filtered using the "MT1" tag
temperature_data=data[data['MV_MT_tag']=='MT1'].loc[:,:'val_t3'].reset_index(drop=True)

#%%
for i in range(2,13):
    temperature_data=pd.concat([temperature_data,data[data['MV_MT_tag']=='MT'+str(i)].loc[:,:'val_t3'].drop(['DateTime','Date','Time'],axis=1).reset_index(drop=True)],axis=1)

# Drop junk columns
temperature_data=temperature_data.dropna(how='all',axis=1)
temperature_data=temperature_data.drop([col for col in temperature_data.columns if "tag" in col], axis=1)

# Set proper column names to the temperature dataframe
temperature_data.columns=t_names

# Filter the temperature data to the proper format
for i in range(1,49):
    temperature_data['MT_'+str(i)]=temperature_data['MT_'+str(i)].astype(str).str.extract(r'(\d{2}.\d{5})')

#%% Clean the voltage dataframe to the proper format
voltage_data.loc[:,'MV_1':'MV_168']=voltage_data.loc[:,'MV_1':'MV_168'].replace('_0',np.nan)
voltage_data.loc[:,'MV_1':'MV_168']=voltage_data.loc[:,'MV_1':'MV_168'].fillna(method='ffill')

for i in range(1,169):
    voltage_data['MV_'+str(i)]=voltage_data['MV_'+str(i)].astype(str).str.extract(r'(\d{5})')

voltage_data.loc[:,'MV_1':'MV_168']=voltage_data.loc[:,'MV_1':'MV_168'].fillna(method='ffill')
voltage_data.loc[:,'MV_1':'MV_168']=voltage_data.loc[:,'MV_1':'MV_168'].fillna(method='bfill')

voltage_data.loc[:,'MV_1':'MV_168']=voltage_data.loc[:,'MV_1':'MV_168'].astype(int)
voltage_data.loc[:,'MV_1':'MV_168']=voltage_data.loc[:,'MV_1':'MV_168']*0.0001

#%% Filter voltage data to remove junk/impossible voltage values
voltage_data.loc[:,'MV_1':'MV_168'] = np.where(voltage_data.loc[:,'MV_1':'MV_168']>5, np.nan, voltage_data.loc[:,'MV_1':'MV_168'])

#%% Calculate the mean voltage and delta voltage
# delta voltage is calculated in mV
voltage_data['Mean_V']=voltage_data.loc[:,'MV_1':'MV_168'].mean(axis=1)
voltage_data['delV']=(voltage_data.loc[:,'MV_1':'MV_168'].max(axis=1)-voltage_data.loc[:,'MV_1':'MV_168'].min(axis=1))*1000

#%% Pre-process temperature data 
temperature_data.loc[:,'MT_1':'MT_48']=temperature_data.loc[:,'MT_1':'MT_48'].astype(float)
temperature_data.loc[:,'MT_1':'MT_48']=temperature_data.loc[:,'MT_1':'MT_48'].fillna(method='ffill')
temperature_data.loc[:,'MT_1':'MT_48']=temperature_data.loc[:,'MT_1':'MT_48'].fillna(method='bfill')
temperature_data['Mean_T']=temperature_data.loc[:,'MT_1':'MT_48'].mean(axis=1)

temperature_data['delT']=temperature_data.loc[:,'MT_1':'MT_48'].max(axis=1)-temperature_data.loc[:,'MT_1':'MT_48'].min(axis=1)

#%% Calculating Capacity_calculated and energy

details['Time_in_sec_s']=(details['DateTime']-details['DateTime'].shift(1))/np.timedelta64(1,'s')# time difference b/w every datapt
details['Time_in_sec']=(details['DateTime']-details['DateTime'].iloc[0])/np.timedelta64(1,'s')


#%%
details['Pack Current']=details['Pack Current'].astype(str).str.extract(r'(\-{0,}\d{1,}.\d{1,})')
details['Pack Current']=details['Pack Current'].astype(float)

#%%
details['State']=np.nan
details['Time_in_sec_s_cap']=np.where((details['Time_in_sec_s']>300),np.nan,details['Time_in_sec_s'])
details['State']=np.where(details['Pack Current']>0,0,details['State']) # 1
details['State']=np.where(details['Pack Current']<0,1,details['State']) # 2
details['State']=details['State'].fillna(method='ffill')
details['State']=details['State'].fillna(method='bfill')

# details['State']=details['State'].astype(str)

#%%
details['Cap_inst']=details['Time_in_sec_s_cap']*abs(details['Pack Current'])/3600
details['Capacity_calculated']=details['Cap_inst'].groupby(details['State']).cumsum()
details['Capacity_calculated_chg']=(details[details['State']==0]['Cap_inst']).cumsum()
details['Capacity_calculated_dchg']=(details[details['State']==1]['Cap_inst']).cumsum()

details['Capacity_calculated_chg']=details['Capacity_calculated_chg'].fillna(method='bfill')
details['Capacity_calculated_dchg']=details['Capacity_calculated_dchg'].fillna(method='bfill')

details['Energy_calculated']=details['Capacity_calculated']*details['Pack Voltage']
details['Energy_calculated_chg']=details['Capacity_calculated_chg']*details['Pack Voltage']
details['Energy_calculated_dchg']=details['Capacity_calculated_dchg']*details['Pack Voltage']

#%%
status_data['Main Relay']=status_data['Main Relay'].fillna(method='ffill')
status_data['Main Relay']=status_data['Main Relay'].astype(str)

#%% Add second table in document
document.add_paragraph().add_run('\nTest Specifications:').bold=True

table2=document.add_table(rows=4,cols=2)
table2.cell(0,0).text='Test Date and Time'
table2.cell(0,1).text=str(status_data['DateTime'].iloc[0])+' to '+str(status_data['DateTime'].iloc[-1])+' ('+ str((status_data['DateTime'].iloc[-1]-status_data['DateTime'].iloc[0])/np.timedelta64(1,'m')) +' minutes)'
table2.cell(1,0).text='Discharging Time Period'
table2.cell(2,0).text='Charging Time Period'
table2.cell(3,0).text='C-rate'

table2.style='Table Grid'

l=7
b=4
x_axis='DateTime'

#%%
status_data['Precharge Relay']=status_data['Precharge Relay'].replace('Open','OFF')
document.add_page_break()
document.add_paragraph().add_run('Precharge Relay:').bold=True

plt.figure(figsize=(l,b))

plt.plot(status_data[x_axis],status_data['Precharge Relay'])
plt.xlabel(x_axis,fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Precharge Relay',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Precharge Relay.png',dpi=1200)

document.add_picture(img_path+'\\Precharge Relay.png')
document.add_paragraph().add_run('As is known already, “open” means the precharge relay is off and “Close” means it is on.')

#%%
document.add_paragraph().add_run('Main Relay:').bold=True

plt.figure(figsize=(l,b))
plt.plot(status_data[x_axis],status_data['Main Relay'])
plt.xlabel(x_axis,fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Main Relay',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Main Relay.png',dpi=1200)

document.add_picture(img_path+'\\Main Relay.png')
document.add_paragraph('As is known already, “Close” means the main relay is on. ')

#%%
document.add_paragraph().add_run('Pack Status:').bold=True

plt.figure(figsize=(l,b))
plt.plot(status_data[x_axis],status_data['Pack Status'])
plt.xlabel(x_axis,fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Pack Status',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Pack Status.png',dpi=1200)

document.add_picture(img_path+'\\Pack Status.png')

#%%
document.add_paragraph().add_run('Pack Current:').bold=True
plt.figure(figsize=(l,b))
plt.plot(details[x_axis],details['Pack Current'])
plt.xlabel(x_axis,fontweight='bold')
plt.ylabel('Current(A)',fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('Pack Current',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Pack Current.png',dpi=1200)

document.add_picture(img_path+'\\Pack Current.png')

#%%
document.add_paragraph().add_run('Pack Capacity (calculated):').bold=True

plt.figure(figsize=(l,b))
plt.plot(details[x_axis],details['Capacity_calculated_dchg'])

plt.xlabel(x_axis,fontweight='bold')
plt.ylabel('Capacity(Ah)',fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('Pack Discharging Capacity (Calculated)',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Pack Discharging Capacity_calculated.png',dpi=1200)

document.add_picture(img_path+'\\Pack Discharging Capacity_calculated.png')
document.add_paragraph('In this discharging session, the capacity expended by the pack is about %sAh ()' %details['Capacity_calculated_dchg'].max())

#%%   

plt.figure(figsize=(l,b))
plt.plot(details[x_axis],details['Capacity_calculated_chg'])

plt.xlabel(x_axis,fontweight='bold')
plt.ylabel('Capacity(Ah)',fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('Pack Charging Capacity (Calculated)',fontweight='bold')
plt.tight_layout()

plt.savefig(img_path+'\\Pack Charging Capacity_calculated.png',dpi=1200)

document.add_picture(img_path+'\\Pack Charging Capacity_calculated.png')
document.add_paragraph('In this charging session, the capacity transferred to the pack is about %sAh (). ' %details['Capacity_calculated_chg'].max())
document.add_paragraph().add_run('Note: These values were calculated from the current (in amperes) values and time (in hours).').underline=True
#%%
document.add_paragraph().add_run('Pack Energy (calculated):').bold=True

plt.figure(figsize=(l,b))
plt.plot(details[x_axis],details['Energy_calculated_dchg']*0.001)

plt.xlabel(x_axis,fontweight='bold')
plt.ylabel('Energy(kWh)',fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('Pack Discharging Energy (Calculated)',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Pack Discharging Energy_calculated.png',dpi=1200)

document.add_picture(img_path+'\\Pack Discharging Energy_calculated.png')
document.add_paragraph('The Pack Energy expended in this discharging session (approx.  minutes) is about %skWh' % (details['Energy_calculated_dchg'].max()*0.001))

#%%
plt.figure(figsize=(l,b))
plt.plot(details[x_axis],details['Energy_calculated_chg']*0.001)

plt.xlabel(x_axis,fontweight='bold')
plt.ylabel('Energy(kWh)',fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('Pack Charging Energy (Calculated)',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Pack Charging Energy_calculated.png',dpi=1200)

document.add_picture(img_path+'\\Pack Charging Energy_calculated.png')
document.add_paragraph('The Pack Energy expended in this charging session () is about %skWh.' % (details['Energy_calculated_chg'].max()*0.001))
document.add_paragraph().add_run('Note: These values were calculated from the calculated capacity (in Ah) values and pack voltage (in volts).').underline=True

#%%
document.add_paragraph().add_run('Pack Energy (BMS):').bold=True

plt.figure(figsize=(l,b))
plt.plot(details[x_axis],details['Discharging Energy'])

plt.xlabel(x_axis,fontweight='bold')
plt.ylabel('Energy(Wh)',fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('Pack Discharging Energy (BMS)',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Pack Discharging Energy_BMS.png',dpi=1200)

document.add_picture(img_path+'\\Pack Discharging Energy_BMS.png')
p_bms_d_engy=document.add_paragraph('The ')
p_bms_d_engy.add_run('discharging pack energy ').bold=True
p_bms_d_engy.add_run('is about ')
p_bms_d_engy.add_run('%sWh.' %details['Discharging Energy'].max()).bold=True

#%%
plt.figure(figsize=(l,b))
plt.plot(details[x_axis],details['Charging Energy'])

plt.xlabel(x_axis,fontweight='bold')
plt.ylabel('Energy(Wh)',fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('Pack Charging Energy (BMS)',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Pack Charging Energy_BMS.png',dpi=1200)

document.add_picture(img_path+'\\Pack Charging Energy_BMS.png')
p_bms_c_engy=document.add_paragraph('The ')
p_bms_c_engy.add_run('charging pack energy ').bold=True
p_bms_c_engy.add_run('is about ')
p_bms_c_engy.add_run('%sWh.' %details['Charging Energy'].max()).bold=True


#%%
document.add_paragraph().add_run('Pack Voltage:').bold=True
plt.figure(figsize=(l,b))
plt.plot(details[x_axis],details['Pack Voltage'])
plt.xlabel(x_axis,fontweight='bold')
plt.ylabel('Voltage(V)',fontweight='bold')

plt.grid(linestyle='dotted')
plt.title('Pack Voltage',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Pack Voltage.png',dpi=1200)

document.add_picture(img_path+'\\Pack Voltage.png')
pv=document.add_paragraph('During ')
pv.add_run('discharge, ').bold=True
pv.add_run('the pack goes from to . The pack voltage at rest after discharge is . At ')
pv.add_run('charge, ').bold=True
pv.add_run('the pack goes from to . ')
pv.add_run('In the end, the voltage stabilizes at %sV during the end rest.' %details['Pack Voltage'].iloc[-1])

#%%
document.add_paragraph().add_run('Cell Voltage:').bold=True

plt.figure(figsize=(l,b))
plt.plot(voltage_data[x_axis],voltage_data.loc[:,'MV_1':'MV_168'])
plt.ylabel('Voltage(V)',fontweight='bold')
plt.xlabel(x_axis,fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Cell Voltage',fontweight='bold')
plt.tight_layout()


#%%
document.add_paragraph().add_run('Average Voltage:').bold=True
#%%
plt.figure(figsize=(l,b))
plt.plot(voltage_data[x_axis],(voltage_data['Mean_V'])) #*0.0001
plt.ylabel('Voltage(V)',fontweight='bold')
plt.xlabel(x_axis,fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Average Voltage',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Average Voltage.png',dpi=1200)

#%%
document.add_picture(img_path+'\\Average Voltage.png')
av=document.add_paragraph('The average cell voltage during ')
av.add_run('discharge ').bold=True
av.add_run('goes from ')
av.add_run('V to V. ').bold=True
av.add_run('The average cell voltage at the rest after discharge is V. At ')
av.add_run('charge').bold=True
av.add_run(', the average cell voltage goes from ')
av.add_run('V to V.').bold=True
av.add_run('In the end rest, it  stabilizes at %sV.' % (voltage_data['Mean_V'].iloc[-1]*0.0001))

#%%
document.add_paragraph().add_run('Voltage Difference (delV):').bold=True
plt.figure(figsize=(l,b))
plt.plot(voltage_data[x_axis],(voltage_data['delV']))
plt.ylabel('Voltage(mV)',fontweight='bold')
plt.xlabel(x_axis,fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Voltage difference (delV)',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\delV.png',dpi=1200)

document.add_picture(img_path+'\\delV.png')
dv=document.add_paragraph('During discharge, the ')
dv.add_run('maximum delV (difference between the highest and lowest cell voltage), ').bold=True
dv.add_run('is about ')
dv.add_run('mV, at the end of discharge. ').bold=True
dv.add_run('At the rest after discharging, the delmV is mV. In the charging session, the maximum delV is mV. ')
dv.add_run('The pack stabilizes at %smV delV at the end rest. '% (voltage_data['delV'].iloc[-1]))

#%%
document.add_paragraph().add_run('Cell Temperature:').bold=True

ct=document.add_paragraph('The maximum temperature reached is around ')
ct.add_run('%s'%(temperature_data.loc[:,'MT_1':'MT_48'].max(axis=1).max())).bold=True
ct.add_run(u'\N{DEGREE SIGN}'+'C.' ).bold=True

plt.figure(figsize=(l,b))
plt.plot(temperature_data[x_axis],temperature_data.loc[:,'MT_1':'MT_48'])
plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
plt.xlabel('DateTime', fontweight = 'bold')
plt.grid(linestyle='dotted')
plt.title('Cell Temperature',fontweight='bold')
plt.tight_layout()

#%%
document.add_paragraph().add_run('Average temperature:').bold=True
plt.figure(figsize=(l,b))
plt.plot(temperature_data[x_axis],temperature_data['Mean_T'])

plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
plt.xlabel(x_axis,fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Average Temperature',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\Average Temperature.png',dpi=1200)

document.add_picture(img_path+'\\Average Temperature.png')
at=document.add_paragraph('As can be seen, the ')
at.add_run('average cell temperature ').bold=True
at.add_run('stays between ')
at.add_run(u'\N{DEGREE SIGN}'+'C to '+u'\N{DEGREE SIGN}'+'C ').bold=True
at.add_run('during ')
at.add_run('discharging').bold=True
at.add_run(' and ')
at.add_run(u'\N{DEGREE SIGN}'+'C to '+u'\N{DEGREE SIGN}'+'C ').bold=True
at.add_run('during ')
at.add_run('charging.').bold=True

#%%
document.add_paragraph().add_run('Temperature Difference (delT):').bold=True
plt.figure(figsize=(l,b))
plt.plot(temperature_data[x_axis],temperature_data['delT'])
plt.ylabel('Temperature('+u'\N{DEGREE SIGN}'+'C)',fontweight='bold')
plt.xlabel(x_axis,fontweight='bold')
plt.grid(linestyle='dotted')
plt.title('Temperature difference (delT)',fontweight='bold')
plt.tight_layout()
plt.savefig(img_path+'\\delT.png',dpi=1200)

document.add_picture(img_path+'\\delT.png')
dt=document.add_paragraph('The ')
dt.add_run('maximum delT ').bold=True
dt.add_run('is about ')
dt.add_run('%s' %temperature_data['delT'].max()).bold=True
dt.add_run(u'\N{DEGREE SIGN}'+'C.').bold=True

#%% Change page margins
from docx.shared import Inches
pageSections = document.sections

for section in pageSections:
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

#%%
voltage_data = voltage_data.drop(columns = ['Date','Time'])
temperature_data = temperature_data.drop(columns = ['Date', 'Time'])
details = details.drop(columns = ['Date', 'Time'])
status_data = status_data.drop(columns = ['Date', 'Time'])
max_min_data = max_min_data.drop(columns = ['Date', 'Time'])

#%%
path=file.rsplit('.',1)[0] +'_modified'+ tag + '.xlsx'
document.save(img_path+'\\'+'LTO_15kW_'+file.rsplit('\\',1)[1].rsplit('.',1)[0]+'_observations'+ tag + '.docx')

#%%
writer = pd.ExcelWriter(path, engine = 'xlsxwriter')
voltage_data.to_excel(writer, sheet_name = 'Cell Voltage',index=False)
temperature_data.to_excel(writer, sheet_name = 'Cell Temperature',index=False)
details.to_excel(writer,sheet_name='Pack Details',index=False)
status_data.to_excel(writer,sheet_name='Pack Status',index=False)
max_min_data.to_excel(writer,sheet_name='Max Min Mean',index=False)
writer.close()
